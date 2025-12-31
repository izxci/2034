import streamlit as st
import zipfile
import xml.etree.ElementTree as ET
import re
from pypdf import PdfReader
from io import BytesIO
import google.generativeai as genai
import importlib.metadata
from docx import Document
from fpdf import FPDF
import urllib.parse
from gtts import gTTS
import speech_recognition as sr
import os
from PIL import Image
import time
import pandas as pd
from datetime import datetime, timedelta, date
import shutil
import json

# --- Sayfa Ayarları ---
st.set_page_config(
    page_title="Hukuk Asistanı AI",
    page_icon="⚖️",
    layout="wide"
)

# --- CSS ---
st.markdown("""
    <style>
    .main { background-color: #f8f9fa; }
    .kanun-kutusu { 
        background-color: #fff3e0; 
        padding: 15px; 
        border-left: 5px solid #ff9800; 
        border-radius: 5px; 
        margin-bottom: 10px;
        white-space: pre-wrap;
    }
    .ictihat-kutusu {
        background-color: #e3f2fd;
        padding: 15px;
        border-left: 5px solid #2196f3;
        border-radius: 5px;
        margin-bottom: 10px;
    }
    .buyur-abi-kutusu {
        background-color: #f3e5f5;
        padding: 15px;
        border-left: 5px solid #9c27b0;
        border-radius: 5px;
        margin-bottom: 10px;
    }
    .alarm-kutusu {
        background-color: #ffebee;
        padding: 15px;
        border-left: 5px solid #f44336;
        border-radius: 5px;
        margin-bottom: 10px;
        font-weight: bold;
        color: #b71c1c;
    }
    .arsiv-kutusu {
        background-color: #e0f2f1;
        padding: 15px;
        border-left: 5px solid #009688;
        border-radius: 5px;
        margin-bottom: 10px;
    }
    .uyap-kutusu {
        background-color: #fce4ec; 
        padding: 15px; 
        border-left: 5px solid #c2185b; 
        border-radius: 5px; 
        margin-bottom: 20px;
    }
    .ozel-sekme {
        border: 1px solid #ddd;
        padding: 20px;
        border-radius: 10px;
        background-color: #ffffff;
    }
    </style>
    """, unsafe_allow_html=True)

# --- KALICILIK (VERİ TABANI) FONKSİYONLARI ---
DURUSMA_FILE = "durusma_kayitlari.json"

def save_durusma_data(data):
    """Duruşma listesini JSON dosyasına kaydeder."""
    serializable_data = []
    for item in data:
        temp = item.copy()
        if isinstance(temp.get('dtstart'), datetime):
            temp['dtstart'] = temp['dtstart'].isoformat()
        serializable_data.append(temp)
    
    try:
        with open(DURUSMA_FILE, "w", encoding="utf-8") as f:
            json.dump(serializable_data, f, ensure_ascii=False, indent=4)
    except Exception as e:
        st.error(f"Kaydetme hatası: {e}")

def load_durusma_data():
    """JSON dosyasından duruşma listesini yükler."""
    if not os.path.exists(DURUSMA_FILE):
        return []
    try:
        with open(DURUSMA_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        for item in data:
            if 'dtstart' in item and item['dtstart']:
                item['dtstart'] = datetime.fromisoformat(item['dtstart'])
        return data
    except:
        return []

# --- YARDIMCI FONKSİYONLAR ---
def parse_udf(file_bytes):
    try:
        with zipfile.ZipFile(file_bytes) as z:
            if 'content.xml' in z.namelist():
                with z.open('content.xml') as f:
                    tree = ET.parse(f)
                    root = tree.getroot()
                    text_content = [elem.text.strip() for elem in root.iter() if elem.text]
                    return " ".join(text_content)
            return "HATA: UDF içeriği okunamadı."
    except Exception as e:
        return f"HATA: {str(e)}"

def parse_pdf(file_bytes):
    try:
        reader = PdfReader(file_bytes)
        text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        if len(text.strip()) < 50: return "" 
        return text
    except Exception as e:
        return ""

def extract_metadata(text):
    if not isinstance(text, str) or text.startswith(("HATA", "UYARI")):
        return {"mahkeme": "-", "esas": "-", "karar": "-", "tarih": "-"}
    
    esas = re.search(r"(?i)Esas\s*No\s*[:\-]?\s*(\d{4}/\d+)", text)
    karar = re.search(r"(?i)Karar\s*No\s*[:\-]?\s*(\d{4}/\d+)", text)
    tarih = re.search(r"(\d{1,2}[./]\d{1,2}[./]\d{4})", text)
    
    mahkeme = "Tespit Edilemedi"
    for line in text.split('\n')[:40]:
        clean = line.strip()
        if ("MAHKEMESİ" in clean.upper() or "DAİRESİ" in clean.upper()) and len(clean) > 5:
            mahkeme = clean
            break
    return {
        "mahkeme": mahkeme,
        "esas": esas.group(1) if esas else "Bulunamadı",
        "karar": karar.group(1) if karar else "Bulunamadı",
        "tarih": tarih.group(1) if tarih else "Bulunamadı"
    }

def parse_ics_data(file_bytes):
    events = []
    try:
        content = file_bytes.getvalue().decode('utf-8')
        lines = content.splitlines()
        current_event = {}
        in_event = False
        
        for line in lines:
            line = line.strip()
            if line == 'BEGIN:VEVENT':
                in_event = True
                current_event = {}
            elif line == 'END:VEVENT':
                in_event = False
                if 'dtstart' in current_event and 'summary' in current_event:
                    events.append(current_event)
            elif in_event:
                if line.startswith('SUMMARY:'):
                    current_event['summary'] = line.split(':', 1)[1]
                elif line.startswith('DTSTART:'):
                    raw_date = line.split(':', 1)[1]
                    try:
                        dt = datetime.strptime(raw_date, '%Y%m%dT%H%M%S')
                        current_event['dtstart'] = dt
                    except: pass
                elif line.startswith('LOCATION:'):
                    current_event['location'] = line.split(':', 1)[1]
                elif line.startswith('DESCRIPTION:'):
                    current_event['description'] = line.split(':', 1)[1]
        return events
    except Exception as e:
        return []

# --- DOSYA OLUŞTURMA FONKSİYONLARI ---
def create_word_file(text):
    doc = Document()
    for line in text.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io

def create_udf_file(text):
    root = ET.Element("content")
    body = ET.SubElement(root, "body")
    for line in text.split('\n'):
        p = ET.SubElement(body, "p")
        p.text = line
    xml_str = ET.tostring(root, encoding='utf-8', method='xml')
    byte_io = BytesIO()
    with zipfile.ZipFile(byte_io, 'w', zipfile.ZIP_DEFLATED) as z:
        z.writestr('content.xml', xml_str)
    byte_io.seek(0)
    return byte_io

def create_pdf_file(text):
    replacements = {
        'ğ': 'g', 'Ğ': 'G', 'ş': 's', 'Ş': 'S', 'ı': 'i', 'İ': 'I',
        'ç': 'c', 'Ç': 'C', 'ü': 'u', 'Ü': 'U', 'ö': 'o', 'Ö': 'O',
        '“': '"', '”': '"', '’': "'", '–': '-', '…': '...'
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    
    text = text.encode('latin-1', 'replace').decode('latin-1')
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="Hukuki Analiz Raporu", ln=1, align='C')
    pdf.ln(10)
    pdf.set_font("Arial", size=11)
    pdf.multi_cell(0, 10, text)
    return pdf.output(dest='S').encode('latin-1', 'ignore')

# --- SES İŞLEME FONKSİYONLARI ---
def text_to_speech(text):
    try:
        tts = gTTS(text=text, lang='tr')
        fp = BytesIO()
        tts.write_to_fp(fp)
        fp.seek(0)
        return fp
    except Exception as e:
        return None

def speech_to_text(audio_bytes):
    r = sr.Recognizer()
    try:
        with sr.AudioFile(audio_bytes) as source:
            audio_data = r.record(source)
            text = r.recognize_google(audio_data, language='tr-TR')
            return text
    except Exception as e:
        return f"Hata: {str(e)}"

# --- MULTIMODAL VE OCR FONKSİYONLARI ---
def perform_ocr_gemini(file_bytes, mime_type, api_key, prompt_text="Bu dosyanın içeriğini tam olarak metne dök."):
    if not api_key: return "API Key Yok"
    genai.configure(api_key=api_key)
    
    if mime_type in ['image/tiff', 'image/tif']:
        try:
            image = Image.open(file_bytes)
            rgb_im = image.convert('RGB')
            buf = BytesIO()
            rgb_im.save(buf, format="JPEG")
            file_bytes = buf
            mime_type = 'image/jpeg'
        except Exception as e:
            return f"TIFF Dönüştürme Hatası: {str(e)}"

    available_models = []
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                if '1.5' in m.name or 'vision' in m.name:
                    available_models.append(m.name)
    except: pass
    
    if not available_models: available_models = ['models/gemini-1.5-flash', 'models/gemini-1.5-pro']

    image_part = {"mime_type": mime_type, "data": file_bytes.getvalue()}
    
    for model_name in available_models:
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content([prompt_text, image_part])
            return response.text
        except: continue
    return "Analiz Başarısız."

def extract_text_from_docx(file_bytes):
    try:
        doc = Document(file_bytes)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        return f"Word Okuma Hatası: {str(e)}"

def read_excel_file(file_bytes):
    try:
        df = pd.read_excel(file_bytes)
        return df.to_string()
    except Exception as e:
        return f"Excel Hatası: {str(e)}"

# --- AKILLI AI MOTORU ---
def get_ai_response(prompt, api_key):
    if not api_key: return "Lütfen API Anahtarı giriniz."
    genai.configure(api_key=api_key)
    candidate_models = ['gemini-1.5-flash', 'gemini-1.5-pro', 'gemini-1.0-pro', 'gemini-pro']
    
    for model_name in candidate_models:
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(prompt)
            return response.text 
        except: continue 
    
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                try:
                    model = genai.GenerativeModel(m.name)
                    response = model.generate_content(prompt)
                    return response.text
                except: continue
    except: pass
    return "Hata: AI yanıt veremedi."

# --- ANA UYGULAMA ---
def main():
    st.title("⚖️ Hukuk Asistanı (v9.0 - Ultra Full Sürüm)")
    
    try:
        lib_ver = importlib.metadata.version("google-generativeai")
    except:
        lib_ver = "Bilinmiyor"

    # --- BAŞLANGIÇTA VERİLERİ YÜKLE ---
    if "durusma_listesi" not in st.session_state:
        st.session_state.durusma_listesi = load_durusma_data()

    if "doc_text" not in st.session_state: st.session_state.doc_text = ""
    if "last_file_id" not in st.session_state: st.session_state.last_file_id = None
    if "messages" not in st.session_state: st.session_state.messages = []
    if "mevzuat_sonuc" not in st.session_state: st.session_state.mevzuat_sonuc = ""
    if "ictihat_sonuc" not in st.session_state: st.session_state.ictihat_sonuc = ""
    if "dilekce_taslak" not in st.session_state: st.session_state.dilekce_taslak = ""
    if "soru_cevap" not in st.session_state: st.session_state.soru_cevap = ""
    if "ses_metni" not in st.session_state: st.session_state.ses_metni = ""
    if "ocr_metni" not in st.session_state: st.session_state.ocr_metni = ""
    if "dalgic_context" not in st.session_state: st.session_state.dalgic_context = ""
    if "dalgic_sonuc" not in st.session_state: st.session_state.dalgic_sonuc = ""
    if "buyur_abi_context" not in st.session_state: st.session_state.buyur_abi_context = ""
    if "buyur_abi_response" not in st.session_state: st.session_state.buyur_abi_response = ""
    
    if "arsiv_context" not in st.session_state: st.session_state.arsiv_context = ""
    if "arsiv_genel_ozet" not in st.session_state: st.session_state.arsiv_genel_ozet = ""
    if "arsiv_soru_cevap" not in st.session_state: st.session_state.arsiv_soru_cevap = ""
    if "arsiv_arama_sonuclari" not in st.session_state: st.session_state.arsiv_arama_sonuclari = []
    if "aktif_dosya_adi" not in st.session_state: st.session_state.aktif_dosya_adi = ""
    if "aktif_dosya_yolu" not in st.session_state: st.session_state.aktif_dosya_yolu = ""
    
    # Yeni State'ler
    if "faiz_sonuc" not in st.session_state: st.session_state.faiz_sonuc = ""
    if "sozlesme_analiz" not in st.session_state: st.session_state.sozlesme_analiz = ""
    if "muvekkil_mesaj" not in st.session_state: st.session_state.muvekkil_mesaj = ""

    ROOT_DIR = "Hukuk_Arsivi"
    if not os.path.exists(ROOT_DIR):
        os.makedirs(ROOT_DIR)

    with st.sidebar:
        st.header("⚙️ Ayarlar")
        api_key = st.text_input("Google Gemini API Key", type="password")
        st.caption(f"Kütüphane Sürümü: {lib_ver}")
        
        st.divider()
        st.header("📁 Dosya Bilgileri")
        input_davaci = st.text_input("Davacı")
        input_davali = st.text_input("Davalı")
        input_mahkeme = st.text_input("Mahkeme")
        input_dosya_no = st.text_input("Dosya No")
        
        if st.button("🗑️ Ekranı Temizle"):
            for key in st.session_state.keys():
                if key != "durusma_listesi":
                    del st.session_state[key]
            st.rerun()

    uploaded_file = st.file_uploader("Dosya Yükle (UDF/PDF)", type=['udf', 'pdf'])

    if uploaded_file and st.session_state.get('last_file_id') != uploaded_file.file_id:
        with st.spinner("Okunuyor..."):
            file_bytes = BytesIO(uploaded_file.getvalue())
            ext = uploaded_file.name.split('.')[-1].lower()
            raw_text = parse_udf(file_bytes) if ext == 'udf' else parse_pdf(file_bytes)
            st.session_state.doc_text = raw_text
            st.session_state.last_file_id = uploaded_file.file_id
            st.session_state.messages = []

    if st.session_state.doc_text.startswith(("HATA", "UYARI")):
        st.warning(st.session_state.doc_text)
    
    auto_data = extract_metadata(st.session_state.doc_text)

    # --- SEKMELER (YENİ ÖZELLİKLER EKLENDİ) ---
    tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9, tab10, tab11, tab12, tab13, tab14, tab15, tab16, tab17, tab18 = st.tabs([
        "📋 Analiz", "💬 Sohbet", "📕 Mevzuat", "⚖️ İçtihat", 
        "✍️ Dilekçe Yaz", "❓ Bana Sor", "🎙️ Sesli Komut", "👁️ OCR", 
        "🤿 Dalgıç", "🙋 Buyur Abi", "⏰ Hatırlatıcı", "🗄️ Arşiv", 
        "🏛️ UYAP Analiz", "🧮 Faiz Hesabı", "⏳ Süre Hesapla", 
        "🕸️ İlişki Ağı", "📝 Sözleşme Analiz", "📧 Müvekkil Bilgi"
    ])

    with tab1:
        col1, col2 = st.columns(2)
        with col1:
            st.write(f"**Mahkeme:** {input_mahkeme or auto_data['mahkeme']}")
            st.write(f"**Dosya No:** {input_dosya_no or auto_data['esas']}")
        with col2:
            st.write(f"**Davacı:** {input_davaci or '-'}")
            st.write(f"**Davalı:** {input_davali or '-'}")
        st.text_area("Metin Önizleme", st.session_state.doc_text, height=150)

    with tab2:
        for msg in st.session_state.messages:
            with st.chat_message(msg["role"]): st.markdown(msg["content"])
        if prompt := st.chat_input("Soru sor..."):
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"): st.markdown(prompt)
            with st.chat_message("assistant"):
                with st.spinner("AI Yanıtlıyor..."):
                    context = f"BELGE: {st.session_state.doc_text[:20000]}\nSORU: {prompt}"
                    reply = get_ai_response(f"Sen bir avukatsın. Şuna cevap ver: {context}", api_key)
                    st.markdown(reply)
                    st.session_state.messages.append({"role": "assistant", "content": reply})

    with tab3:
        c1, c2 = st.columns([3,1])
        q = c1.text_input("Kanun Madde No", key="mq")
        if c2.button("Getir", key="mb") and q:
            with st.spinner("Aranıyor..."):
                res = get_ai_response(f"GÖREV: '{q}' maddesini tam metin yaz.", api_key)
                st.session_state.mevzuat_sonuc = res
        if st.session_state.mevzuat_sonuc:
            st.markdown(f"<div class='kanun-kutusu'>{st.session_state.mevzuat_sonuc}</div>", unsafe_allow_html=True)

    with tab4:
        c3, c4 = st.columns([3,1])
        iq = c3.text_input("İçtihat Konusu", key="iq")
        if c4.button("Ara", key="ib") and iq:
            with st.spinner("Taranıyor..."):
                res = get_ai_response(f"GÖREV: '{iq}' hakkında Yargıtay kararlarını özetle.", api_key)
                st.session_state.ictihat_sonuc = res
        if st.session_state.ictihat_sonuc:
            st.markdown(f"<div class='ictihat-kutusu'>{st.session_state.ictihat_sonuc}</div>", unsafe_allow_html=True)

    with tab5:
        st.subheader("✍️ Otomatik Savunma/Cevap Dilekçesi")
        if not st.session_state.doc_text or st.session_state.doc_text.startswith(("HATA", "UYARI")):
            st.info("Dilekçe oluşturmak için önce sol menüden bir dosya yükleyin.")
        else:
            col_d1, col_d2 = st.columns([2, 1])
            with col_d1:
                dilekce_turu = st.selectbox("Dilekçe Türü", ["Cevap Dilekçesi", "İtiraz Dilekçesi", "Beyan Dilekçesi"])
                ozel_talimat = st.text_area("Özel Savunma Stratejisi (Opsiyonel)", placeholder="Örn: Zamanaşımı itirazında bulun...")
            with col_d2:
                st.write("")
                st.write("")
                if st.button("Dilekçeyi Yaz (AI)", type="primary"):
                    if not api_key: st.error("API Key gerekli!")
                    else:
                        with st.spinner("Dilekçe yazılıyor..."):
                            mahkeme = input_mahkeme or auto_data['mahkeme']
                            dosya = input_dosya_no or auto_data['esas']
                            davaci = input_davaci or "Davacı"
                            davali = input_davali or "Davalı"
                            prompt = f"""
                            GÖREV: Aşağıdaki metne dayanarak profesyonel bir {dilekce_turu} yaz.
                            BİLGİLER: Mahkeme: {mahkeme}, Dosya: {dosya}, Davacı: {davaci}, Davalı: {davali}, Ek Talimat: {ozel_talimat}
                            KARŞI TARAFIN DİLEKÇESİ (ÖZET): {st.session_state.doc_text[:20000]}
                            KURALLAR: Resmi Türk hukuk dilekçesi formatında olsun.
                            """
                            res = get_ai_response(prompt, api_key)
                            st.session_state.dilekce_taslak = res
            if st.session_state.dilekce_taslak:
                st.divider()
                st.subheader("📄 Dilekçe Taslağı")
                btn_col1, btn_col2 = st.columns(2)
                word_file = create_word_file(st.session_state.dilekce_taslak)
                with btn_col1:
                    st.download_button("💾 Word Olarak İndir (.docx)", word_file, "Dilekce.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                udf_file = create_udf_file(st.session_state.dilekce_taslak)
                with btn_col2:
                    st.download_button("💾 UDF Olarak İndir (.udf)", udf_file, "Dilekce.udf", "application/zip")
                st.text_area("Dilekçe Metni", st.session_state.dilekce_taslak, height=500)

    with tab6:
        st.subheader("❓ Hukuki Soru & WhatsApp Paylaşımı")
        col_s1, col_s2 = st.columns([3, 1])
        with col_s1:
            kullanici_sorusu = st.text_area("Hukuki Sorunuzu Yazın", height=100, placeholder="Örn: Kiracı kirayı ödemezse tahliye süreci nasıl işler?")
        with col_s2:
            telefon_no = st.text_input("WhatsApp No (905xxxxxxxxx)", placeholder="905551234567")
            if st.button("Analiz Et ve Hazırla", type="primary"):
                if not api_key: st.error("API Key giriniz.")
                elif not kullanici_sorusu: st.warning("Lütfen bir soru yazın.")
                else:
                    with st.spinner("Mevzuat ve İçtihatlar taranıyor..."):
                        prompt = f"""
                        GÖREV: Aşağıdaki hukuki soruyu detaylıca cevapla.
                        SORU: {kullanici_sorusu}
                        KURALLAR: 1. İlgili KANUN MADDELERİNİ belirt. 2. YARGITAY İÇTİHATLARINDAN örnek ver. 3. Net hukuki görüş bildir.
                        """
                        res = get_ai_response(prompt, api_key)
                        st.session_state.soru_cevap = res
        if st.session_state.soru_cevap:
            st.divider()
            st.markdown(f"<div class='ictihat-kutusu'><b>💡 Hukuki Görüş:</b><br>{st.session_state.soru_cevap}</div>", unsafe_allow_html=True)
            pdf_data = create_pdf_file(st.session_state.soru_cevap)
            encoded_text = urllib.parse.quote(f"*Hukuki Soru:* {kullanici_sorusu}\n\n*Cevap:*\n{st.session_state.soru_cevap}")
            wa_link = f"https://wa.me/{telefon_no}?text={encoded_text}"
            col_btn1, col_btn2 = st.columns(2)
            with col_btn1: st.download_button("📄 Cevabı PDF Olarak İndir", pdf_data, "Hukuki_Gorus.pdf", "application/pdf")
            with col_btn2:
                if telefon_no: st.link_button("📲 Cevabı WhatsApp ile Gönder", wa_link)
                else: st.warning("WhatsApp butonu için telefon no giriniz.")

    with tab7:
        st.subheader("🎙️ Sesli Asistan")
        col_audio1, col_audio2 = st.columns(2)
        with col_audio1:
            st.markdown("##### 🗣️ Metni Seslendir")
            text_to_read = st.text_area("Okunacak Metni Yazın:", height=150)
            if st.button("🔊 Seslendir"):
                if text_to_read:
                    with st.spinner("Ses oluşturuluyor..."):
                        audio_fp = text_to_speech(text_to_read)
                        if audio_fp: st.audio(audio_fp, format='audio/mp3')
                        else: st.error("Hata oluştu.")
        with col_audio2:
            st.markdown("##### 📝 Sesi Yazıya Çevir")
            audio_input = st.file_uploader("Ses Dosyası (WAV/MP3)", type=["wav", "mp3"])
            if audio_input and st.button("📝 Yazıya Dök"):
                with st.spinner("Analiz ediliyor..."):
                    text_result = speech_to_text(audio_input)
                    st.session_state.ses_metni = text_result
            if st.session_state.ses_metni:
                st.success("Sonuç:")
                st.text_area("", st.session_state.ses_metni, height=150)

    with tab8:
        st.subheader("👁️ OCR (Resim/PDF -> Metin)")
        ocr_file = st.file_uploader("Dosya Yükle", type=['png', 'jpg', 'jpeg', 'pdf', 'docx', 'tif', 'tiff'])
        if ocr_file and st.button("🔍 Metni Ayıkla (OCR)", type="primary"):
            if not api_key: st.error("API Key gerekli.")
            else:
                with st.spinner("İşleniyor..."):
                    ftype = ocr_file.name.split('.')[-1].lower()
                    if ftype == 'docx': res = extract_text_from_docx(ocr_file)
                    else:
                        mime = "application/pdf" if ftype == 'pdf' else "image/tiff" if ftype in ['tif', 'tiff'] else "image/jpeg"
                        ocr_file.seek(0)
                        res = perform_ocr_gemini(ocr_file, mime, api_key)
                    st.session_state.ocr_metni = res
        if st.session_state.ocr_metni:
            st.text_area("OCR Sonucu:", st.session_state.ocr_metni, height=400)
            word_ocr = create_word_file(st.session_state.ocr_metni)
            st.download_button("💾 Word İndir", word_ocr, "ocr.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    with tab9:
        st.subheader("🤿 Dalgıç Modu (Çoklu Dosya Analizi)")
        st.info("Birden fazla dosyayı aynı anda yükleyin. Sistem hepsini okuyup, birleştirip sorularınızı yanıtlar.")
        dalgic_files = st.file_uploader("Dosyaları Sürükleyin (Max 30 Dosya)", type=['udf', 'pdf', 'docx', 'doc', 'txt', 'png', 'jpg', 'jpeg', 'mp4', 'tif', 'tiff'], accept_multiple_files=True)
        if dalgic_files:
            if st.button("🚀 Dosyaları İşle ve Hafızaya Al", type="primary"):
                if not api_key: st.error("API Key giriniz.")
                else:
                    full_context = ""
                    progress_bar = st.progress(0)
                    for i, file in enumerate(dalgic_files):
                        file_bytes = BytesIO(file.read())
                        ext = file.name.split('.')[-1].lower()
                        extracted_text = ""
                        try:
                            if ext == 'udf': 
                                extracted_text = parse_udf(file_bytes)
                            elif ext == 'txt': 
                                extracted_text = file_bytes.read().decode('utf-8', errors='ignore')
                            elif ext in ['docx', 'doc']: 
                                extracted_text = extract_text_from_docx(file_bytes)
                            elif ext == 'pdf':
                                extracted_text = parse_pdf(file_bytes)
                                if not extracted_text:
                                    file_bytes.seek(0)
                                    extracted_text = perform_ocr_gemini(file_bytes, "application/pdf", api_key)
                            elif ext in ['png', 'jpg', 'jpeg', 'img', 'tif', 'tiff']:
                                mime = "image/tiff" if ext in ['tif', 'tiff'] else "image/jpeg"
                                extracted_text = perform_ocr_gemini(file_bytes, mime, api_key)
                            elif ext == 'mp4':
                                extracted_text = perform_ocr_gemini(file_bytes, "video/mp4", api_key, "Video içeriğini özetle.")
                            
                            full_context += f"\n\n--- DOSYA: {file.name} ---\n{extracted_text}"
                        except Exception as e:
                            full_context += f"\nHATA ({file.name}): {str(e)}"
                        
                        progress_bar.progress((i + 1) / len(dalgic_files))
                    st.session_state.dalgic_context = full_context
                    st.success(f"Veriler hafızaya alındı! ({len(full_context)} karakter)")
        if st.session_state.dalgic_context:
            st.divider()
            dalgic_soru = st.text_area("Dosyalar Hakkında Soru Sorun:", placeholder="Örn: Bu dosyalardaki tüm tanık ifadelerindeki çelişkileri listele.")
            if st.button("Analiz Et ve Yanıtla"):
                if not dalgic_soru: st.warning("Soru yazın.")
                else:
                    with st.spinner("Dalgıç derinlere iniyor..."):
                        prompt = f"GÖREV: Aşağıdaki dosya içeriklerine göre cevapla.\nSORU: {dalgic_soru}\nİÇERİK: {st.session_state.dalgic_context[:500000]}"
                        res = get_ai_response(prompt, api_key)
                        st.session_state.dalgic_sonuc = res
            if st.session_state.dalgic_sonuc:
                st.markdown(f"<div class='kanun-kutusu'>{st.session_state.dalgic_sonuc}</div>", unsafe_allow_html=True)
                col_d1, col_d2 = st.columns(2)
                with col_d1: st.download_button("📕 PDF İndir", create_pdf_file(st.session_state.dalgic_sonuc), "Dalgic.pdf", "application/pdf")
                with col_d2: st.download_button("📘 Word İndir", create_word_file(st.session_state.dalgic_sonuc), "Dalgic.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    with tab10:
        st.subheader("🙋 Buyur Abi (Genel Asistan & Çoklu Format)")
        st.info("Hukuk, kodlama, yemek tarifi veya günlük sohbet... Ne istersen sor. Ayrıca Excel, Ses, Video dahil her türlü dosyayı yükleyip analiz ettirebilirsin.")
        col_ba1, col_ba2 = st.columns([1, 2])
        with col_ba1:
            st.markdown("#### 📁 Dosya & Ses Girişi")
            buyur_files = st.file_uploader("Dosya Ekle (Excel, Ses, Video, Resim vb.)", 
                                           type=['pdf','udf','doc','docx','txt','xls','xlsx','xlt','xml','jpg','png','jpeg','mp3','mp4','wav','tif','tiff'],
                                           accept_multiple_files=True)
            st.markdown("#### 🎙️ Sesli Soru Sor")
            audio_prompt = st.file_uploader("Ses Kaydı Yükle (Soru olarak)", type=['wav', 'mp3', 'ogg'], key="voice_prompt")
        with col_ba2:
            st.markdown("#### 💬 Sohbet Alanı")
            user_text_input = st.text_area("Sorunu Yaz Abi:", height=150, placeholder="Örn: Bu Excel dosyasındaki ciroları topla veya yüklediğim ses kaydını özetle...")
            if st.button("🚀 Gönder Gelsin", type="primary"):
                if not api_key: st.error("Önce sol menüden API Anahtarını girmen lazım abi.")
                else:
                    context_data = ""
                    voice_text = ""
                    with st.spinner("Dosyalar ve sesler inceleniyor..."):
                        if audio_prompt:
                            voice_text = speech_to_text(audio_prompt)
                            st.info(f"🎤 Sesli Sorun: {voice_text}")
                        if buyur_files:
                            for file in buyur_files:
                                f_bytes = BytesIO(file.read())
                                ext = file.name.split('.')[-1].lower()
                                try:
                                    if ext in ['xls', 'xlsx', 'xlt']: context_data += f"\n--- EXCEL ({file.name}) ---\n{read_excel_file(f_bytes)}"
                                    elif ext in ['txt', 'xml', 'py', 'js', 'html']: context_data += f"\n--- METİN ({file.name}) ---\n{f_bytes.read().decode('utf-8', errors='ignore')}"
                                    elif ext in ['doc', 'docx']: context_data += f"\n--- WORD ({file.name}) ---\n{extract_text_from_docx(f_bytes)}"
                                    elif ext == 'pdf':
                                        pdf_txt = parse_pdf(f_bytes)
                                        if not pdf_txt:
                                            f_bytes.seek(0)
                                            pdf_txt = perform_ocr_gemini(f_bytes, "application/pdf", api_key)
                                        context_data += f"\n--- PDF ({file.name}) ---\n{pdf_txt}"
                                    elif ext in ['jpg', 'png', 'jpeg', 'img', 'tif', 'tiff']:
                                        mime = "image/tiff" if ext in ['tif', 'tiff'] else "image/jpeg"
                                        ocr_res = perform_ocr_gemini(f_bytes, mime, api_key, "Bu resimde ne var?")
                                        context_data += f"\n--- RESİM ({file.name}) ---\n{ocr_res}"
                                    elif ext in ['mp3', 'wav', 'mp4']:
                                        mime = "video/mp4" if ext == 'mp4' else "audio/mp3"
                                        media_res = perform_ocr_gemini(f_bytes, mime, api_key, "Bu kaydı analiz et ve içeriğini dök.")
                                        context_data += f"\n--- MEDYA ({file.name}) ---\n{media_res}"
                                    elif ext == 'udf': context_data += f"\n--- UDF ({file.name}) ---\n{parse_udf(f_bytes)}"
                                except Exception as e: context_data += f"\n⚠️ {file.name} okunurken hata: {str(e)}"
                    if not user_text_input and not voice_text and not context_data: st.warning("Abi boş gönderdin.")
                    else:
                        final_prompt = f"GÖREV: Yardımsever asistan ol.\nSORU: {user_text_input}\nSESLİ SORU: {voice_text}\nDOSYALAR: {context_data[:100000]}"
                        with st.spinner("Hazırlıyorum abi..."):
                            resp = get_ai_response(final_prompt, api_key)
                            st.session_state.buyur_abi_response = resp
            if st.session_state.buyur_abi_response:
                st.markdown(f"<div class='buyur-abi-kutusu'>{st.session_state.buyur_abi_response}</div>", unsafe_allow_html=True)
                b_col1, b_col2 = st.columns(2)
                with b_col1: st.download_button("📄 PDF Olarak Al", create_pdf_file(st.session_state.buyur_abi_response), "Cevap.pdf", "application/pdf")
                with b_col2: st.download_button("📝 Word Olarak Al", create_word_file(st.session_state.buyur_abi_response), "Cevap.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    with tab11:
        st.subheader("⏰ Duruşma Hatırlatıcı & Takvim")
        st.info("UYAP'tan aldığınız .ics (Takvim) dosyalarını buraya yükleyin. Yaklaşan duruşmaları otomatik listeler ve 24 saat kalanlar için ALARM verir.")
        col_h1, col_h2 = st.columns([1, 3])
        with col_h1:
            ics_file = st.file_uploader("Takvim Dosyası (.ics) Yükle", type=['ics'])
            if ics_file:
                if st.button("📅 Takvime Ekle", type="primary"):
                    events = parse_ics_data(BytesIO(ics_file.getvalue()))
                    if events:
                        count = 0
                        for evt in events:
                            exists = any(e['dtstart'] == evt['dtstart'] and e['summary'] == evt['summary'] for e in st.session_state.durusma_listesi)
                            if not exists:
                                st.session_state.durusma_listesi.append(evt)
                                count += 1
                        save_durusma_data(st.session_state.durusma_listesi)
                        st.success(f"{count} yeni duruşma eklendi!")
                    else: st.error("Dosya okunamadı.")
            st.divider()
            if st.button("🗑️ Tüm Listeyi Temizle"):
                st.session_state.durusma_listesi = []
                save_durusma_data([])
                st.rerun()
        with col_h2:
            if not st.session_state.durusma_listesi: st.info("Henüz eklenmiş bir duruşma yok.")
            else:
                sorted_events = sorted(st.session_state.durusma_listesi, key=lambda x: x['dtstart'])
                now = datetime.now()
                st.write(f"**Toplam Duruşma Sayısı:** {len(sorted_events)}")
                for evt in sorted_events:
                    dt = evt['dtstart']
                    diff = dt - now
                    tarih_str = dt.strftime("%d.%m.%Y %H:%M")
                    is_alarm = timedelta(0) < diff < timedelta(hours=24)
                    is_past = diff < timedelta(0)
                    if is_past:
                        with st.expander(f"✅ (GEÇMİŞ) {tarih_str} - {evt.get('summary', 'Başlıksız')}"):
                            st.write(f"**Mahkeme:** {evt.get('location', '-')}")
                            st.write(f"**Detay:** {evt.get('description', '-')}")
                    elif is_alarm:
                        st.markdown(f"""<div class="alarm-kutusu">🚨 ALARM: DURUŞMAYA AZ KALDI!<br>📅 {tarih_str}<br>⚖️ {evt.get('summary', 'Başlıksız')}<br>📍 {evt.get('location', '-')}</div>""", unsafe_allow_html=True)
                        with st.expander("Detayları Gör"): st.write(f"**Açıklama:** {evt.get('description', '-')}")
                    else:
                        st.markdown(f"""<div class="normal-durusma">📅 <b>{tarih_str}</b> (Kalan: {diff.days} gün)<br>⚖️ {evt.get('summary', 'Başlıksız')}<br>📍 {evt.get('location', '-')}</div>""", unsafe_allow_html=True)

    with tab12:
        st.subheader("🗄️ Doküman Yönetimi ve Arşivleme")
        st.info(f"Verileriniz bilgisayarınızda '{ROOT_DIR}' klasöründe saklanır. TIF, PDF, Resim dahil tüm dosyaları okur.")

        ar_tab1, ar_tab2, ar_tab3 = st.tabs(["📂 Yeni Dava Dosyası Aç", "📎 Dosya Yükle", "🔍 Arşivde Ara & Analiz"])

        with ar_tab1:
            st.markdown("#### Yeni Dava Klasörü Oluştur")
            c_tur, c_mah = st.columns(2)
            with c_tur: dava_turu = st.selectbox("Dava Türü", ["Hukuk Davaları", "Ceza Davaları", "İcra Dosyaları", "İdari Davalar"])
            with c_mah: yeni_mahkeme = st.text_input("Mahkeme Adı", placeholder="Örn: Ankara 1. Asliye Hukuk")
            c_esas, c_taraf = st.columns(2)
            with c_esas: yeni_esas = st.text_input("Dosya/Esas No", placeholder="Örn: 2024-123")
            with c_taraf: yeni_taraflar = st.text_input("Taraf Bilgileri", placeholder="Örn: Ahmet Yılmaz vs Mehmet Demir")

            if st.button("📁 Klasörü Oluştur"):
                if yeni_mahkeme and yeni_esas:
                    safe_mah = "".join([c for c in yeni_mahkeme if c.isalnum() or c in (' ', '-', '_')]).strip()
                    safe_esas = "".join([c for c in yeni_esas if c.isalnum() or c in (' ', '-', '_')]).strip()
                    target_path = os.path.join(ROOT_DIR, dava_turu, safe_mah, safe_esas)
                    try:
                        os.makedirs(target_path, exist_ok=True)
                        with open(os.path.join(target_path, "Dosya_Bilgileri.txt"), "w", encoding="utf-8") as f:
                            f.write(f"Dava Türü: {dava_turu}\nMahkeme: {yeni_mahkeme}\nEsas: {yeni_esas}\nTaraflar: {yeni_taraflar}\nOluşturma: {datetime.now()}")
                        st.success(f"✅ Klasör Başarıyla Oluşturuldu:\n{target_path}")
                    except Exception as e: st.error(f"Hata: {str(e)}")
                else: st.warning("Lütfen Mahkeme ve Esas No giriniz.")

        with ar_tab2:
            st.markdown("#### Mevcut Dosyaya Evrak Ekle")
            if os.path.exists(ROOT_DIR):
                turler = [d for d in os.listdir(ROOT_DIR) if os.path.isdir(os.path.join(ROOT_DIR, d))]
                if not turler: st.warning("Henüz hiç dava klasörü yok.")
                else:
                    secilen_tur = st.selectbox("Dava Türü Seç", turler)
                    tur_path = os.path.join(ROOT_DIR, secilen_tur)
                    mahkemeler = [d for d in os.listdir(tur_path) if os.path.isdir(os.path.join(tur_path, d))]
                    if mahkemeler:
                        secilen_mah = st.selectbox("Mahkeme Seç", mahkemeler)
                        mah_path = os.path.join(tur_path, secilen_mah)
                        dosyalar = [d for d in os.listdir(mah_path) if os.path.isdir(os.path.join(mah_path, d))]
                        if dosyalar:
                            secilen_dosya = st.selectbox("Dosya No Seç", dosyalar)
                            final_path = os.path.join(mah_path, secilen_dosya)
                            st.info(f"Seçilen Klasör: {final_path}")
                            yuklenen_evraklar = st.file_uploader("Evrakları Yükle", 
                                                               type=['pdf','doc','docx','udf','png','jpg','mp3','mp4','wav','txt','tif','tiff'],
                                                               accept_multiple_files=True)
                            if st.button("💾 Evrakları Kaydet"):
                                if yuklenen_evraklar:
                                    for evrak in yuklenen_evraklar:
                                        with open(os.path.join(final_path, evrak.name), "wb") as f:
                                            f.write(evrak.getbuffer())
                                    st.success(f"✅ {len(yuklenen_evraklar)} adet dosya başarıyla kaydedildi!")
                                else: st.warning("Dosya seçmediniz.")
                        else: st.warning("Bu mahkemede dosya yok.")
                    else: st.warning("Bu türde mahkeme yok.")
            else: st.error("Ana arşiv klasörü bulunamadı.")

        with ar_tab3:
            st.markdown("#### 🔍 Arşivde Arama ve Yapay Zeka Analizi")
            
            if st.session_state.aktif_dosya_yolu:
                st.markdown(f"<div class='arsiv-kutusu'><b>📂 ÇALIŞILAN DOSYA: {st.session_state.aktif_dosya_adi}</b><br>Şu an sadece bu dosyadaki evraklar hafızada.</div>", unsafe_allow_html=True)
                
                if st.button("⬅️ Dosyayı Kapat ve Listeye Dön"):
                    st.session_state.aktif_dosya_yolu = ""
                    st.session_state.arsiv_context = ""
                    st.session_state.arsiv_genel_ozet = ""
                    st.session_state.arsiv_soru_cevap = ""
                    st.rerun()
                
                st.divider()
                col_analiz, col_soru = st.columns(2)
                
                with col_analiz:
                    st.markdown("### 📊 Analiz Et")
                    st.info("Bu klasördeki belgeleri özetler.")
                    if st.button("Dosyayı Analiz Et", type="primary", use_container_width=True):
                        if not api_key: st.error("API Key gerekli.")
                        else:
                            with st.spinner("Sadece bu dosyadaki evraklar analiz ediliyor..."):
                                prompt = f"GÖREV: Bu dava dosyasının içeriğini özetle, hukuki durumu analiz et.\nİÇERİK: {st.session_state.arsiv_context[:500000]}"
                                res = get_ai_response(prompt, api_key)
                                st.session_state.arsiv_genel_ozet = res
                    
                    if st.session_state.arsiv_genel_ozet:
                        st.markdown(f"<div class='kanun-kutusu'>{st.session_state.arsiv_genel_ozet}</div>", unsafe_allow_html=True)
                        
                        st.markdown("###### 📥 Raporu İndir / Paylaş")
                        c_down1, c_down2 = st.columns(2)
                        with c_down1: st.download_button("📄 PDF", create_pdf_file(st.session_state.arsiv_genel_ozet), "Analiz.pdf", "application/pdf")
                        with c_down2: st.download_button("📝 Word", create_word_file(st.session_state.arsiv_genel_ozet), "Analiz.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        
                        wa_no_analiz = st.text_input("WhatsApp No", placeholder="905...", key="wa_analiz")
                        if wa_no_analiz:
                            encoded_msg = urllib.parse.quote(f"*Dosya Analizi:*\n{st.session_state.arsiv_genel_ozet}")
                            st.link_button("📲 WhatsApp'tan Gönder", f"https://wa.me/{wa_no_analiz}?text={encoded_msg}")

                with col_soru:
                    st.markdown("### ❓ Soru Sor")
                    st.info("Sadece bu dosya ile ilgili sorular sorun.")
                    arsiv_soru = st.text_input("Sorunuzu yazın", placeholder="Örn: Bilirkişi raporu ne zaman gelmiş?")
                    
                    if st.button("Soruyu Cevapla", use_container_width=True):
                        if not api_key: st.error("API Key gerekli.")
                        elif not arsiv_soru: st.warning("Soru yazın.")
                        else:
                            with st.spinner("Bu dosyadaki belgeler taranıyor..."):
                                prompt = f"""
                                soruyu cevapla.
                                SORU: {arsiv_soru}
                                DOSYA İÇERİĞİ:
                                {st.session_state.arsiv_context[:500000]}
                                """
                                res = get_ai_response(prompt, api_key)
                                st.session_state.arsiv_soru_cevap = res
                    
                    if st.session_state.arsiv_soru_cevap:
                        st.markdown(f"<div class='kanun-kutusu'>{st.session_state.arsiv_soru_cevap}</div>", unsafe_allow_html=True)
                        
                        st.markdown("###### 📥 Cevabı İndir / Paylaş")
                        q_down1, q_down2 = st.columns(2)
                        with q_down1: st.download_button("📄 PDF", create_pdf_file(st.session_state.arsiv_soru_cevap), "Cevap.pdf", "application/pdf")
                        with q_down2: st.download_button("📝 Word", create_word_file(st.session_state.arsiv_soru_cevap), "Cevap.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                        wa_no_soru = st.text_input("WhatsApp No", placeholder="905...", key="wa_soru")
                        if wa_no_soru:
                            encoded_msg = urllib.parse.quote(f"*Soru:* {arsiv_soru}\n\n*Cevap:*\n{st.session_state.arsiv_soru_cevap}")
                            st.link_button("📲 WhatsApp'tan Gönder", f"https://wa.me/{wa_no_soru}?text={encoded_msg}")

            else:
                arama_terimi = st.text_input("Aranacak Kelime (Dosya No, Mahkeme veya Dosya Adı)", placeholder="Örn: 2024-123 veya Ahmet Yılmaz")
                
                if st.button("🔎 Ara"):
                    st.session_state.arsiv_arama_sonuclari = []
                    bulunanlar = []
                    for root, dirs, files in os.walk(ROOT_DIR):
                        if arama_terimi.lower() in root.lower():
                            bulunanlar.append({"tip": "KLASÖR", "yol": root, "dosyalar": files})
                        for file in files:
                            if arama_terimi.lower() in file.lower():
                                bulunanlar.append({"tip": "DOSYA", "yol": os.path.join(root, file), "dosya_adi": file})
                    st.session_state.arsiv_arama_sonuclari = bulunanlar

                if st.session_state.arsiv_arama_sonuclari:
                    st.success(f"{len(st.session_state.arsiv_arama_sonuclari)} sonuç bulundu.")
                    for sonuc in st.session_state.arsiv_arama_sonuclari:
                        if sonuc["tip"] == "KLASÖR":
                            with st.expander(f"📁 {sonuc['yol']}"):
                                st.write(f"İçerik: {len(sonuc['dosyalar'])} dosya")
                                
                                if st.button(f"📂 Bu Dosyayı Aç ve Çalış ({os.path.basename(sonuc['yol'])})", key=sonuc['yol']):
                                    full_text = ""
                                    if not api_key: st.error("Lütfen önce API Key giriniz.")
                                    else:
                                        st.session_state.arsiv_context = ""
                                        st.session_state.arsiv_genel_ozet = ""
                                        st.session_state.arsiv_soru_cevap = ""
                                        
                                        with st.spinner("Sadece seçilen klasördeki dosyalar okunuyor..."):
                                            sadece_bu_klasordeki_dosyalar = [f for f in os.listdir(sonuc['yol']) if os.path.isfile(os.path.join(sonuc['yol'], f))]
                                            
                                            for f_name in sadece_bu_klasordeki_dosyalar:
                                                f_path = os.path.join(sonuc['yol'], f_name)
                                                ext = f_name.split('.')[-1].lower()
                                                try:
                                                    with open(f_path, 'rb') as f:
                                                        file_content = BytesIO(f.read())
                                                    
                                                    if ext == 'txt':
                                                        full_text += f"\n--- {f_name} ---\n{file_content.getvalue().decode('utf-8', errors='ignore')}"
                                                    elif ext == 'pdf':
                                                        pdf_text = parse_pdf(file_content)
                                                        if not pdf_text:
                                                            file_content.seek(0)
                                                            pdf_text = perform_ocr_gemini(file_content, "application/pdf", api_key)
                                                        full_text += f"\n--- {f_name} ---\n{pdf_text}"
                                                    elif ext in ['docx', 'doc']:
                                                        full_text += f"\n--- {f_name} ---\n{extract_text_from_docx(file_content)}"
                                                    elif ext == 'udf':
                                                        full_text += f"\n--- {f_name} ---\n{parse_udf(file_content)}"
                                                    elif ext in ['png', 'jpg', 'jpeg', 'tif', 'tiff']:
                                                        mime = "image/tiff" if ext in ['tif', 'tiff'] else "image/jpeg"
                                                        file_content.seek(0)
                                                        ocr_res = perform_ocr_gemini(file_content, mime, api_key)
                                                        full_text += f"\n--- {f_name} ---\n{ocr_res}"
                                                except Exception as e:
                                                    full_text += f"\n--- {f_name} (HATA) ---\n{str(e)}"
                                            
                                            st.session_state.arsiv_context = full_text
                                            st.session_state.aktif_dosya_adi = os.path.basename(sonuc['yol'])
                                            st.session_state.aktif_dosya_yolu = sonuc['yol']
                                            st.rerun()

    # --- TAB 13: UYAP ANALİZ ---
    with tab13:
        st.subheader("🏛️ UYAP Toplu Dosya Analizi")
        st.info("UYAP'tan indirdiğiniz ZIP dosyalarını yükleyin. Sistem son 5 evrağı analiz eder.")
        uyap_zips = st.file_uploader("UYAP Dosyalarını Yükle (ZIP)", type=['zip'], accept_multiple_files=True)
        
        if uyap_zips and st.button("🚀 Dosyaları Analiz Et", type="primary"):
            if not api_key: st.error("API Key gerekli.")
            else:
                progress_bar = st.progress(0)
                for idx, zip_file in enumerate(uyap_zips):
                    dosya_adi = zip_file.name
                    st.markdown(f"### 📂 {dosya_adi}")
                    with st.spinner(f"{dosya_adi} inceleniyor..."):
                        try:
                            with zipfile.ZipFile(zip_file) as z:
                                files_info = []
                                for info in z.infolist():
                                    if not info.is_dir():
                                        files_info.append({'name': info.filename, 'date': datetime(*info.date_time)})
                                sorted_files = sorted(files_info, key=lambda x: x['date'], reverse=True)[:5]
                                
                                file_context = ""
                                for f_info in sorted_files:
                                    fname = f_info['name']
                                    fdate = f_info['date'].strftime('%d.%m.%Y')
                                    with z.open(fname) as f:
                                        file_bytes = BytesIO(f.read())
                                        ext = fname.split('.')[-1].lower()
                                        content = ""
                                        try:
                                            if ext == 'udf': content = parse_udf(file_bytes)
                                            elif ext == 'pdf': content = parse_pdf(file_bytes)
                                            elif ext in ['docx', 'doc']: content = extract_text_from_docx(file_bytes)
                                            elif ext == 'txt': content = file_bytes.read().decode('utf-8', errors='ignore')
                                        except: content = "Okunamadı"
                                        file_context += f"\n--- {fname} ({fdate}) ---\n{content[:5000]}"
                                
                                prompt = f"GÖREV: Bu dava dosyasının SON 5 evrağına göre durumu özetle.\nEVRAKLAR:\n{file_context}"
                                analiz_sonucu = get_ai_response(prompt, api_key)
                                st.markdown(f"<div class='uyap-kutusu'>{analiz_sonucu}</div>", unsafe_allow_html=True)
                        except Exception as e: st.error(f"Hata: {str(e)}")
                    progress_bar.progress((idx + 1) / len(uyap_zips))

    # --- TAB 14: FAİZ HESABI ---
    with tab14:
        st.subheader("🧮 Yasal Faiz Hesaplama")
        st.info("Basit faiz mantığıyla çalışır. Resmi hesaplamalar için kontrol ediniz.")
        
        col_f1, col_f2 = st.columns(2)
        with col_f1:
            ana_para = st.number_input("Ana Para (TL)", min_value=0.0, value=10000.0, step=100.0)
            baslangic_tarihi = st.date_input("Faiz Başlangıç Tarihi")
        with col_f2:
            bitis_tarihi = st.date_input("Faiz Bitiş Tarihi (Bugün)", value=datetime.now())
            faiz_orani = st.number_input("Yıllık Faiz Oranı (%)", value=9.0, help="Yasal Faiz: %9, Ticari: %15+ (Değişken)")
        
        if st.button("Hesapla"):
            if bitis_tarihi > baslangic_tarihi:
                gun_farki = (bitis_tarihi - baslangic_tarihi).days
                faiz_tutari = (ana_para * faiz_orani * gun_farki) / 36500
                toplam_tutar = ana_para + faiz_tutari
                
                st.markdown(f"""
                <div class="ozel-sekme">
                    <h4>📊 Hesaplama Sonucu</h4>
                    <ul>
                        <li><b>Gün Sayısı:</b> {gun_farki} gün</li>
                        <li><b>İşlemiş Faiz:</b> {faiz_tutari:,.2f} TL</li>
                        <li><b>Ana Para:</b> {ana_para:,.2f} TL</li>
                        <li><h3 style='color:#d32f2f'>TOPLAM: {toplam_tutar:,.2f} TL</h3></li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)
            else:
                st.error("Bitiş tarihi başlangıçtan büyük olmalı.")

    # --- TAB 15: SÜRE HESAPLAYICI ---
    with tab15:
        st.subheader("⏳ HMK/CMK Süre Hesaplayıcı")
        st.info("Adli tatil ve hafta sonlarını dikkate alır.")
        
        col_t1, col_t2 = st.columns(2)
        with col_t1:
            teblig_tarihi = st.date_input("Tebliğ Tarihi")
            sure_gun = st.number_input("Süre (Gün)", min_value=1, value=14)
        with col_t2:
            adli_tatil_dahil = st.checkbox("Adli Tatile Denk Gelirse Uzat (20 Temmuz - 31 Ağustos)", value=True)
        
        if st.button("Son Günü Bul"):
            son_gun = teblig_tarihi + timedelta(days=sure_gun)
            
            # Adli Tatil Kontrolü (Basit Mantık)
            if adli_tatil_dahil:
                yil = son_gun.year
                tatil_bas = date(yil, 7, 20)
                tatil_bit = date(yil, 8, 31)
                
                # Eğer süre adli tatil içinde bitiyorsa
                if tatil_bas <= son_gun <= tatil_bit:
                    son_gun = date(yil, 9, 7) # Adli tatil bitiminden 1 hafta sonraya uzar (HMK m.104)
                    st.warning("⚠️ Süre Adli Tatile denk geldiği için 7 Eylül'e uzatıldı!")

            # Hafta Sonu Kontrolü
            if son_gun.weekday() == 5: # Cumartesi
                son_gun += timedelta(days=2)
                st.info("📆 Süre Cumartesiye geldiği için Pazartesiye uzadı.")
            elif son_gun.weekday() == 6: # Pazar
                son_gun += timedelta(days=1)
                st.info("📆 Süre Pazara geldiği için Pazartesiye uzadı.")
            
            kalan_gun = (son_gun - date.today()).days
            
            st.markdown(f"""
            <div class="alarm-kutusu" style="text-align:center;">
                <h2>SON GÜN: {son_gun.strftime('%d.%m.%Y')}</h2>
                <p>({son_gun.strftime('%A')})</p>
                <h4>Kalan Süre: {kalan_gun} Gün</h4>
            </div>
            """, unsafe_allow_html=True)

    # --- TAB 16: İLİŞKİ AĞI ---
    with tab16:
        st.subheader("🕸️ Dosya İlişki Ağı (Relationship Graph)")
        st.info("Yüklenen dosyadaki kişileri ve aralarındaki bağlantıyı görselleştirir.")
        
        if not st.session_state.doc_text:
            st.warning("Lütfen önce sol menüden bir dosya yükleyin.")
        else:
            if st.button("İlişki Ağını Çiz", type="primary"):
                if not api_key: st.error("API Key gerekli.")
                else:
                    with st.spinner("Kişiler analiz ediliyor..."):
                        prompt = f"""
                        GÖREV: Bu metindeki kişileri ve rollerini (Davacı, Davalı, Tanık, Vekil vb.) tespit et.
                        ÇIKTI FORMATI: Graphviz DOT formatında kod ver. Sadece kodu ver.
                        Örnek: digraph G {{ "Ahmet" [label="Ahmet (Davacı)"]; "Mehmet" [label="Mehmet (Davalı)"]; "Ahmet" -> "Mehmet" [label="Dava Açtı"]; }}
                        METİN: {st.session_state.doc_text[:50000]}
                        """
                        dot_code = get_ai_response(prompt, api_key)
                        
                        # Temizlik
                        dot_code = dot_code.replace("```dot", "").replace("```", "").strip()
                        
                        try:
                            st.graphviz_chart(dot_code)
                            st.success("İlişki ağı oluşturuldu.")
                        except:
                            st.error("Grafik oluşturulamadı. AI çıktısı hatalı olabilir.")
                            st.code(dot_code)

    # --- TAB 17: SÖZLEŞME ANALİZ ---
    with tab17:
        st.subheader("📝 Sözleşme Risk Analizi")
        sozlesme_file = st.file_uploader("Sözleşme Yükle (PDF/Word)", type=['pdf', 'docx'])
        
        if sozlesme_file and st.button("Sözleşmeyi İncele"):
            if not api_key: st.error("API Key gerekli.")
            else:
                with st.spinner("Sözleşme taranıyor..."):
                    s_bytes = BytesIO(sozlesme_file.getvalue())
                    s_ext = sozlesme_file.name.split('.')[-1].lower()
                    s_text = extract_text_from_docx(s_bytes) if s_ext == 'docx' else parse_pdf(s_bytes)
                    
                    prompt = f"""
                    GÖREV: Bu sözleşmeyi bir avukat gözüyle incele.
                    ÇIKTI FORMATI:
                    1. 🚩 **RİSKLİ MADDELER**: Aleyhe olan maddeler.
                    2. ⚠️ **EKSİK HUSUSLAR**: Sözleşmede olması gereken ama olmayanlar.
                    3. ✅ **ÖNERİLER**: Nasıl düzeltilmeli?
                    
                    SÖZLEŞME: {s_text[:50000]}
                    """
                    st.session_state.sozlesme_analiz = get_ai_response(prompt, api_key)
        
        if st.session_state.sozlesme_analiz:
            st.markdown(st.session_state.sozlesme_analiz)
            st.download_button("📥 Raporu İndir", create_word_file(st.session_state.sozlesme_analiz), "Sozlesme_Analiz.docx")

    # --- TAB 18: MÜVEKKİL BİLGİLENDİRME ---
    with tab18:
        st.subheader("📧 Müvekkil Bilgilendirme Mesajı")
        st.info("Hukuki metni, müvekkilin anlayacağı nazik bir dile çevirir.")
        
        durum_metni = st.text_area("Hukuki Gelişme / Karar Metni", placeholder="Örn: Dosya bilirkişiden döndü, kusur oranı %50 verildi...")
        hitap = st.text_input("Müvekkil Adı", placeholder="Ahmet Bey")
        
        if st.button("Mesaj Taslağı Oluştur"):
            if not api_key: st.error("API Key gerekli.")
            elif not durum_metni: st.warning("Metin giriniz.")
            else:
                with st.spinner("Yazılıyor..."):
                    prompt = f"""
                    GÖREV: Bu hukuki durumu müvekkilim {hitap}'e anlatmak için kısa, anlaşılır, profesyonel ve nazik bir mesaj yaz.
                    Hukuki terim kullanma.
                    DURUM: {durum_metni}
                    """
                    st.session_state.muvekkil_mesaj = get_ai_response(prompt, api_key)
        
        if st.session_state.muvekkil_mesaj:
            st.markdown(f"<div class='buyur-abi-kutusu'>{st.session_state.muvekkil_mesaj}</div>", unsafe_allow_html=True)
            if st.button("Kopyala (Simüle)"):
                st.toast("Mesaj panoya kopyalandı!")

if __name__ == "__main__":
    main()
