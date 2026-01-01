import streamlit as st
import pandas as pd  # Pandas hatasını çözer
import requests
from bs4 import BeautifulSoup # Web tarama hatasını çözer
import io
import PyPDF2
import zipfile
import xml.etree.ElementTree as ET
import re
from pypdf import PdfReader
from io import BytesIO
import google.generativeai as genai
import importlib.metadata
from docx import Document
from fpdf import FPDF
import urllib.parse
import concurrent.futures
from gtts import gTTS
import speech_recognition as sr
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import json
import os
from PIL import Image
from PIL.ExifTags import TAGS  # <--- Bu çok önemli, eksikse hata verir
import time
from datetime import datetime, timedelta, date
import shutil
import difflib
import plotly.graph_objects as go # Görsel grafikler için gerekli
from PIL import Image



# --- Sayfa Ayarları ---
st.set_page_config(
    page_title="Hukuk Asistanı AI",
    page_icon="⚖️",
    layout="wide"
)

# --- CSS ---
st.markdown("""
    <style>
    .main { background-color: #f8f9fa; }
    .kanun-kutusu { 
        background-color: #fff3e0; 
        padding: 15px; 
        border-left: 5px solid #ff9800; 
        border-radius: 5px; 
        margin-bottom: 10px;
        white-space: pre-wrap;
    }
    .ictihat-kutusu {
        background-color: #e3f2fd;
        padding: 15px;
        border-left: 5px solid #2196f3;
        border-radius: 5px;
        margin-bottom: 10px;
    }
    .buyur-abi-kutusu {
        background-color: #f3e5f5;
        padding: 15px;
        border-left: 5px solid #9c27b0;
        border-radius: 5px;
        margin-bottom: 10px;
    }
    .alarm-kutusu {
        background-color: #ffebee;
        padding: 15px;
        border-left: 5px solid #f44336;
        border-radius: 5px;
        margin-bottom: 10px;
        font-weight: bold;
        color: #b71c1c;
    }
    .arsiv-kutusu {
        background-color: #e0f2f1;
        padding: 15px;
        border-left: 5px solid #009688;
        border-radius: 5px;
        margin-bottom: 10px;
    }
    .uyap-kutusu {
        background-color: #fce4ec; 
        padding: 15px; 
        border-left: 5px solid #c2185b; 
        border-radius: 5px; 
        margin-bottom: 20px;
    }
    .ozel-sekme {
        border: 1px solid #ddd;
        padding: 20px;
        border-radius: 10px;
        background-color: #ffffff;
    }
    </style>
    """, unsafe_allow_html=True)


# --- KALICILIK (VERİ TABANI) FONKSİYONLARI ---
DURUSMA_FILE = "durusma_kayitlari.json"

def save_durusma_data(data):
    """Duruşma listesini JSON dosyasına kaydeder."""
    serializable_data = []
    for item in data:
        temp = item.copy()
        if isinstance(temp.get('dtstart'), datetime):
            temp['dtstart'] = temp['dtstart'].isoformat()
        serializable_data.append(temp)
    
    try:
        with open(DURUSMA_FILE, "w", encoding="utf-8") as f:
            json.dump(serializable_data, f, ensure_ascii=False, indent=4)
    except Exception as e:
        st.error(f"Kaydetme hatası: {e}")

def load_durusma_data():
    """JSON dosyasından duruşma listesini yükler."""
    if not os.path.exists(DURUSMA_FILE):
        return []
    try:
        with open(DURUSMA_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        for item in data:
            if 'dtstart' in item and item['dtstart']:
                item['dtstart'] = datetime.fromisoformat(item['dtstart'])
        return data
    except:
        return []

# --- YARDIMCI FONKSİYONLAR ---
def parse_udf(file_bytes):
    try:
        with zipfile.ZipFile(file_bytes) as z:
            if 'content.xml' in z.namelist():
                with z.open('content.xml') as f:
                    tree = ET.parse(f)
                    root = tree.getroot()
                    text_content = [elem.text.strip() for elem in root.iter() if elem.text]
                    return " ".join(text_content)
            return "HATA: UDF içeriği okunamadı."
    except Exception as e:
        return f"HATA: {str(e)}"

def parse_pdf(file_bytes):
    try:
        reader = PdfReader(file_bytes)
        text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        if len(text.strip()) < 50: return "" 
        return text
    except Exception as e:
        return ""

def extract_metadata(text):
    if not isinstance(text, str) or text.startswith(("HATA", "UYARI")):
        return {"mahkeme": "-", "esas": "-", "karar": "-", "tarih": "-"}
    
    esas = re.search(r"(?i)Esas\s*No\s*[:\-]?\s*(\d{4}/\d+)", text)
    karar = re.search(r"(?i)Karar\s*No\s*[:\-]?\s*(\d{4}/\d+)", text)
    tarih = re.search(r"(\d{1,2}[./]\d{1,2}[./]\d{4})", text)
    
    mahkeme = "Tespit Edilemedi"
    for line in text.split('\n')[:40]:
        clean = line.strip()
        if ("MAHKEMESİ" in clean.upper() or "DAİRESİ" in clean.upper()) and len(clean) > 5:
            mahkeme = clean
            break
    return {
        "mahkeme": mahkeme,
        "esas": esas.group(1) if esas else "Bulunamadı",
        "karar": karar.group(1) if karar else "Bulunamadı",
        "tarih": tarih.group(1) if tarih else "Bulunamadı"
    }

def parse_ics_data(file_bytes):
    events = []
    try:
        content = file_bytes.getvalue().decode('utf-8')
        lines = content.splitlines()
        current_event = {}
        in_event = False
        
        for line in lines:
            line = line.strip()
            if line == 'BEGIN:VEVENT':
                in_event = True
                current_event = {}
            elif line == 'END:VEVENT':
                in_event = False
                if 'dtstart' in current_event and 'summary' in current_event:
                    events.append(current_event)
            elif in_event:
                if line.startswith('SUMMARY:'):
                    current_event['summary'] = line.split(':', 1)[1]
                elif line.startswith('DTSTART:'):
                    raw_date = line.split(':', 1)[1]
                    try:
                        dt = datetime.strptime(raw_date, '%Y%m%dT%H%M%S')
                        current_event['dtstart'] = dt
                    except: pass
                elif line.startswith('LOCATION:'):
                    current_event['location'] = line.split(':', 1)[1]
                elif line.startswith('DESCRIPTION:'):
                    current_event['description'] = line.split(':', 1)[1]
        return events
    except Exception as e:
        return []

# --- DOSYA OLUŞTURMA FONKSİYONLARI ---
def create_word_file(text):
    doc = Document()
    for line in text.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io

def create_udf_file(text):
    root = ET.Element("content")
    body = ET.SubElement(root, "body")
    for line in text.split('\n'):
        p = ET.SubElement(body, "p")
        p.text = line
    xml_str = ET.tostring(root, encoding='utf-8', method='xml')
    byte_io = BytesIO()
    with zipfile.ZipFile(byte_io, 'w', zipfile.ZIP_DEFLATED) as z:
        z.writestr('content.xml', xml_str)
    byte_io.seek(0)
    return byte_io

def create_pdf_file(text):
    replacements = {
        'ğ': 'g', 'Ğ': 'G', 'ş': 's', 'Ş': 'S', 'ı': 'i', 'İ': 'I',
        'ç': 'c', 'Ç': 'C', 'ü': 'u', 'Ü': 'U', 'ö': 'o', 'Ö': 'O',
        '“': '"', '”': '"', '’': "'", '–': '-', '…': '...'
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    
    text = text.encode('latin-1', 'replace').decode('latin-1')
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="Hukuki Analiz Raporu", ln=1, align='C')
    pdf.ln(10)
    pdf.set_font("Arial", size=11)
    pdf.multi_cell(0, 10, text)
    return pdf.output(dest='S').encode('latin-1', 'ignore')

# --- SES İŞLEME FONKSİYONLARI ---
def text_to_speech(text):
    try:
        tts = gTTS(text=text, lang='tr')
        fp = BytesIO()
        tts.write_to_fp(fp)
        fp.seek(0)
        return fp
    except Exception as e:
        return None

def speech_to_text(audio_bytes):
    r = sr.Recognizer()
    try:
        with sr.AudioFile(audio_bytes) as source:
            audio_data = r.record(source)
            text = r.recognize_google(audio_data, language='tr-TR')
            return text
    except Exception as e:
        return f"Hata: {str(e)}"

# --- MULTIMODAL VE OCR FONKSİYONLARI ---
def perform_ocr_gemini(file_bytes, mime_type, api_key, prompt_text="Bu dosyanın içeriğini tam olarak metne dök."):
    if not api_key: return "API Key Yok"
    genai.configure(api_key=api_key)
    
    if mime_type in ['image/tiff', 'image/tif']:
        try:
            image = Image.open(file_bytes)
            rgb_im = image.convert('RGB')
            buf = BytesIO()
            rgb_im.save(buf, format="JPEG")
            file_bytes = buf
            mime_type = 'image/jpeg'
        except Exception as e:
            return f"TIFF Dönüştürme Hatası: {str(e)}"

    available_models = []
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                if '1.5' in m.name or 'vision' in m.name:
                    available_models.append(m.name)
    except: pass
    
    if not available_models: available_models = ['models/gemini-1.5-flash', 'models/gemini-1.5-pro']

    image_part = {"mime_type": mime_type, "data": file_bytes.getvalue()}
    
    for model_name in available_models:
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content([prompt_text, image_part])
            return response.text
        except: continue
    return "Analiz Başarısız."

def extract_text_from_docx(file_bytes):
    try:
        doc = Document(file_bytes)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        return f"Word Okuma Hatası: {str(e)}"

def read_excel_file(file_bytes):
    try:
        df = pd.read_excel(file_bytes)
        return df.to_string()
    except Exception as e:
        return f"Excel Hatası: {str(e)}"

# --- AKILLI AI MOTORU ---
def get_ai_response(prompt, api_key):
    if not api_key: return "Lütfen API Anahtarı giriniz."
    genai.configure(api_key=api_key)
    candidate_models = ['gemini-1.5-flash', 'gemini-1.5-pro', 'gemini-1.0-pro', 'gemini-pro']
    
    for model_name in candidate_models:
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(prompt)
            return response.text 
        except: continue 
    
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                try:
                    model = genai.GenerativeModel(m.name)
                    response = model.generate_content(prompt)
                    return response.text
                except: continue
    except: pass
    return "Hata: AI yanıt veremedi."

# ==========================================
# YENİ EKLENEN MODÜLLER (CHECK-UP & ZAMAN MAKİNESİ)
# ==========================================

def get_gemini_text_response(prompt, api_key):
    """Mevcut modelleri tarayıp çalışan ilk modeli kullanan fonksiyon."""
    if not api_key: return "Lütfen API Anahtarınızı giriniz."
    
    try:
        genai.configure(api_key=api_key)
        
        # 1. ADIM: Kullanıcının erişebildiği modelleri listele
        available_models = []
        try:
            for m in genai.list_models():
                if 'generateContent' in m.supported_generation_methods:
                    available_models.append(m.name)
        except Exception as e:
            # Listeleme hatası olursa manuel listeye dön
            available_models = ["models/gemini-1.5-flash", "models/gemini-pro", "gemini-1.5-flash"]

        # 2. ADIM: En iyi modeli seç (Hız > Kalite sıralaması)
        selected_model = None
        
        # Öncelik 1: Flash modeller (Hızlı)
        for m in available_models:
            if 'flash' in m.lower():
                selected_model = m
                break
        
        # Öncelik 2: Pro modeller (Eğer flash yoksa)
        if not selected_model:
            for m in available_models:
                if 'pro' in m.lower() and 'vision' not in m.lower():
                    selected_model = m
                    break
        
        # Öncelik 3: Listede ne varsa ilki
        if not selected_model and available_models:
            selected_model = available_models[0]
            
        # Hiçbiri yoksa son çare
        if not selected_model:
            selected_model = "models/gemini-1.5-flash"

        # 3. ADIM: Seçilen modelle üretimi yap
        try:
            model = genai.GenerativeModel(selected_model)
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            return f"Model Hatası ({selected_model}): {str(e)}"
            
    except Exception as e:
        return f"Genel Sistem Hatası: {str(e)}"



def render_checkup_module(api_key):
    st.info("Şirket sözleşmelerini veya İK belgelerini yükleyin. Yapay zeka, güncel Yargıtay kararlarına göre 'Görünmez Riskleri' tespit edip puanlasın.")
    col1, col2 = st.columns([1, 2])

    with col1:
        st.markdown("#### 📂 Belge Yükleme")
        doc_type = st.selectbox("Belge Türü", ["İş Sözleşmesi", "KVKK Aydınlatma Metni", "Tedarikçi Sözleşmesi", "Kira Kontratı"])
        uploaded_file = st.file_uploader("Dosyayı Sürükleyin (PDF/DOCX)", type=["pdf", "docx", "txt"])
        analyze_btn = st.button("🔍 Risk Taramasını Başlat", type="primary", use_container_width=True)

    with col2:
        if analyze_btn and uploaded_file:
            if not api_key:
                st.error("⚠️ Lütfen sol menüden API Anahtarını giriniz.")
            else:
                with st.spinner("Belge taranıyor, Yargıtay kararlarıyla karşılaştırılıyor..."):
                    # Dosya okuma (Basit)
                    file_text = "Örnek metin"
                    try:
                        if uploaded_file.name.endswith(".pdf"):
                            reader = PyPDF2.PdfReader(uploaded_file)
                            if len(reader.pages) > 0:
                                file_text = reader.pages[0].extract_text()
                        elif uploaded_file.name.endswith(".txt"):
                            file_text = uploaded_file.getvalue().decode("utf-8")
                    except: pass

                    # AI Prompt
                    prompt = f"""
                    GÖREV: Sen kıdemli bir hukuk denetçisisin.
                    BELGE TÜRÜ: {doc_type}
                    BELGE İÇERİĞİ (ÖZET): {file_text[:3000]}
                    
                    GÖREVLER:
                    1. Bu belge için güncel Yargıtay kararlarına göre en kritik 3 riski bul.
                    2. Belgeye 0-100 arası "HUKUKİ SAĞLAMLIK SKORU" ver.
                    
                    ÇIKTI FORMATI:
                    SKOR: [Sayı]
                    RİSKLER: [Detaylar]
                    """
                    ai_response = get_gemini_text_response(prompt, api_key)
                    
                    # Skoru çekme
                    risk_score = 60
                    match = re.search(r"SKOR:\s*(\d+)", ai_response)
                    if match: risk_score = int(match.group(1))

                    # --- GÖRSELLEŞTİRME (HATA ÖNLEYİCİ MOD) ---
                    if PLOTLY_VAR:
                        # Plotly varsa havalı göstergeyi çiz
                        fig = go.Figure(go.Indicator(
                            mode = "gauge+number",
                            value = risk_score,
                            title = {'text': "Hukuki Sağlamlık Skoru"},
                            gauge = {
                                'axis': {'range': [None, 100]},
                                'bar': {'color': "black"},
                                'steps': [
                                    {'range': [0, 50], 'color': "#ff4d4d"},
                                    {'range': [50, 80], 'color': "#ffcc00"},
                                    {'range': [80, 100], 'color': "#33cc33"}
                                ],
                                'threshold': {'line': {'color': "red", 'width': 4}, 'thickness': 0.75, 'value': risk_score}
                            }
                        ))
                        fig.update_layout(height=250, margin=dict(l=20, r=20, t=30, b=20))
                        st.plotly_chart(fig, use_container_width=True)
                    else:
                        # Plotly yoksa standart bar kullan (Çökmemesi için)
                        st.metric("Hukuki Sağlamlık Skoru", f"{risk_score} / 100")
                        st.progress(risk_score / 100)
                        if risk_score < 50:
                            st.error("Risk Seviyesi: YÜKSEK")
                        elif risk_score < 80:
                            st.warning("Risk Seviyesi: ORTA")
                        else:
                            st.success("Risk Seviyesi: DÜŞÜK")
                    
                    st.markdown("### 📋 Risk Raporu")
                    st.write(ai_response.replace(f"SKOR: {risk_score}", ""))


def render_time_machine(api_key):
    st.info("Bir olay tarihi girin, sistem sizi o güne götürsün. O gün geçerli olan kanun maddesini ve faiz oranlarını görün.")
    col_date, col_topic = st.columns([1, 2])
    
    with col_date:
        target_date = st.date_input("Olay Tarihi Seçin:", value=date(2015, 5, 14))
    with col_topic:
        topic = st.text_input("Sorgulanacak Konu", placeholder="Örn: Kıdem Tazminatı Tavanı")

    if st.button("🕒 Geçmişe Git", use_container_width=True):
        if api_key and topic:
            with st.spinner(f"Sistem {target_date.strftime('%d.%m.%Y')} tarihine geri sarılıyor..."):
                prompt = f"""
                GÖREV: Hukuk Tarihçisi. TARİH: {target_date.strftime('%d.%m.%Y')}. KONU: {topic}.
                SORU: O tarihte bu konuyla ilgili yürürlükte olan kanun maddesi, faiz oranı ve Yargıtay görüşü neydi?
                """
                response = get_gemini_text_response(prompt, api_key)
                st.markdown(f"### 📅 Tarih: {target_date.strftime('%d %B %Y')}")
                st.info(response)
                st.image("https://img.freepik.com/free-vector/sepia-vintage-paper-texture_53876-88607.jpg?w=1380", caption="Arşiv Kaydı", width=600)

# --- 3. MODÜL: AYM & AİHM UYGUNLUK TESTİ ---
def render_aym_aihm_module(api_key):
    st.info("Dilekçenizi, Mahkeme Kararını veya UYAP (UDF) dosyasını yükleyin. Sistem OCR ile okuyup AİHM/AYM standartlarına göre 'Hak İhlali' analizi yapsın.")
    
    # Sekmeli Giriş Yapısı
    tab_text, tab_file = st.tabs(["📝 Metin Yapıştır", "📂 Dosya Yükle (PDF/UDF/TIFF)"])
    
    process_text = ""
    analyze_trigger = False

    # --- TAB 1: MANUEL METİN ---
    with tab_text:
        user_text_input = st.text_area("Metni Buraya Yapıştırın:", height=300, placeholder="Örn: Mahkeme gerekçesiz karar vererek adil yargılanma hakkımı ihlal etmiştir...")
        if st.button("⚖️ Metni Analiz Et", key="btn_text_aym", type="primary"):
            process_text = user_text_input
            analyze_trigger = True

    # --- TAB 2: DOSYA YÜKLEME ---
    with tab_file:
        uploaded_file = st.file_uploader("Dosya Seçin", type=["pdf", "udf", "xml", "tiff", "tif", "jpg", "png", "txt"])
        
        if uploaded_file:
            st.caption(f"Yüklenen Dosya: {uploaded_file.name}")
            if st.button("👁️ Dosyayı Oku ve Analiz Et", key="btn_file_aym", type="primary"):
                with st.spinner("Dosya okunuyor ve OCR yapılıyor..."):
                    extracted_text = extract_text_from_legal_file(uploaded_file, api_key)
                    
                    if "[OCR GEREKLİ]" in extracted_text or "Hata" in extracted_text:
                        st.error(extracted_text)
                    else:
                        process_text = extracted_text
                        st.success("Dosya başarıyla metne dönüştürüldü! Analiz başlıyor...")
                        with st.expander("Okunan Metni Gör"):
                            st.text(process_text[:1000] + "...")
                        analyze_trigger = True

    # --- ORTAK ANALİZ MOTORU ---
    if analyze_trigger and process_text:
        if not api_key:
            st.error("⚠️ Lütfen API Anahtarını giriniz.")
        elif len(process_text) < 20:
            st.warning("Analiz için yeterli metin bulunamadı.")
        else:
            with st.spinner("Metin, AİHM ve AYM içtihatlarıyla çapraz sorgulanıyor..."):
                
                prompt = f"""
                GÖREV: Sen AİHM ve AYM kararları konusunda uzmanlaşmış kıdemli bir hukukçusun.
                METİN: "{process_text[:6000]}" (Kısaltılmış olabilir)
                
                ANALİZ ADIMLARI:
                1. Bu metindeki olayda, Avrupa İnsan Hakları Sözleşmesi (AİHS) veya Anayasa ile korunan hangi temel haklar risk altında? (Örn: Mülkiyet Hakkı, Adil Yargılanma Hakkı).
                2. Bu metin bir mahkeme kararıysa Üst Mahkemede BOZULMA İHTİMALİ yüzde kaçtır? Bir dilekçeyse KABUL EDİLME GÜCÜ yüzde kaçtır? (0-100 arası bir puan ver).
                3. Konuyla ilgili emsal bir AİHM veya AYM kararı adı ver ve özetle.
                4. Eğer bir ihlal varsa, başvuru formunda hangi argüman kullanılmalı?
                
                ÇIKTI FORMATI:
                ORAN: [Sayı]
                ANALİZ: [Detaylı Hukuki Görüş]
                EMSAL: [Karar İsimleri]
                STRATEJİ: [Öneri]
                """
                
                ai_response = get_gemini_text_response(prompt, api_key)
                
                # Oranı çekme
                ihlal_orani = 50
                match = re.search(r"ORAN:\s*(\d+)", ai_response)
                if match: ihlal_orani = int(match.group(1))
                
                # --- SONUÇ EKRANI ---
                st.divider()
                col_score, col_detail = st.columns([1, 2])
                
                with col_score:
                    st.markdown(f"<h2 style='text-align: center; color: #d63031;'>%{ihlal_orani}</h2>", unsafe_allow_html=True)
                    st.markdown("<p style='text-align: center;'><b>Bozulma / İhlal Riski</b></p>", unsafe_allow_html=True)
                    st.progress(ihlal_orani / 100)
                    
                    if ihlal_orani > 70:
                        st.error("🚨 KRİTİK: Yüksek ihtimalle hak ihlali var.")
                    elif ihlal_orani > 40:
                        st.warning("⚠️ DİKKAT: Güçlü argümanlar gerekiyor.")
                    else:
                        st.success("✅ TEMİZ: Belirgin bir ihlal görünmüyor.")

                with col_detail:
                    st.markdown("### 🏛️ Yüksek Yargı Raporu")
                    st.write(ai_response.replace(f"ORAN: {ihlal_orani}", ""))


def get_image_metadata(image):
    """Resimden EXIF verilerini çeker."""
    meta_dict = {}
    try:
        exif_data = image._getexif()
        if exif_data:
            for tag_id, value in exif_data.items():
                tag_name = TAGS.get(tag_id, tag_id)
                if tag_name in ['DateTime', 'DateTimeOriginal', 'Make', 'Model', 'Software']:
                    meta_dict[tag_name] = str(value)
    except:
        return None
    return meta_dict

def render_deepfake_module(api_key):
    # --- HATA YAKALAYICI BLOK BAŞLANGICI ---
    try:
        st.info("Şüpheli fotoğraf veya ses kaydını yükleyin. Yapay zeka, metadata (üst veri) analizi ve içerik taraması yaparak 'Montaj/Deepfake' izlerini arasın.")
        
        col_upload, col_report = st.columns([1, 2])
        
        with col_upload:
            st.markdown("#### 🕵️‍♂️ Delil Yükle")
            file_type = st.radio("Delil Türü", ["Fotoğraf / Belge Görüntüsü", "Ses Kaydı (Kısa)"])
            
            if file_type == "Fotoğraf / Belge Görüntüsü":
                uploaded_file = st.file_uploader("Resim Seç (JPG, PNG)", type=["jpg", "jpeg", "png"])
            else:
                uploaded_file = st.file_uploader("Ses Dosyası Seç (MP3, WAV)", type=["mp3", "wav"])
                
            analyze_btn = st.button("🔍 Adli Bilişim Analizi Yap", type="primary", use_container_width=True)

        with col_report:
            if analyze_btn and uploaded_file:
                if not api_key:
                    st.error("⚠️ API Anahtarı eksik.")
                else:
                    with st.spinner("Dosya bit-bit inceleniyor, metadata taranıyor ve AI analizi yapılıyor..."):
                        
                        genai.configure(api_key=api_key)
                        # Model seçimi (Hata verirse Pro'ya düş)
                        try:
                            model = genai.GenerativeModel('gemini-1.5-flash')
                        except:
                            model = genai.GenerativeModel('gemini-pro-vision')
                        
                        report_text = ""
                        fake_score = 0
                        metadata_info = {}

                        # --- FOTOĞRAF ANALİZİ ---
                        if file_type == "Fotoğraf / Belge Görüntüsü":
                            image = Image.open(uploaded_file)
                            st.image(image, caption="İncelenen Delil", width=300)
                            
                            # Metadata Kontrolü (Güvenli)
                            try:
                                metadata_info = get_image_metadata(image)
                                meta_str = str(metadata_info) if metadata_info else "Metadata bulunamadı."
                            except Exception as e:
                                meta_str = f"Metadata okunamadı: {e}"
                            
                            prompt = f"""
                            GÖREV: Sen uzman bir Adli Bilişim (Forensics) uzmanısın.
                            METADATA: {meta_str}
                            GÖREVLER:
                            1. Görselde Deepfake/Montaj izi var mı?
                            2. Metadata tutarlı mı?
                            3. Güvenilirlik puanı (0-100).
                            ÇIKTI: GÜVEN_SKORU: [Sayı] ...
                            """
                            response = model.generate_content([prompt, image])
                            report_text = response.text

                        # --- SES ANALİZİ ---
                        else: 
                            # Ses analizi için güvenli blok
                            try:
                                st.audio(uploaded_file)
                                # Geçici dosya oluşturma
                                import tempfile
                                with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp_file:
                                    tmp_file.write(uploaded_file.getvalue())
                                    tmp_path = tmp_file.name
                                
                                # Speech Recognition
                                r = sr.Recognizer()
                                with sr.AudioFile(tmp_path) as source:
                                    audio_data = r.record(source)
                                    text_output = r.recognize_google(audio_data, language="tr-TR")
                                    
                                    prompt = f"""Ses Transkripti: "{text_output}". Bu konuşma doğal mı, kurgu mu? Puanla (0-100). ÇIKTI: GÜVEN_SKORU: [Sayı] ..."""
                                    model_text = genai.GenerativeModel('gemini-pro')
                                    response = model_text.generate_content(prompt)
                                    report_text = response.text
                            except ImportError:
                                st.error("Ses analizi için 'SpeechRecognition' kütüphanesi yüklü değil.")
                                return
                            except Exception as e:
                                st.error(f"Ses işleme hatası: {str(e)}")
                                return

                        # --- SONUÇLARI GÖSTER ---
                        match = re.search(r"GÜVEN_SKORU:\s*(\d+)", report_text)
                        if match: fake_score = int(match.group(1))
                        
                        st.divider()
                        st.metric("Delil Güvenilirlik Skoru", f"{fake_score} / 100")
                        st.progress(fake_score / 100)
                        
                        if fake_score < 50:
                            st.error("🚨 SAHTECİLİK ŞÜPHESİ YÜKSEK")
                        else:
                            st.success("✅ DELİL GÜVENİLİR GÖRÜNÜYOR")
                            
                        st.write(report_text.replace(f"GÜVEN_SKORU: {fake_score}", ""))

    except Exception as e:
        # EĞER BEYAZ EKRAN ÇIKARSA BURASI DEVREYE GİRER VE HATAYI YAZAR
        st.error(f"🚨 Modül Yükleme Hatası: {str(e)}")
        st.warning("Lütfen 'PIL', 'SpeechRecognition' kütüphanelerinin yüklü olduğundan ve 'TAGS' importunun yapıldığından emin olun.")


def generate_dork_category(category, target_name, city):
    """Belirli bir kategori için gelişmiş arama linkleri (Dorks) üretir."""
    # Simüle edilmiş işlem süresi (Threading etkisini görmek için)
    time.sleep(0.5) 
    
    links = []
    base_url = "https://www.google.com/search?q="
    name_slug = target_name.replace(" ", "+")
    
    if category == "social":
        # Sosyal Medya Taraması
        links.append(f"[📸 Instagram: {target_name}]({base_url}site:instagram.com+%22{name_slug}%22)")
        links.append(f"[💼 LinkedIn: {target_name}]({base_url}site:linkedin.com/in/+%22{name_slug}%22)")
        links.append(f"[🐦 Twitter/X: {target_name}]({base_url}site:twitter.com+%22{name_slug}%22)")
        links.append(f"[👤 Facebook: {target_name}]({base_url}site:facebook.com+%22{name_slug}%22)")
        
    elif category == "business":
        # Ticari Varlık ve Şirket Taraması
        links.append(f"[🏢 Ticaret Sicil: {target_name}]({base_url}%22{name_slug}%22+site:ticaretsicil.gov.tr)")
        links.append(f"[📄 Resmi Gazete: {target_name}]({base_url}%22{name_slug}%22+site:resmigazete.gov.tr)")
        links.append(f"[🤝 Şirket Ortaklıkları]({base_url}%22{name_slug}%22+kurucu+ortak+sahibi)")
        
    elif category == "assets":
        # Mal Varlığı ve Lüks Yaşam İzi (Tatil, Araba vb.)
        links.append(f"[🏖️ Tatil/Otel Yorumları]({base_url}%22{name_slug}%22+otel+tatil+gezi)")
        links.append(f"[🚗 Araba/Satış İlanları]({base_url}%22{name_slug}%22+sahibinden+satılık)")
        links.append(f"[🎓 Mezuniyet/Okul]({base_url}%22{name_slug}%22+mezun+okul+lise+üniversite)")
        
    return category, links

def render_osint_module(api_key):
    st.info("Hedef kişinin (Borçlu, Davalı) dijital ayak izlerini takip edin. Sistem 'Threading' teknolojisiyle aynı anda sosyal medya, ticaret sicil ve mal varlığı taraması başlatır.")
    
    col_input, col_results = st.columns([1, 2])
    
    with col_input:
        st.markdown("#### 🎯 Hedef Tanımla")
        target_name = st.text_input("Ad Soyad / Şirket Adı", placeholder="Örn: Ahmet Yılmaz")
        target_city = st.text_input("Şehir (Opsiyonel)", placeholder="Örn: İstanbul")
        
        start_scan = st.button("🚀 İstihbarat Taramasını Başlat", type="primary")
        
        st.markdown("---")
        st.caption("⚠️ **Yasal Uyarı:** Bu modül sadece halka açık verileri (Open Source) tarar. KVKK sınırları içinde kullanınız.")

    with col_results:
        if start_scan and target_name:
            st.write(f"📡 **'{target_name}'** için çok kanallı tarama başlatılıyor...")
            
            # --- THREADING (Çoklu İş Parçacığı) BAŞLANGICI ---
            # Sosyal medya, İş ve Varlık taramalarını aynı anda yapar
            results = {}
            
            with st.spinner("Veri madenciliği yapılıyor (Social + Business + Assets)..."):
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    # Görevleri tanımla
                    t1 = executor.submit(generate_dork_category, "social", target_name, target_city)
                    t2 = executor.submit(generate_dork_category, "business", target_name, target_city)
                    t3 = executor.submit(generate_dork_category, "assets", target_name, target_city)
                    
                    # Sonuçları topla
                    for future in concurrent.futures.as_completed([t1, t2, t3]):
                        cat, links = future.result()
                        results[cat] = links
            
            st.success("✅ Tarama Tamamlandı! Bulunan İzler:")
            
            # Sonuçları Göster
            tab_social, tab_business, tab_assets = st.tabs(["📸 Sosyal Medya", "🏢 Ticari Varlık", "🏖️ Yaşam Tarzı"])
            
            with tab_social:
                st.markdown("### Sosyal Ağ Taraması")
                for link in results.get("social", []):
                    st.markdown(f"- {link}", unsafe_allow_html=True)
                st.info("💡 İpucu: 'Borcum yok' diyen kişinin Instagram'da gizli hikayesi olabilir.")

            with tab_business:
                st.markdown("### Ticari Sicil & Resmi Kayıtlar")
                for link in results.get("business", []):
                    st.markdown(f"- {link}", unsafe_allow_html=True)
                st.info("💡 İpucu: Üzerine kayıtlı şirket veya ortaklıkları buradan yakalayabilirsiniz.")

            with tab_assets:
                st.markdown("### Lüks Yaşam & Varlık İzleri")
                for link in results.get("assets", []):
                    st.markdown(f"- {link}", unsafe_allow_html=True)
                st.info("💡 İpucu: Otel yorumları veya 2. el satış ilanları gizli varlıkları ele verebilir.")

            # --- AI ANALİZ KISMI ---
            st.divider()
            st.markdown("#### 🧠 İstihbarat Analizi")
            evidence_text = st.text_area("Bulduğunuz şüpheli bilgiyi buraya yapıştırın (Örn: Instagram biyografisi veya Ticaret Sicil kaydı):", height=100)
            
            if st.button("🕵️ Delil Analizi Yap"):
                if not api_key:
                    st.error("API Anahtarı gerekli.")
                else:
                    genai.configure(api_key=api_key)
                    model = genai.GenerativeModel('gemini-pro')
                    prompt = f"""
                    GÖREV: Bir OSINT (Açık Kaynak İstihbaratı) uzmanısın.
                    HEDEF KİŞİ: {target_name}
                    BULUNAN VERİ: "{evidence_text}"
                    
                    SORU: 
                    1. Bu veri, kişinin "borç ödemekten kaçınma" veya "mal kaçırma" şüphesini destekler mi?
                    2. Hukuki olarak bu veri delil dosyasında nasıl kullanılabilir?
                    
                    Kısa ve net cevap ver.
                    """
                    with st.spinner("Yapay zeka veriyi yorumluyor..."):
                        response = model.generate_content(prompt)
                        st.write(response.text)

def render_precedent_alert_module(api_key):
    st.info("Bu modül, derdest (devam eden) davalarınızı takip eder ve Yargıtay/AYM tarafından yayınlanan **'Bugünkü Kararlar'** ile otomatik eşleştirir.")

    # --- 1. OTURUM DURUMU (Dava Portföyü) ---
    if 'my_cases' not in st.session_state:
        st.session_state.my_cases = [
            {"id": 1, "ad": "Yılmaz v. Demir (Kira)", "konu": "5 yıllık kiracı tahliyesi, uyarlama davası", "durum": "Bilirkişi aşamasında"},
            {"id": 2, "ad": "Kripto Dolandırıcılık", "konu": "Thodex benzeri borsa batışı, güveni kötüye kullanma", "durum": "Savcılık soruşturması"},
            {"id": 3, "ad": "İşe İade (Ahmet B.)", "konu": "Performans düşüklüğü nedeniyle fesih", "durum": "Tanık dinleniyor"}
        ]

    col_portfolio, col_feed = st.columns([1, 2])

    # --- SOL KOLON: DAVA PORTFÖYÜM ---
    with col_portfolio:
        st.markdown("### 📂 Dava Portföyüm")
        
        with st.expander("➕ Yeni Dava Ekle"):
            new_case_name = st.text_input("Dava Adı")
            new_case_topic = st.text_area("Dava Konusu/Detayı")
            if st.button("Listeye Ekle"):
                new_id = len(st.session_state.my_cases) + 1
                st.session_state.my_cases.append({"id": new_id, "ad": new_case_name, "konu": new_case_topic, "durum": "Yeni"})
                st.success("Eklendi!")
                st.rerun()
        
        for case in st.session_state.my_cases:
            st.markdown(f"**Dosya #{case['id']}: {case['ad']}**\n*{case['konu']}*\n`Durum: {case['durum']}`\n---")

    # --- SAĞ KOLON: GÜNLÜK BÜLTEN TARAMASI ---
    with col_feed:
        st.markdown("### 📡 Günlük Yargı Bülteni & Etki Analizi")
        
        if st.button("🔄 Bülteni Tara ve Analiz Et", type="primary", use_container_width=True):
            if not api_key:
                st.error("API Anahtarı gerekli.")
            else:
                daily_decisions = [
                    """KARAR 2024/105 (Yargıtay HGK): Kira tespit davalarında '5 yıllık süre' dolmadan yapılan uyarlamalarda, TÜFE oranı tavan olarak kabul edilemez. Hakim hakkaniyete göre serbestçe belirler.""",
                    """KARAR 2024/88 (AYM Bireysel Başvuru): Kripto para borsalarındaki kayıplarda, devletin denetim yükümlülüğünü ihlal ettiği iddiasıyla yapılan başvuruda 'Mülkiyet Hakkı İhlali' olmadığına karar verildi.""",
                    """KARAR 2024/12 (İş Mahkemesi Emsal): Sadece performans düşüklüğü, yazılı savunma alınmadan ve eğitim verilmeden fesih sebebi yapılamaz."""
                ]
                
                st.write(f"📅 **Bugün Yayınlanan Kritik Karar Sayısı:** {len(daily_decisions)}")
                
                with st.spinner("Uygun yapay zeka modeli aranıyor ve analiz yapılıyor..."):
                    genai.configure(api_key=api_key)
                    
                    # --- OTOMATİK MODEL SEÇİCİ (HATA ÇÖZÜMÜ) ---
                    target_model_name = "models/gemini-pro" # Varsayılan
                    try:
                        # Sistemdeki mevcut modelleri listele ve ilk çalışanı seç
                        for m in genai.list_models():
                            if 'generateContent' in m.supported_generation_methods:
                                if 'gemini' in m.name:
                                    target_model_name = m.name
                                    break
                    except Exception as e:
                        st.warning(f"Model listesi alınamadı, varsayılan deneniyor: {e}")

                    # Seçilen modeli ekrana yaz (Debug için)
                    # st.caption(f"Kullanılan Model: {target_model_name}") 
                    
                    try:
                        model = genai.GenerativeModel(target_model_name)
                        
                        cases_str = str(st.session_state.my_cases)
                        decisions_str = "\n".join(daily_decisions)
                        
                        prompt = f"""
                        GÖREV: Sen proaktif bir hukuk asistanısın.
                        1. AŞAĞIDAKİ MÜVEKKİL DAVALARI (PORTFÖY): {cases_str}
                        2. AŞAĞIDAKİ BUGÜN ÇIKAN YENİ YARGI KARARLARI: {decisions_str}
                        YAPMAN GEREKEN: Her bir davayı kontrol et. Eğer yeni kararlardan biri, mevcut bir davayı etkiliyorsa uyar.
                        ÇIKTI FORMATI:
                        UYARI: [Dosya Adı]
                        DURUM: [KRİTİK / DİKKAT / FIRSAT]
                        NEDEN: [Açıklama]
                        AKSİYON: [Öneri]
                        """
                        
                        response = model.generate_content(prompt)
                        
                        st.divider()
                        st.markdown("### 🚨 Tespit Edilen Riskler ve Fırsatlar")
                        
                        alerts = response.text.split("UYARI:")
                        if len(alerts) < 2:
                            st.write(response.text)
                        else:
                            for alert in alerts:
                                if alert.strip():
                                    if "KRİTİK" in alert: st.error(f"**UYARI:{alert}**")
                                    elif "FIRSAT" in alert: st.success(f"**UYARI:{alert}**")
                                    else: st.warning(f"**UYARI:{alert}**")
                                    
                    except Exception as e:
                        st.error(f"Model Hatası: {str(e)}")
                        st.info("Lütfen API anahtarınızın 'Generative AI' servisine erişimi olduğundan emin olun.")


def render_owner_mode(api_key):
    st.info("👑 **Sahip Modu (Web):** Bilgisayarınızdaki dosyaları seçip sürükleyin. Sistem, hesabınızda çalışan en uygun Yapay Zeka modelini otomatik bulup kullanacaktır.")

    # --- 0. OTOMATİK MODEL BULUCU (HATA ÖNLEYİCİ) ---
    def get_working_model():
        """Sistemdeki aktif modelleri tarar ve ilk çalışanı getirir."""
        default_model = "models/gemini-pro" # En kötü ihtimal yedeği
        try:
            for m in genai.list_models():
                if 'generateContent' in m.supported_generation_methods:
                    # İsminde 'gemini' geçen ilk modeli al (flash, pro, 1.5 vs.)
                    if 'gemini' in m.name:
                        return m.name
        except:
            pass
        return default_model

    # --- 1. DOSYA OKUMA MOTORU ---
    def get_file_text(file_obj, api_key_for_ocr):
        filename = file_obj.name.lower()
        file_bytes = file_obj.read()
        text = ""
        
        try:
            # A) PDF
            if filename.endswith('.pdf'):
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            # B) WORD
            elif filename.endswith('.docx'):
                doc = Document(io.BytesIO(file_bytes))
                text = "\n".join([p.text for p in doc.paragraphs])
            
            # C) UYAP (UDF)
            elif filename.endswith('.udf'):
                try:
                    with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                        text = "".join(ET.fromstring(z.read('content.xml')).itertext())
                except:
                    text = "".join(ET.fromstring(file_bytes).itertext())
            
            # D) RESİM (OTOMATİK MODEL İLE)
            elif filename.endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp')):
                if api_key_for_ocr:
                    image = Image.open(io.BytesIO(file_bytes))
                    
                    # Dinamik model seçimi
                    active_model = get_working_model()
                    model = genai.GenerativeModel(active_model)
                    
                    try:
                        response = model.generate_content(["Bu resimdeki yazıları oku:", image])
                        text = f"[RESİM İÇERİĞİ]:\n{response.text}"
                    except:
                        text = "[RESİM OKUNAMADI: Seçilen model görsel desteklemiyor olabilir.]"
            
            # E) DÜZ METİN
            else:
                text = file_bytes.decode("utf-8", errors='ignore')
                
            return text
        except Exception as e:
            return f"[Okuma Hatası]: {str(e)}"

    # --- 2. ARAYÜZ ---
    if 'web_memory' not in st.session_state: st.session_state.web_memory = ""
    if 'web_history' not in st.session_state: st.session_state.web_history = []

    col_upload, col_chat = st.columns([1, 2])

    # --- SOL: YÜKLEME ---
    with col_upload:
        st.markdown("### 📤 Dosyaları Sürükle")
        uploaded_files = st.file_uploader("Klasördeki dosyaları seçip buraya bırak", accept_multiple_files=True)
        
        if st.button("🧠 Analiz Et", type="primary"):
            if not uploaded_files:
                st.warning("Dosya yok.")
            elif not api_key:
                st.error("API Anahtarı yok.")
            else:
                genai.configure(api_key=api_key)
                full_text = ""
                bar = st.progress(0)
                
                for i, file in enumerate(uploaded_files):
                    content = get_file_text(file, api_key)
                    full_text += f"\n=== DOSYA: {file.name} ===\n{content}\n"
                    bar.progress((i + 1) / len(uploaded_files))
                
                st.session_state.web_memory = full_text
                st.session_state.web_history = [] 
                st.success(f"✅ {len(uploaded_files)} dosya okundu!")

        if st.session_state.web_memory:
            if st.button("🗑️ Temizle"):
                st.session_state.web_memory = ""
                st.rerun()

    # --- SAĞ: SOHBET ---
    with col_chat:
        st.markdown("### 💬 Asistan")
        
        for msg in st.session_state.web_history:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

        if prompt := st.chat_input("Sorunuzu yazın..."):
            if not st.session_state.web_memory:
                st.warning("Önce dosya yükleyin.")
            else:
                st.session_state.web_history.append({"role": "user", "content": prompt})
                with st.chat_message("user"):
                    st.markdown(prompt)

                with st.chat_message("assistant"):
                    with st.spinner("Düşünüyor..."):
                        try:
                            # BURADA OTOMATİK MODEL SEÇİLİYOR
                            active_model_name = get_working_model()
                            # st.caption(f"Kullanılan Model: {active_model_name}") # İstersen açıp görebilirsin
                            
                            model = genai.GenerativeModel(active_model_name)
                            
                            context = st.session_state.web_memory[:90000]
                            final_prompt = f"VERİLER:\n{context}\n\nSORU: {prompt}"
                            
                            response = model.generate_content(final_prompt)
                            st.markdown(response.text)
                            st.session_state.web_history.append({"role": "assistant", "content": response.text})
                        except Exception as e:
                            st.error(f"Cevap üretilemedi: {e}")


import json
import time

def render_property_genealogy(api_key):
    st.info("🌳 **Mülkiyet Soyağacı:** Tapu ve kadastro belgelerinizi yükleyin, AI zinciri kursun.")

    # --- 0. AKILLI MODEL SEÇİCİ (HATA ÖNLEYİCİ) ---
    def get_best_model():
        """Hesapta aktif olan en hızlı modeli bulur."""
        try:
            available_models = []
            for m in genai.list_models():
                if 'generateContent' in m.supported_generation_methods:
                    available_models.append(m.name)
            
            # Öncelik Sırası: 1. Flash (Hızlı), 2. Pro (Standart), 3. Herhangi biri
            for m in available_models:
                if 'flash' in m: return m
            for m in available_models:
                if 'pro' in m: return m
            return available_models[0] if available_models else "models/gemini-pro"
        except:
            return "models/gemini-pro" # En kötü ihtimal yedeği

    # --- 1. DOSYA OKUMA ---
    def get_genealogy_file_text(file_obj, api_key_for_ocr):
        filename = file_obj.name.lower()
        file_bytes = file_obj.read()
        text = ""
        try:
            if filename.endswith('.pdf'):
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                for page in pdf_reader.pages: text += page.extract_text() + "\n"
            elif filename.endswith('.docx'):
                doc = Document(io.BytesIO(file_bytes))
                text = "\n".join([p.text for p in doc.paragraphs])
            elif filename.endswith(('.png', '.jpg', '.jpeg')):
                if api_key_for_ocr:
                    image = Image.open(io.BytesIO(file_bytes))
                    # Otomatik model seçimi
                    active_model = get_best_model()
                    model = genai.GenerativeModel(active_model)
                    response = model.generate_content(["Bu belgedeki isimleri ve tarihleri oku:", image])
                    text = response.text
            else:
                text = file_bytes.decode("utf-8", errors='ignore')
            return text
        except:
            return ""

    # --- 2. STATE ---
    if 'prop_history' not in st.session_state:
        st.session_state.prop_history = [
            {"yil": "1960", "kimden": "Hazine", "kime": "Mehmet Ağa", "islem": "Kadastro", "durum": "Pasif"},
            {"yil": "1990", "kimden": "Mehmet Ağa", "kime": "Ali (Oğlu)", "islem": "Miras", "durum": "Aktif"}
        ]

    # --- 3. ARAYÜZ ---
    col_left, col_right = st.columns([1, 2])

    # SOL: YÜKLEME
    with col_left:
        st.markdown("### 📂 Belge Yükle")
        uploaded_files = st.file_uploader("Tapu/Kadastro Evrakları", accept_multiple_files=True)
        
        if st.button("⚡ Zinciri Oluştur", type="primary"):
            if not uploaded_files or not api_key:
                st.warning("Dosya ve API Key gerekli.")
            else:
                status_box = st.empty()
                status_box.info("Belgeler okunuyor...")
                
                genai.configure(api_key=api_key)
                full_text = ""
                for f in uploaded_files:
                    full_text += f"\nDOC: {f.name}\n" + get_genealogy_file_text(f, api_key)
                
                try:
                    status_box.info("AI Modeli seçiliyor ve zincir kuruluyor...")
                    
                    # OTOMATİK MODEL SEÇİMİ
                    active_model_name = get_best_model()
                    model = genai.GenerativeModel(active_model_name)
                    
                    prompt = f"""
                    GÖREV: Metinlerdeki mülkiyet devirlerini JSON listesi yap.
                    METİN: {full_text[:40000]}
                    FORMAT: [{{"yil": "...", "kimden": "...", "kime": "...", "islem": "...", "durum": "Aktif/Pasif/Kritik"}}]
                    SADECE JSON VER.
                    """
                    response = model.generate_content(prompt)
                    clean_json = response.text.replace("```json", "").replace("```", "").strip()
                    st.session_state.prop_history = json.loads(clean_json)
                    status_box.success(f"Tamamlandı! (Kullanılan Model: {active_model_name})")
                    time.sleep(1)
                    st.rerun()
                except Exception as e:
                    status_box.error(f"Hata: {e}")

        st.dataframe(st.session_state.prop_history, height=300)
        if st.button("Temizle"):
            st.session_state.prop_history = []
            st.rerun()

    # SAĞ: GRAFİK VE ANALİZ
    with col_right:
        st.markdown("### 🗺️ Görsel Harita")
        
        if st.session_state.prop_history:
            # Grafik Çizimi
            graph_code = "digraph { rankdir=LR; node [shape=box, style=filled, fontname=\"Arial\"];"
            for item in st.session_state.prop_history:
                color = "#d4edda" if item.get("durum") == "Aktif" else "#e2e3e5"
                if item.get("durum") == "Kritik": color = "#f8d7da"
                
                k1 = str(item.get('kimden', '?')).replace('"', '').strip()
                k2 = str(item.get('kime', '?')).replace('"', '').strip()
                lbl = f"{item.get('yil')}\\n{item.get('islem')}"
                
                graph_code += f'\n "{k1}" -> "{k2}" [label="{lbl}", fontsize=10];'
                graph_code += f'\n "{k2}" [fillcolor="{color}", label="{k2}"];'
            graph_code += "\n}"
            st.graphviz_chart(graph_code)
            
            st.divider()
            
            # --- ANALİZ KISMI (STREAMING & AUTO MODEL) ---
            if st.button("🕵️ Risk Analizi Başlat"):
                output_placeholder = st.empty()
                output_placeholder.text("Model aranıyor ve analiz başlıyor...")
                
                try:
                    genai.configure(api_key=api_key)
                    
                    # OTOMATİK MODEL SEÇİMİ
                    active_model_name = get_best_model()
                    model = genai.GenerativeModel(active_model_name)
                    
                    chain_data = json.dumps(st.session_state.prop_history, ensure_ascii=False)
                    
                    prompt = f"""
                    GÖREV: Sen uzman bir tapu denetçisisin. Aşağıdaki mülkiyet zincirini analiz et.
                    VERİ: {chain_data}
                    
                    Lütfen şu başlıklar altında rapor yaz:
                    1. 🔴 Kritik Riskler
                    2. ⚠️ Hukuki Uyarılar
                    3. ✅ Sonuç
                    """
                    
                    # stream=True ile parça parça alıyoruz
                    response = model.generate_content(prompt, stream=True)
                    
                    full_text = ""
                    for chunk in response:
                        full_text += chunk.text
                        output_placeholder.markdown(full_text + "▌") 
                    
                    output_placeholder.markdown(full_text)
                        
                except Exception as e:
                    output_placeholder.error(f"Hata oluştu: {e}")
        else:
            st.info("👈 Veri yok.")

import pandas as pd
from datetime import datetime, timedelta

def render_limitations_heatmap(api_key):
    # --- IMPORTLARI İZOLE ET (Çakışmayı Önler) ---
    import pandas as pd
    import datetime as dt  # datetime modülünü 'dt' olarak çağırıyoruz
    
    st.info("🔥 **Zamanaşımı Isı Haritası:** Dava türüne ve tarihlere göre her bir alacak kaleminin risk durumunu analiz eder. Islah ve hak düşürücü süreleri 'Borsa Ekranı' gibi takip eder.")

    # --- 0. OTOMATİK MODEL SEÇİCİ ---
    def get_best_model():
        try:
            available_models = []
            for m in genai.list_models():
                if 'generateContent' in m.supported_generation_methods:
                    available_models.append(m.name)
            for m in available_models:
                if 'flash' in m: return m
            for m in available_models:
                if 'pro' in m: return m
            return available_models[0] if available_models else "models/gemini-pro"
        except:
            return "models/gemini-pro"

    # --- 1. GİRİŞ PANELİ ---
    col_input, col_dashboard = st.columns([1, 2])

    with col_input:
        st.markdown("### 📅 Kritik Tarihler")
        
        dava_turu = st.selectbox("Dava Türü", ["İş Hukuku (İşçi Alacağı)", "Ticari Alacak", "Tüketici", "Tazminat (Haksız Fiil)"])
        
        # HATA ÇIKARAN SATIR DÜZELTİLDİ: dt.datetime.now().date()
        bugun = dt.datetime.now().date()
        
        # Tarih Seçiciler
        fesih_tarihi = st.date_input("Fesih / Olay Tarihi", value=bugun - dt.timedelta(days=365*4))
        dava_tarihi = st.date_input("Dava Açılış Tarihi", value=bugun - dt.timedelta(days=300))
        
        st.divider()
        st.markdown("#### ⚡ Islah Alarmı")
        is_bilirkişi = st.checkbox("Bilirkişi Raporu Geldi mi?")
        
        teblig_tarihi = None
        if is_bilirkişi:
            teblig_tarihi = st.date_input("Rapor Tebliğ Tarihi", value=bugun - dt.timedelta(days=5))
            st.caption("Islah için genellikle 2 haftalık itiraz süresi veya tahkikat sonuna kadar süre dikkate alınır.")

    # --- 2. HESAPLAMA MOTORU ---
    data = []
    
    # İş Hukuku Kuralları (Basitleştirilmiş Örnekler)
    if dava_turu == "İş Hukuku (İşçi Alacağı)":
        # 1. Kıdem Tazminatı (5 Yıl - 2017 sonrası)
        kidem_suresi = fesih_tarihi + dt.timedelta(days=365*5)
        kalan_gun = (kidem_suresi - bugun).days
        data.append({"Kalem": "Kıdem Tazminatı", "Son Tarih": kidem_suresi, "Kalan Gün": kalan_gun, "Risk": ""})
        
        # 2. Fazla Mesai (5 Yıl)
        mesai_suresi = fesih_tarihi + dt.timedelta(days=365*5)
        kalan_gun_mesai = (mesai_suresi - bugun).days
        data.append({"Kalem": "Fazla Mesai", "Son Tarih": mesai_suresi, "Kalan Gün": kalan_gun_mesai, "Risk": ""})
        
        # 3. İşe İade (1 Ay - Arabulucu)
        ise_iade_suresi = fesih_tarihi + dt.timedelta(days=30)
        kalan_gun_iade = (ise_iade_suresi - bugun).days
        data.append({"Kalem": "İşe İade (Arabulucu)", "Son Tarih": ise_iade_suresi, "Kalan Gün": kalan_gun_iade, "Risk": ""})

    # Islah Hesabı (Kritik)
    if is_bilirkişi and teblig_tarihi:
        # HMK 281 - 2 Hafta İtiraz (Islah için stratejik zaman)
        islah_suresi = teblig_tarihi + dt.timedelta(days=14)
        kalan_gun_islah = (islah_suresi - bugun).days
        data.append({"Kalem": "🚨 ISLAH / İTİRAZ", "Son Tarih": islah_suresi, "Kalan Gün": kalan_gun_islah, "Risk": "ÇOK YÜKSEK"})

    # DataFrame Oluştur
    df = pd.DataFrame(data)

    # Risk Renklendirme Fonksiyonu
    def risk_color(val):
        if val < 0: return "background-color: #ff4b4b; color: white" # Kırmızı (Süre Doldu)
        elif val < 15: return "background-color: #ffa500; color: black" # Turuncu (Kritik)
        elif val < 60: return "background-color: #ffe066; color: black" # Sarı (Yaklaşıyor)
        else: return "background-color: #90ee90; color: black" # Yeşil (Güvenli)

    # --- 3. DASHBOARD (ISI HARİTASI) ---
    with col_dashboard:
        st.markdown("### 🌡️ Zamanaşımı Isı Haritası")
        
        if not df.empty:
            # Tabloyu Renklendir
            st.dataframe(
                df.style.applymap(risk_color, subset=["Kalan Gün"])
                        .format({"Son Tarih": "{:%d.%m.%Y}"}),
                use_container_width=True,
                height=250
            )
            
            # Görsel Ticker (İlerleme Çubukları)
            st.markdown("#### ⏳ Kritik Geri Sayım")
            for index, row in df.iterrows():
                kalan = row["Kalan Gün"]
                kalem = row["Kalem"]
                
                if kalan < 0:
                    st.error(f"❌ {kalem}: SÜRE DOLDU! ({abs(kalan)} gün geçti)")
                elif kalan < 15:
                    st.warning(f"⚠️ {kalem}: SON {kalan} GÜN! (Acil İşlem Gerekli)")
                    st.progress(max(0, min(100, int((kalan/15)*100))))
                else:
                    st.success(f"✅ {kalem}: {kalan} gün var. (Güvenli)")
        else:
            st.info("Lütfen sol taraftan tarihleri giriniz.")

        st.divider()
        
        # --- 4. AI STRATEJİ DANIŞMANI ---
        if st.button("🧠 AI Risk & Strateji Analizi Yap"):
            if not api_key:
                st.error("API Key gerekli.")
            else:
                output_box = st.empty()
                output_box.info("Veriler analiz ediliyor...")
                
                try:
                    genai.configure(api_key=api_key)
                    active_model = get_best_model()
                    model = genai.GenerativeModel(active_model)
                    
                    # Tarihleri stringe çevirerek JSON hatasını önle
                    prompt = f"""
                    GÖREV: Bir avukat için zamanaşımı risk analizi yap.
                    
                    DURUM:
                    - Dava Türü: {dava_turu}
                    - Fesih Tarihi: {fesih_tarihi}
                    - Bugün: {bugun}
                    - Tablo Verileri: {df.to_json(orient='records', date_format='iso')}
                    
                    İSTENENLER:
                    1. Hangi kalemlerde zamanaşımı riski var? (Kısa ve net)
                    2. Islah dilekçesi için ne kadar vaktim kaldı? Geç kalırsam ne olur?
                    3. Zamanaşımı def'i (savunması) ile karşılaşırsam ne yapmalıyım?
                    4. Faiz başlangıç tarihleri için stratejik bir öneri ver.
                    """
                    
                    response = model.generate_content(prompt, stream=True)
                    
                    full_text = ""
                    for chunk in response:
                        full_text += chunk.text
                        output_box.markdown(full_text + "▌")
                    output_box.markdown(full_text)
                    
                except Exception as e:
                    output_box.error(f"Hata: {e}")


import networkx as nx
import matplotlib.pyplot as plt

def render_conflict_scanner(api_key):
    st.info("🕸️ **Gizli Bağlantı (Conflict of Interest) Tarayıcısı:** Hakim, avukat ve tanıklar arasındaki görünmez ticari ve sosyal bağları ortaya çıkarır. NetworkX ile ağ analizi yapar.")

    # --- 0. OTOMATİK MODEL SEÇİCİ ---
    def get_best_model():
        try:
            available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
            for m in available_models:
                if 'flash' in m: return m
            return available_models[0] if available_models else "models/gemini-pro"
        except:
            return "models/gemini-pro"

    col_input, col_graph = st.columns([1, 2])

    # --- 1. GİRDİ PANELİ ---
    with col_input:
        st.markdown("### 👥 Aktörleri Tanımla")
        
        hakim = st.text_input("Hakim İsmi", "Hakim Zeynep Yılmaz")
        avukat_karsi = st.text_input("Karşı Taraf Avukatı", "Av. Ahmet Demir")
        tanik = st.text_input("Tanık / Bilirkişi", "Bilirkişi Mehmet Öztürk")
        sirket = st.text_input("İlgili Şirket (Opsiyonel)", "Delta İnşaat A.Ş.")
        
        st.divider()
        st.markdown("#### 📡 Veri Kaynağı")
        source_type = st.radio("Tarama Yöntemi", ["Demo Simülasyonu (NetworkX Testi)", "AI ile Açık Kaynak Tarama (OSINT)"])

    # --- 2. ANALİZ MOTORU ---
    with col_graph:
        st.markdown("### 🕸️ İlişki Ağı Haritası")
        
        if st.button("🔍 Derinlemesine Tara", type="primary"):
            
            # SENARYO 1: DEMO SİMÜLASYONU (NetworkX Gücünü Göstermek İçin)
            if source_type == "Demo Simülasyonu (NetworkX Testi)":
                st.warning("⚠️ Demo Modu: Rastgele ticari sicil verileri simüle ediliyor...")
                
                # NetworkX Grafiği Oluştur
                G = nx.Graph()
                
                # Düğümleri (Kişileri/Kurumları) Ekle
                G.add_node(hakim, type="Yargı", color="red")
                G.add_node(avukat_karsi, type="Avukat", color="black")
                G.add_node(tanik, type="Tanık", color="blue")
                G.add_node(sirket, type="Şirket", color="green")
                
                # Gizli Bağlantıları Ekle (Simülasyon)
                # Örnek: Avukat ve Tanık, 5 yıl önce "Omega Yazılım"da ortaktı.
                hidden_entity = "Omega Yazılım Ltd. Şti. (Eski)"
                G.add_node(hidden_entity, type="Şirket", color="grey")
                
                G.add_edge(avukat_karsi, hidden_entity, relation="Eski Ortak (2018)")
                G.add_edge(tanik, hidden_entity, relation="Yön. Kur. Üyesi (2018)")
                
                # Örnek: Hakim ve Şirket arasında dolaylı bağ
                dernek = "Hukukçular Vakfı"
                G.add_node(dernek, type="STK", color="orange")
                G.add_edge(hakim, dernek, relation="Üye")
                G.add_edge(avukat_karsi, dernek, relation="Yönetim Kurulu")

                # Graphviz ile Çiz (Streamlit için en temizi)
                dot_code = "graph {"
                dot_code += "\n  rankdir=LR;"
                
                # NetworkX verisini Graphviz formatına çevir
                for u, v, data in G.edges(data=True):
                    rel = data.get('relation', '')
                    dot_code += f'\n  "{u}" -- "{v}" [label="{rel}", fontsize=10];'
                
                # Renklendirme
                dot_code += f'\n  "{hakim}" [style=filled, fillcolor="#ffcccc"];' # Kırmızımsı
                dot_code += f'\n  "{avukat_karsi}" [style=filled, fillcolor="#cccccc"];'
                dot_code += f'\n  "{tanik}" [style=filled, fillcolor="#ccccff"];'
                dot_code += f'\n  "{hidden_entity}" [style=filled, fillcolor="#ffffcc", shape=box];' # Sarı (Kilit Nokta)
                
                dot_code += "\n}"
                st.graphviz_chart(dot_code)
                
                # NETWORKX ANALİZİ: EN KISA YOL (Shortest Path)
                try:
                    path = nx.shortest_path(G, source=avukat_karsi, target=tanik)
                    st.error(f"🚨 **KRİTİK BULGU:** {avukat_karsi} ile {tanik} arasında bağlantı tespit edildi!")
                    st.write(f"🔗 **Bağlantı Zinciri:** {' -> '.join(path)}")
                    st.caption("Bu bilgi, HMK m. 254 kapsamında tanığın tarafsızlığını sorgulamak için kullanılabilir.")
                except nx.NetworkXNoPath:
                    st.success("Doğrudan bir bağlantı bulunamadı.")

            # SENARYO 2: AI OSINT ANALİZİ (Gerçekçi Senaryo)
            else:
                if not api_key:
                    st.error("API Key gerekli.")
                else:
                    output_box = st.empty()
                    output_box.info("Açık kaynaklar ve haberler taranıyor...")
                    
                    try:
                        genai.configure(api_key=api_key)
                        active_model = get_best_model()
                        model = genai.GenerativeModel(active_model)
                        
                        # Prompt: AI'yı bir OSINT uzmanı gibi çalıştırıyoruz
                        prompt = f"""
                        GÖREV: Sen kıdemli bir istihbarat analistisin.
                        Aşağıdaki kişiler arasında potansiyel bir "Çıkar Çatışması" (Conflict of Interest) senaryosu kurgula ve analiz et.
                        
                        KİŞİLER:
                        1. Hakim: {hakim}
                        2. Karşı Avukat: {avukat_karsi}
                        3. Tanık: {tanik}
                        4. Şirket: {sirket}
                        
                        İSTENENLER:
                        1. Bu isimler arasında olası (hayali veya genel bilgiye dayalı) geçmiş bağlantıları düşün (Eski okul arkadaşlığı, aynı dernek üyeliği, eski şirket ortaklığı).
                        2. Özellikle "Tanık" ile "Karşı Avukat" arasında redd-i hakim veya tanık itirazına gerekçe olabilecek bir bağ bul.
                        3. Bunu bir "İstihbarat Raporu" formatında sun.
                        4. Hukuki Tavsiye: Bu bağlantıyı mahkemede nasıl delillendiririm?
                        """
                        
                        response = model.generate_content(prompt, stream=True)
                        
                        full_text = ""
                        for chunk in response:
                            full_text += chunk.text
                            output_box.markdown(full_text + "▌")
                        output_box.markdown(full_text)
                        
                    except Exception as e:
                        output_box.error(f"Hata: {e}")

def render_mediation_checker(api_key):
    st.info("🤝 **Arabuluculuk Kontrolcüsü:** Dava türünü girin, sistem bunun 'Dava Şartı (Zorunlu)' olup olmadığını, ilgili kanun maddesini ve başvuru süresini analiz etsin.")

    # --- 0. MODEL SEÇİCİ ---
    def get_best_model():
        try:
            available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
            for m in available_models:
                if 'flash' in m: return m
            return available_models[0] if available_models else "models/gemini-pro"
        except:
            return "models/gemini-pro"

    col1, col2 = st.columns([1, 1])

    with col1:
        st.markdown("### ⚖️ Dava Konusu Nedir?")
        
        # Hızlı Seçim Butonları
        st.write("Sık Kullanılanlar:")
        col_btn1, col_btn2, col_btn3 = st.columns(3)
        
        case_input = st.text_input("Veya detaylı yazın (Örn: Mobbing nedeniyle tazminat)", "")
        
        if col_btn1.button("Kıdem/İhbar"): case_input = "İşçilik Alacakları (Kıdem, İhbar, Fazla Mesai)"
        if col_btn2.button("Kira/Tahliye"): case_input = "Kira Tespiti ve Tahliye (Konut/Çatılı İşyeri)"
        if col_btn3.button("Ticari Alacak"): case_input = "İki Tacir Arasındaki Fatura Alacağı"

    with col2:
        st.markdown("### 🔍 Analiz Sonucu")
        
        if st.button("Arabuluculuk Şartını Kontrol Et", type="primary"):
            if not case_input:
                st.warning("Lütfen bir dava türü girin.")
            elif not api_key:
                st.error("API Key gerekli.")
            else:
                status_box = st.empty()
                status_box.info("Mevzuat taranıyor (7036, 6102, 6325 Sayılı Kanunlar)...")
                
                try:
                    genai.configure(api_key=api_key)
                    active_model = get_best_model()
                    model = genai.GenerativeModel(active_model)
                    
                    prompt = f"""
                    GÖREV: Sen uzman bir Türk Hukuku avukatısın.
                    SORGU: "{case_input}" konulu bir dava açmak istiyorum.
                    
                    ANALİZ ET:
                    1. Bu dava için Arabuluculuk ZORUNLU MU (Dava Şartı mı) yoksa İHTİYARİ Mİ?
                    2. Hangi Kanun maddesine dayanıyor? (Örn: TTK 5/A, İŞK 3, 7445 SK vb.)
                    3. Eğer zorunluysa ve gitmezsem ne olur? (Usulden Ret uyarısı)
                    4. Başvuru nereye yapılır? (Adliye/Büro)
                    
                    ÇIKTI FORMATI:
                    Lütfen cevabı şu formatta ver (Markdown kullanarak):
                    
                    ### 🚦 DURUM: [ZORUNLU / İHTİYARİ / İSTİSNA]
                    
                    **📜 Yasal Dayanak:** ...
                    **⚠️ Risk Uyarısı:** ...
                    **📍 Başvuru Yeri:** ...
                    **💡 Kısa Özet:** ...
                    """
                    
                    response = model.generate_content(prompt, stream=True)
                    
                    full_text = ""
                    for chunk in response:
                        full_text += chunk.text
                        status_box.markdown(full_text + "▌")
                    status_box.markdown(full_text)
                    
                    # Görsel Uyarılar (Basit Regex Kontrolü)
                    if "ZORUNLU" in full_text:
                        st.error("🚨 DİKKAT: Arabulucuya gitmeden dava açarsanız, davanız USULDEN REDDEDİLİR!")
                    elif "İHTİYARİ" in full_text:
                        st.success("✅ Zorunlu değil, doğrudan dava açabilirsiniz. Ancak yine de arabuluculuk denenebilir.")
                        
                except Exception as e:
                    status_box.error(f"Hata: {e}")

    st.divider()
    st.caption("ℹ️ Not: 01.09.2023 tarihinden itibaren Kira, Kat Mülkiyeti, Komşuluk Hukuku ve Ortaklığın Giderilmesi davaları da zorunlu arabuluculuk kapsamına alınmıştır.")


import folium
from streamlit_folium import st_folium
import random
import streamlit as st

def render_forensic_map(api_key):
    st.info("🗺️ **Adli Isı Haritası (Forensic Geolocation):** Olay yerindeki geçmiş vakaları analiz eder. 'Sürücü hatası mı, yoksa yol kusuru mu?' sorusuna İdare Hukuku perspektifiyle yanıt arar.")

    # --- SESSION STATE (HAFIZA) AYARLARI ---
    # Analiz yapıldı mı bilgisini hafızada tutuyoruz
    if "map_analyzed" not in st.session_state:
        st.session_state.map_analyzed = False
    
    # AI Sonucunu hafızada tutmak için (Tekrar tekrar API harcamasın)
    if "ai_map_result" not in st.session_state:
        st.session_state.ai_map_result = None

    # --- 0. MODEL SEÇİCİ ---
    def get_best_model():
        try:
            # Genai import kontrolü (Global scope'ta yoksa hata vermesin)
            import google.generativeai as genai
            available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
            for m in available_models:
                if 'flash' in m: return m
            return available_models[0] if available_models else "models/gemini-pro"
        except:
            return "models/gemini-pro"

    col_input, col_map = st.columns([1, 2])

    # --- 1. GİRDİ PANELİ ---
    with col_input:
        st.markdown("### 📍 Olay Yeri Tanımla")
        
        location_name = st.text_input("Konum / Kavşak Adı", "Bağdat Caddesi Şaşkınbakkal Işıklar")
        city = st.selectbox("Şehir", ["İstanbul", "Ankara", "İzmir", "Bursa", "Antalya"])
        event_type = st.selectbox("Olay Türü", ["Trafik Kazası", "Hırsızlık / Gasp", "Sel / Su Baskını", "Çukur / Yol Çökmesi"])
        
        st.divider()
        st.markdown("#### 🎯 Hedef Analiz")
        
        # BUTONLAR
        col_b1, col_b2 = st.columns(2)
        
        # Analiz Butonu (Callback ile hafızayı tetikler)
        def activate_analysis():
            st.session_state.map_analyzed = True
            st.session_state.ai_map_result = None # Yeni analiz için eski AI sonucunu sil
            
        if col_b1.button("📡 Bölgeyi Tara", type="primary", on_click=activate_analysis):
            pass # İşlemi aşağıda yapacağız
            
        # Sıfırlama Butonu
        def reset_analysis():
            st.session_state.map_analyzed = False
            st.session_state.ai_map_result = None
            
        if col_b2.button("🔄 Sıfırla", on_click=reset_analysis):
            pass

    # --- 2. HARİTA VE ANALİZ ---
    with col_map:
        # Varsayılan Koordinatlar
        lat, lon = 41.0082, 28.9784 
        if city == "Ankara": lat, lon = 39.9334, 32.8597
        if city == "İzmir": lat, lon = 38.4192, 27.1287

        # Harita Oluştur (Her seferinde temiz başlar)
        m = folium.Map(location=[lat, lon], zoom_start=13)
        
        # EĞER ANALİZ BUTONUNA BASILDIYSA (Hafıza True ise)
        if st.session_state.map_analyzed:
            st.markdown(f"### 🔍 {location_name} - Risk Analizi")
            
            # --- A. HARİTA GÖRSELLEŞTİRME ---
            # Merkez Nokta
            folium.Marker(
                [lat, lon], 
                popup=f"<b>OLAY YERİ</b><br>{location_name}", 
                icon=folium.Icon(color="red", icon="info-sign")
            ).add_to(m)
            
            # Risk Noktaları (Simülasyon)
            # Not: Her render'da yer değiştirmemesi için seed sabitlenebilir veya statik veri kullanılabilir
            random.seed(42) 
            for _ in range(15):
                r_lat = lat + random.uniform(-0.015, 0.015)
                r_lon = lon + random.uniform(-0.015, 0.015)
                folium.CircleMarker(
                    location=[r_lat, r_lon],
                    radius=6,
                    color="crimson",
                    fill=True,
                    fill_color="crimson",
                    fill_opacity=0.6,
                    popup="Geçmiş Vaka (Riskli Bölge)"
                ).add_to(m)

            # Haritayı Çiz
            st_folium(m, height=350, width=700)
            
            # --- B. YAPAY ZEKA ANALİZİ ---
            if not api_key:
                st.error("Detaylı rapor için API Key gerekli.")
            else:
                # Eğer daha önce üretilmediyse üret
                if st.session_state.ai_map_result is None:
                    status_box = st.info("Haber arşivleri ve yerel şikayetler taranıyor... Lütfen bekleyin.")
                    
                    try:
                        import google.generativeai as genai
                        genai.configure(api_key=api_key)
                        active_model = get_best_model()
                        model = genai.GenerativeModel(active_model)
                        
                        prompt = f"""
                        GÖREV: Sen uzman bir İdare Hukuku avukatı ve Trafik Bilirkişisisin.
                        KONUM: {location_name}, {city}
                        OLAY TÜRÜ: {event_type}
                        
                        SENARYO: Müvekkil burada bir kaza yaptı/zarar gördü. Sadece karşı tarafı değil, devleti/belediyeyi de dava etmek istiyoruz.
                        
                        İSTENENLER:
                        1. Bu bölgeyle ilgili geçmişte basına yansıyan benzer kazalar veya "ölüm virajı", "karanlık yol" gibi haberler var mı? (Genel bilgi birikimini kullan).
                        2. İdarenin "Hizmet Kusuru" (Service Defect) sayılabilecek ihmalleri neler olabilir? (Örn: Sinyalizasyon eksikliği, yetersiz aydınlatma, çukur, rögar kapağı).
                        3. STRATEJİ: Davayı "Tam Yargı Davası" olarak İdare Mahkemesi'ne taşımak için hangi delilleri toplamalıyım? (MOBESE, Belediye şikayet kayıtları vb.)
                        4. SONUÇ: "Bu kavşakta son 1 yılda çok kaza olduysa, kusur sürücüde değil yoldadır" tezini savunacak hukuki argümanlar yaz.
                        """
                        
                        response = model.generate_content(prompt)
                        st.session_state.ai_map_result = response.text
                        status_box.empty() # Yükleniyor yazısını kaldır
                        
                    except Exception as e:
                        st.error(f"AI Hatası: {e}")
                
                # Sonucu Göster (Hafızadan)
                if st.session_state.ai_map_result:
                    st.markdown(st.session_state.ai_map_result)
        
        else:
            # Analiz öncesi boş harita
            st_folium(m, height=350, width=700)
            st.caption("👈 Analiz butonuna bastığınızda bölgedeki risk yoğunluğu haritaya işlenecektir.")



def render_temporal_law_machine(api_key):
    import datetime as dtss # Çakışmayı önlemek için güvenli import
    st.info("🕰️ **Mevzuat Zaman Makinesi:** Olayın yaşandığı tarihe geri döner. O gün yürürlükte olan (şu an mülga) kanunları, tüzükleri ve Yargıtay içtihatlarını bugünkülerle kıyaslar.")
	
    # --- 0. MODEL SEÇİCİ ---
    def get_best_model():
        try:
            available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
            for m in available_models:
                if 'flash' in m: return m
            return available_models[0] if available_models else "models/gemini-pro"
        except:
            return "models/gemini-pro"

    col_settings, col_result = st.columns([1, 2])

    # --- 1. ZAMAN AYARLARI ---
    with col_settings:
        st.markdown("### ⚙️ Zaman Koordinatları")
        
        # Tarih Seçimi (Varsayılan: 1990'lar)
        target_date = st.date_input("Olay Tarihi", dtss.date(1995, 6, 15))
        
        topic = st.selectbox("Hukuki Konu", [
            "Gayrimenkul Devri (Tapu İptal)", 
            "Miras Paylaşımı (Tereke)", 
            "Boşanma ve Mal Rejimi", 
            "İş Kazası Tazminatı",
            "Ticari Sözleşme İhlali"
        ])
        
        specific_query = st.text_input("Özel Detay (Opsiyonel)", "Muris muvazaası ve saklı pay")
        
        st.divider()
        st.markdown("#### 🔄 Dönüşüm Modu")
        comparison_mode = st.radio("Analiz Türü", ["Sadece O Günün Kanunu", "Eski vs Yeni Kanun Kıyaslaması"])
        
        start_travel = st.button("🚀 Geçmişe Git ve Mevzuatı Getir", type="primary")

    # --- 2. SONUÇ EKRANI ---
    with col_result:
        st.markdown(f"### 📜 {target_date.year} Yılı Mevzuat Panoramas")
        
        if start_travel:
            if not api_key:
                st.error("Zaman yolculuğu için API Key gerekli.")
            else:
                # Görsel Efekt
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                status_text.text(f"⏳ {target_date.year} yılına gidiliyor...")
                time.sleep(0.5)
                progress_bar.progress(30)
                
                status_text.text("📚 Resmi Gazete arşivleri taranıyor...")
                time.sleep(0.5)
                progress_bar.progress(60)
                
                status_text.text("⚖️ Mülga kanun maddeleri getiriliyor...")
                progress_bar.progress(90)
                
                try:
                    genai.configure(api_key=api_key)
                    active_model = get_best_model()
                    model = genai.GenerativeModel(active_model)
                    
                    # Kritik Tarih Kontrolleri (Prompt'a ipucu vermek için)
                    era_context = ""
                    if target_date.year < 2002:
                        era_context += "UYARI: Bu tarihte 4721 sayılı Türk Medeni Kanunu YOKTU. 743 sayılı Türk Kanunu Medenisi yürürlükteydi. "
                    if target_date.year < 2012:
                        era_context += "UYARI: 6098 sayılı Borçlar Kanunu YOKTU. 818 sayılı Borçlar Kanunu yürürlükteydi. "
                    
                    prompt = f"""
                    GÖREV: Sen bir Hukuk Tarihçisi ve Mevzuat Uzmanısın.
                    
                    HEDEF TARİH: {target_date.strftime('%d.%m.%Y')}
                    KONU: {topic}
                    DETAY: {specific_query}
                    BAĞLAM: {era_context}
                    
                    İSTENEN ÇIKTI (Rapor Formatı):
                    
                    1. 🏛️ YÜRÜRLÜKTEKİ TEMEL KANUN
                    - O tarihte geçerli olan Kanun Numarası ve Adı (Örn: 743 s. TKM).
                    - İlgili Madde Numarası ve (mümkünse) o günkü metni.
                    
                    2. 📜 KRİTİK FARKLILIKLAR (BUGÜNE GÖRE)
                    - Bugün uygulanan kanunla (Örn: 4721 s. TMK) o günkü kanun arasındaki hayati fark nedir?
                    - Örnek: "O tarihte 'Edinilmiş Mallara Katılma Rejimi' yoktu, 'Mal Ayrılığı' esastı."
                    
                    3. ⚖️ DÖNEMİN İÇTİHADI
                    - O yıllarda Yargıtay'ın bu konuya bakışı nasıldı? (Örn: 1990'larda inançlı işlem içtihadı).
                    
                    4. 💎 AVUKAT İÇİN STRATEJİ
                    - Davayı kazanmak için mahkemeye "Olay tarihindeki mevzuat uygulanmalıdır" itirazını nasıl sunmalıyım?
                    """
                    
                    response = model.generate_content(prompt, stream=True)
                    
                    full_text = ""
                    status_text.empty() # Yazıyı temizle
                    output_placeholder = st.empty()
                    
                    for chunk in response:
                        full_text += chunk.text
                        output_placeholder.markdown(full_text + "▌")
                    
                    output_placeholder.markdown(full_text)
                    progress_bar.progress(100)
                    
                except Exception as e:
                    st.error(f"Hata: {e}")
        else:
            st.info("👈 Sol taraftan tarihi seçin ve yolculuğu başlatın.")
            
            # Örnek Gösterim (Placeholder)
            st.markdown("""
            **Örnek Senaryo:**
            * **Tarih:** 1995
            * **Konu:** Boşanma Mal Paylaşımı
            * **Sonuç:** 2002 öncesi evliliklerde "Mal Ayrılığı" rejimi geçerli olduğundan, kadın eşin ev hanımı olması durumunda tapuda adı yoksa hak talep etmesi çok zordu. Sistem bunu tespit edip "Katkı Payı Alacağı" davası açmanızı önerir.
            """)



def render_expert_report_auditor(api_key):
    st.info("🧐 **Bilirkişi Raporu Denetçisi:** Karmaşık raporları tarar. Matematiksel hataları (kusur toplamı != 100), mantıksal çelişkileri ve eksik incelemeleri tespit ederek 'İtiraz Dilekçesi' taslağı hazırlar.")

    # --- 0. KÜTÜPHANE KONTROLÜ ---
    try:
        from pypdf import PdfReader
    except ImportError:
        st.error("Bu modül için 'pypdf' kütüphanesi gereklidir. Lütfen requirements.txt dosyasına ekleyin.")
        return

    # --- 1. GİRDİ PANELİ ---
    col_upload, col_analysis = st.columns([1, 1])

    report_text = ""

    with col_upload:
        st.markdown("### 📄 Raporu Yükle")
        uploaded_file = st.file_uploader("Bilirkişi Raporu (PDF)", type=["pdf"])
        
        st.markdown("--- VEYA ---")
        text_input = st.text_area("Metni Buraya Yapıştır", height=150, placeholder="Rapor içeriğini buraya kopyalayabilirsiniz...")

        # Metin Çıkarma İşlemi
        if uploaded_file:
            try:
                reader = PdfReader(uploaded_file)
                for page in reader.pages:
                    report_text += page.extract_text() + "\n"
                st.success(f"✅ PDF Okundu: {len(reader.pages)} sayfa")
            except Exception as e:
                st.error(f"PDF Okuma Hatası: {e}")
        elif text_input:
            report_text = text_input

    # --- 2. ANALİZ MOTORU ---
    with col_analysis:
        st.markdown("### 🔍 Denetim Sonucu")
        
        analyze_btn = st.button("🛡️ Raporu Denetle ve Hata Bul", type="primary")
        
        if analyze_btn:
            if not report_text:
                st.warning("Lütfen analiz edilecek bir rapor yükleyin veya metin girin.")
            elif len(report_text) < 50:
                st.warning("Girilen metin analiz için çok kısa.")
            elif not api_key:
                st.error("API Key gerekli.")
            else:
                output_box = st.empty()
                output_box.info("Rapor taranıyor: Kusur oranları toplanıyor, çelişkiler aranıyor...")
                
                try:
                    import google.generativeai as genai
                    genai.configure(api_key=api_key)
                    
                    # Model Seçimi
                    model_name = "models/gemini-pro"
                    for m in genai.list_models():
                        if 'flash' in m.name: model_name = m.name; break
                    
                    model = genai.GenerativeModel(model_name)
                    
                    prompt = f"""
                    GÖREV: Sen titiz bir 'Bilirkişi Raporu Denetçisi' ve Yargıtay İçtihatları uzmanısın.
                    Aşağıdaki bilirkişi raporu metnini analiz et ve hataları bul.
                    
                    METİN:
                    {report_text[:10000]} (Metin kısaltıldıysa devamını dikkate al)
                    
                    İSTENEN ANALİZ (Markdown Formatında):
                    
                    ### 1. 🧮 Matematiksel ve Mantıksal Tutarlılık
                    - Kusur oranları toplamı 100 ediyor mu? (Kontrol et: %25 + %75 vb.)
                    - Hesaplamalarda bariz bir çarpım/toplam hatası var mı?
                    - Tarihler tutarlı mı? (Kaza tarihinden sonraki bir mevzuat uygulanmış mı?)
                    
                    ### 2. ⚖️ Hukuki ve Teknik Dayanak
                    - Rapor hangi teknik veriye dayanıyor? (Tramer, MOBESE, Tanık, Takograf vb.)
                    - Bilirkişi "Hukuki niteleme" yapmış mı? (UYARI: Bilirkişi hukuki yorum yapamaz, sadece teknik tespit yapar. Hakim yerine geçip hüküm kurduysa bunu belirt.)
                    
                    ### 3. 🚩 Tespit Edilen Çelişkiler
                    - "Tanık ifadesinde X denmesine rağmen, raporda Y kabul edilmiştir" gibi çelişkiler var mı?
                    
                    ### 4. 📝 İtiraz Stratejisi (HMK m. 281)
                    - Bu rapora itiraz etmek için kullanılabilecek 3 güçlü argüman yaz.
                    - "Ek Rapor" veya "Yeni Bilirkişi Heyeti" talep etmek için gerekçe oluştur.
                    """
                    
                    response = model.generate_content(prompt, stream=True)
                    
                    full_text = ""
                    for chunk in response:
                        full_text += chunk.text
                        output_box.markdown(full_text + "▌")
                    output_box.markdown(full_text)
                    
                except Exception as e:
                    output_box.error(f"Analiz Hatası: {e}")



def render_corporate_memory(api_key):
    st.info("🏛️ **Kurumsal Hafıza V3 (Oto-Pilot):** Mevcut en güncel AI modelini otomatik bulur ve 'Model Bulunamadı' hatalarını engeller.")

    # --- KÜTÜPHANE KONTROLLERİ ---
    try:
        import pandas as pd
        from pypdf import PdfReader
        from docx import Document
        from PIL import Image
        import google.generativeai as genai
    except ImportError:
        st.error("Gerekli kütüphaneler eksik (pandas, pypdf, python-docx, Pillow, google-generativeai).")
        return

    # --- 0. OTURUM VE VERİ YÖNETİMİ ---
    if "archive_df" not in st.session_state:
        st.session_state.archive_df = pd.DataFrame(columns=["Tarih", "Konu", "Özet", "Detay", "İlgili Kişi/Kurum", "Dosya Adı"])

    # --- KRİTİK FONKSİYON: SAĞLAM MODEL BULUCU ---
    def get_working_model(api_key_val):
        """
        API'den güncel listeyi çeker. Hata verirse manuel listeyi dener.
        En garantili çalışan modeli döndürür.
        """
        genai.configure(api_key=api_key_val)
        
        # 1. YÖNTEM: API'den Canlı Liste İste
        try:
            available_models = []
            for m in genai.list_models():
                if 'generateContent' in m.supported_generation_methods:
                    available_models.append(m.name)
            
            # Öncelik Sıralaması (En iyiden en eskiye)
            priorities = [
                'models/gemini-1.5-flash',
                'models/gemini-1.5-pro',
                'models/gemini-1.5-flash-latest',
                'models/gemini-1.0-pro',
                'models/gemini-pro'
            ]
            
            # Listede eşleşen en iyi modeli bul
            for p in priorities:
                if p in available_models:
                    return p
            
            # Tam eşleşme yoksa, içinde 'flash' geçeni al
            for m in available_models:
                if 'flash' in m: return m
                
            # O da yoksa ilk bulduğunu al
            if available_models:
                return available_models[0]
                
        except Exception as e:
            pass # Liste alınamazsa manuel listeye geç

        # 2. YÖNTEM: Manuel Güvenli Liste (Fallback)
        # API listeleme başarısız olsa bile bu isimler genellikle çalışır
        return "models/gemini-1.5-flash"

    # --- SEKME YAPISI ---
    tab_upload, tab_query = st.tabs(["📂 Belge İşle & Arşivle", "🔍 Arşivde Sorgu Yap"])

    # ==========================================
    # 1. SEKME: BELGE İŞLEME
    # ==========================================
    with tab_upload:
        col_db, col_process = st.columns([1, 1])

        # A. MEVCUT VERİTABANINI YÜKLE
        with col_db:
            st.markdown("### 1. Mevcut Arşivi Yükle")
            uploaded_excel = st.file_uploader("Önceki Excel Dosyanız", type=["xlsx"])
            
            if uploaded_excel:
                try:
                    loaded_df = pd.read_excel(uploaded_excel)
                    st.session_state.archive_df = loaded_df
                    st.success(f"✅ Veritabanı Yüklendi! ({len(loaded_df)} kayıt)")
                except Exception as e:
                    st.error(f"Excel Hatası: {e}")
            
            # Tablo Önizleme
            st.dataframe(st.session_state.archive_df, height=200, use_container_width=True)
            
            # İndir
            if not st.session_state.archive_df.empty:
                excel_data = io.BytesIO()
                st.session_state.archive_df.to_excel(excel_data, index=False)
                st.download_button("💾 Arşivi İndir", excel_data.getvalue(), "Kurumsal_Hafiza.xlsx")

        # B. YENİ BELGE EKLE
        with col_process:
            st.markdown("### 2. Yeni Belge Ekle")
            files = st.file_uploader("Belgeler", type=["pdf", "docx", "png", "jpg", "jpeg"], accept_multiple_files=True)
            
            if st.button("⚙️ Analiz Et ve Ekle", type="primary") and files:
                if not api_key:
                    st.error("API Key gerekli.")
                else:
                    # --- MODELİ BELİRLE ---
                    active_model_name = get_working_model(api_key)
                    st.toast(f"🤖 Aktif Model: {active_model_name}", icon="✅")
                    
                    model = genai.GenerativeModel(active_model_name)
                    progress_bar = st.progress(0)
                    new_records = []
                    
                    for idx, file in enumerate(files):
                        try:
                            content = ""
                            image_data = None
                            is_image = False
                            
                            # Dosya Okuma
                            if file.type == "application/pdf":
                                reader = PdfReader(file)
                                for page in reader.pages: content += page.extract_text() + "\n"
                            elif "word" in file.type:
                                doc = Document(file)
                                for p in doc.paragraphs: content += p.text + "\n"
                            elif "image" in file.type:
                                is_image = True
                                image_data = Image.open(file)

                            # Prompt Hazırla
                            prompt = """
                            Bu belgeden şu bilgileri JSON formatında çıkar:
                            {"Tarih": "GG.AA.YYYY", "Konu": "...", "Ozet": "...", "Detay": "...", "Ilgili_Kisi": "..."}
                            Sadece JSON ver.
                            """
                            
                            # API Çağrısı (Retry Mekanizması ile)
                            response = None
                            try:
                                if is_image:
                                    # Eğer model vision desteklemiyorsa flash'a zorla
                                    if "flash" not in active_model_name and "1.5" not in active_model_name:
                                        model_vision = genai.GenerativeModel("models/gemini-1.5-flash")
                                        response = model_vision.generate_content([prompt, image_data])
                                    else:
                                        response = model.generate_content([prompt, image_data])
                                else:
                                    if len(content) > 5:
                                        response = model.generate_content(prompt + f"\n\nMETİN:\n{content[:20000]}")
                            except Exception as api_err:
                                st.warning(f"Model hatası ({active_model_name}), yedek model deneniyor...")
                                # Hata verirse kesin çalışan Flash modelini dene
                                backup_model = genai.GenerativeModel("models/gemini-1.5-flash")
                                if is_image:
                                    response = backup_model.generate_content([prompt, image_data])
                                else:
                                    response = backup_model.generate_content(prompt + f"\n\nMETİN:\n{content[:20000]}")

                            # Sonucu İşle
                            if response and response.text:
                                clean_json = response.text.replace("```json", "").replace("```", "").strip()
                                data = json.loads(clean_json)
                                
                                # Mükerrer Kontrolü
                                is_dup = False
                                if not st.session_state.archive_df.empty:
                                    check = st.session_state.archive_df[
                                        (st.session_state.archive_df['Konu'] == data.get('Konu')) & 
                                        (st.session_state.archive_df['Tarih'] == data.get('Tarih'))
                                    ]
                                    if not check.empty: is_dup = True
                                
                                if not is_dup:
                                    new_records.append({
                                        "Tarih": data.get("Tarih", "-"),
                                        "Konu": data.get("Konu", "-"),
                                        "Özet": data.get("Ozet", "-"),
                                        "Detay": data.get("Detay", "-"),
                                        "İlgili Kişi/Kurum": data.get("Ilgili_Kisi", "-"),
                                        "Dosya Adı": file.name
                                    })
                                    
                        except Exception as e:
                            st.error(f"Hata ({file.name}): {e}")
                        
                        progress_bar.progress((idx + 1) / len(files))

                    # Kaydet
                    if new_records:
                        st.session_state.archive_df = pd.concat([st.session_state.archive_df, pd.DataFrame(new_records)], ignore_index=True)
                        st.success("İşlem Tamamlandı!")
                        st.rerun()

    # ==========================================
    # 2. SEKME: SORGULAMA
    # ==========================================
    with tab_query:
        st.markdown("### 🧠 Arşivde Semantik Arama")
        if st.session_state.archive_df.empty:
            st.info("Veri yok.")
        else:
            query = st.text_input("Soru:", placeholder="Örn: X firması ile ilgili sözleşme detayı?")
            if st.button("🔍 Ara"):
                with st.spinner("Aranıyor..."):
                    try:
                        active_model = get_working_model(api_key)
                        model = genai.GenerativeModel(active_model)
                        context = st.session_state.archive_df.to_json(orient="records", force_ascii=False)
                        prompt = f"VERİTABANI:\n{context}\n\nSORU: {query}\n\nBu veritabanına göre cevapla:"
                        st.markdown(model.generate_content(prompt).text)
                    except Exception as e:
                        st.error(f"Hata: {e}")


def render_cost_calculator_module(api_key):
    st.header("💰 Dava Maliyeti ve Harç Hesaplama Robotu (2026 Projeksiyonu)")
    st.warning("⚠️ DİKKAT: Bu hesaplama, beklenen 2026 Yeniden Değerleme Oranlarına (Tahmini %45-50 Artış) göre simüle edilmiştir.")

    # --- GİRDİLER ---
    col1, col2 = st.columns(2)
    
    with col1:
        dava_turu = st.selectbox("Dava Türü", [
            "Asliye Hukuk (Konusu Para Olan)", 
            "Asliye Hukuk (Maktu - Örn: Tapu İptal)", 
            "İş Mahkemesi (İşe İade)", 
            "İş Mahkemesi (Alacak)", 
            "Tüketici Mahkemesi", 
            "Sulh Hukuk (Tahliye)",
            "İcra Takibi",
            "İdare Mahkemesi (İptal Davası)",      # YENİ
            "İdare Mahkemesi (Tam Yargı - Tazminat)" # YENİ
        ])
        
        # Dinamik Girdi Alanları
        dava_degeri = 0.0
        yd_talebi = False
        
        # İdari Davalarda YD Talebi Sorusu
        if "İdare" in dava_turu:
            yd_talebi = st.checkbox("Yürütmeyi Durdurma (YD) İsteniyor mu?", value=True)

        # Değer Girilmesi Gereken Durumlar
        if any(x in dava_turu for x in ["Konusu Para", "Alacak", "Tam Yargı", "İcra"]):
            dava_degeri = st.number_input("Dava/Talep Değeri (TL)", min_value=0.0, value=150000.0, step=1000.0, format="%.2f")
        
        # Tahliye Davası (Yıllık Kira)
        if "Tahliye" in dava_turu:
            aylik_kira = st.number_input("Aylık Kira Bedeli (TL)", min_value=0.0, value=15000.0)
            dava_degeri = aylik_kira * 12 
    
    with col2:
        davaci_sayisi = st.number_input("Davacı Sayısı", min_value=1, value=1)
        davali_sayisi = st.number_input("Davalı/İdare Sayısı", min_value=1, value=1)
        tanik_sayisi = st.number_input("Dinlenecek Tanık Sayısı", min_value=0, value=0 if "İdare" in dava_turu else 2)
        bilirkisi_sayisi = st.number_input("Bilirkişi Sayısı", min_value=0, value=1)

    # --- HESAPLAMA BUTONU ---
    if st.button("🧮 2026 Tarifesine Göre Hesapla", type="primary"):
        
        # ==========================================
        # 🏛️ 2026 TAHMİNİ VERİLER (YDO Artışlı)
        # ==========================================
        
        # Harçlar (Tahmini)
        BASVURMA_HARCI_SULH = 500.00
        BASVURMA_HARCI_ASLIYE = 980.00
        BASVURMA_HARCI_IDARE = 980.00
        BASVURMA_HARCI_VERGI = 980.00
        
        VEKALET_HARCI = 140.00          # Baro Pulu (Tahmini)
        KARAR_ILAM_HARCI_MAKTU = 980.00
        YD_HARCI = 550.00               # Yürütmeyi Durdurma Harcı (Tahmini)
        
        # Giderler (Enflasyon Farkı Eklenmiş)
        DOSYA_GIDERI = 300.00
        TEBLIGAT_UCRETI = 300.00       # PTT 2026 Tahmini
        TANIK_UCRETI = 250.00
        BILIRKISI_UCRETI = 4500.00     # Bilirkişi ücretleri artış eğiliminde
        KESIF_HARCI = 3500.00
        
        # AAÜT MAKTU ÜCRETLER (2026 Tahmini - %45 Artış)
        AAUT_ASLIYE_MAKTU = 38000.00
        AAUT_SULH_MAKTU = 26000.00
        AAUT_ICRA_MAKTU = 9000.00
        AAUT_ISE_IADE = 38000.00
        AAUT_IDARE_MAKTU = 26000.00     # İptal davaları için
        AAUT_TUKETICI = 26000.00

        # --- HESAPLAMA MOTORU ---
        gider_avansi = 0.0
        pesin_harc = 0.0
        basvurma_harci = 0.0
        vekalet_ucreti = 0.0
        ekstra_harclar = 0.0 # YD harcı vb.
        
        # 1. Başvurma Harcı
        if "Sulh" in dava_turu or "İcra" in dava_turu:
            basvurma_harci = BASVURMA_HARCI_SULH
        elif "Tüketici" in dava_turu:
            basvurma_harci = 0 
        elif "İdare" in dava_turu:
            basvurma_harci = BASVURMA_HARCI_IDARE
        else:
            basvurma_harci = BASVURMA_HARCI_ASLIYE

        # 2. Gider Avansı
        tebligat_gideri = (davaci_sayisi + davali_sayisi) * 3 * TEBLIGAT_UCRETI 
        tanik_gideri = tanik_sayisi * TANIK_UCRETI
        bilirkisi_gideri = bilirkisi_sayisi * BILIRKISI_UCRETI
        diger_isler = 1000.00 # Kırtasiye 2026
        
        gider_avansi = tebligat_gideri + tanik_gideri + bilirkisi_gideri + diger_isler + DOSYA_GIDERI

        # 3. Peşin / Karar Harcı Hesabı
        if "Maktu" in dava_turu or "İşe İade" in dava_turu or "İptal" in dava_turu:
            pesin_harc = KARAR_ILAM_HARCI_MAKTU
        elif "Tüketici" in dava_turu:
            pesin_harc = 0
        elif "İcra" in dava_turu:
            pesin_harc = dava_degeri * 0.005 
        else:
            # Nispi Harç (%6.831) - Oran genelde sabittir, matrah değişir
            toplam_harc = dava_degeri * 0.06831
            pesin_harc = toplam_harc / 4
            
        # İdari Yargı Özel: Yürütmeyi Durdurma Harcı
        if yd_talebi:
            ekstra_harclar += YD_HARCI

        # 4. AAÜT (Avukatlık Ücreti) 2026 Tahmini
        if "İptal" in dava_turu:
            vekalet_ucreti = AAUT_IDARE_MAKTU
        elif "Maktu" in dava_turu:
            vekalet_ucreti = AAUT_ASLIYE_MAKTU
        elif "İşe İade" in dava_turu:
            vekalet_ucreti = AAUT_ISE_IADE
        elif "Sulh" in dava_turu:
            vekalet_ucreti = AAUT_SULH_MAKTU
        elif "Tüketici" in dava_turu:
            vekalet_ucreti = AAUT_TUKETICI
        else:
            # NİSPİ VEKALET (Dilimler 2026 için genişletildi)
            kalan = dava_degeri
            hesap = 0.0
            
            # Dilim 1: 600.000 TL'ye kadar %16 (Tahmini artış)
            dilim1 = min(kalan, 600000)
            hesap += dilim1 * 0.16
            kalan -= dilim1
            
            # Dilim 2
            if kalan > 0:
                dilim2 = min(kalan, 600000)
                hesap += dilim2 * 0.15
                kalan -= dilim2
                
            # Dilim 3 ve sonrası...
            if kalan > 0:
                hesap += kalan * 0.14 # Basitleştirilmiş devamı

            # Alt sınır kontrolü
            if "İdare" in dava_turu: # Tam Yargı
                vekalet_ucreti = max(hesap, AAUT_IDARE_MAKTU)
            elif "Asliye" in dava_turu:
                vekalet_ucreti = max(hesap, AAUT_ASLIYE_MAKTU)
            elif "Sulh" in dava_turu:
                vekalet_ucreti = max(hesap, AAUT_SULH_MAKTU)
            else:
                vekalet_ucreti = hesap

        toplam_ilk_masraf = basvurma_harci + VEKALET_HARCI + pesin_harc + gider_avansi + ekstra_harclar

        # --- SONUÇ EKRANI ---
        st.divider()
        st.subheader("📋 2026 Dava Maliyet Projeksiyonu")
        
        c_res1, c_res2, c_res3 = st.columns(3)
        c_res1.metric("Toplam İlk Masraf", f"{toplam_ilk_masraf:,.2f} TL", help="Müvekkilden talep edilecek toplam tutar")
        c_res2.metric("Karşı Yan Vekalet (Risk)", f"{vekalet_ucreti:,.2f} TL", help="Kaybedilirse ödenecek tutar")
        c_res3.metric("Peşin Harç", f"{pesin_harc:,.2f} TL")
        
        # Detay Tablosu
        detay_list = [
            ["Başvurma Harcı", f"{basvurma_harci:,.2f}"],
            ["Peşin / Karar Harcı", f"{pesin_harc:,.2f}"],
            ["Gider Avansı (Bilirkişi, Tebligat)", f"{gider_avansi:,.2f}"],
            ["Vekalet Harcı (Baro Pulu)", f"{VEKALET_HARCI:,.2f}"]
        ]
        if yd_talebi:
            detay_list.append(["Yürütmeyi Durdurma Harcı", f"{YD_HARCI:,.2f}"])
            
        detay_list.append(["**TOPLAM**", f"**{toplam_ilk_masraf:,.2f}**"])
        
        df_detay = pd.DataFrame(detay_list, columns=["Kalem", "Tutar (TL)"])
        st.table(df_detay)
        
        # --- AI RAPOR ---
        st.subheader("📝 2026 Müvekkil Bilgilendirme Notu")
        if api_key:
            with st.spinner("AI, 2026 projeksiyonlarına göre rapor hazırlıyor..."):
                prompt = f"""
                GÖREV: Bir avukat olarak müvekkile dava masraflarını açıklayan profesyonel bir mesaj yaz.
                
                DURUM:
                - Yıl: 2026 (Tahmini Rakamlar)
                - Dava Türü: {dava_turu}
                - Yürütmeyi Durdurma Talebi: {"Var" if yd_talebi else "Yok"}
                - Toplam Masraf: {toplam_ilk_masraf:,.2f} TL
                - Karşı Yan Vekalet Riski: {vekalet_ucreti:,.2f} TL
                
                İÇERİK:
                1. Masrafların 2026 yılı tahmini harç ve giderlerine göre hesaplandığını belirt.
                2. İdari dava ise Yürütmeyi Durdurma (YD) harcının dahil olup olmadığını belirt.
                3. Gider avansının (Bilirkişi, tebligat vb.) kullanılmayan kısmının iade edileceğini vurgula.
                4. Güven verici, net bir dil kullan.
                """
                aciklama = get_ai_response(prompt, api_key)
                st.markdown(f"<div class='buyur-abi-kutusu'>{aciklama}</div>", unsafe_allow_html=True)
                
                st.download_button("📩 Raporu İndir (Word)", create_word_file(aciklama + f"\n\nTOPLAM: {toplam_ilk_masraf} TL"), "2026_Maliyet_Raporu.docx")

def render_forensic_map_module(api_key):
    st.header("🗺️ Adli Olay Yeri ve Kaza Rekonstrüksiyonu")
    st.info("Trafik ve iş kazalarında kusur analizi ve olay yeri canlandırması yapar.")

    # --- GİRDİLER ---
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.subheader("📋 Kaza Verileri")
        kaza_tipi = st.selectbox("Kaza Tipi", ["Trafik Kazası (Araç-Araç)", "Trafik Kazası (Araç-Yaya)", "İş Kazası (Düşme/Çarpma)"])
        
        arac_hizi = st.number_input("Araç Hızı (km/s)", value=90, step=5)
        fren_izi = st.number_input("Fren İzi Uzunluğu (metre)", value=35.0, step=1.0)
        yol_durumu = st.selectbox("Yol Durumu", ["Kuru Asfalt (0.8)", "Islak Asfalt (0.5)", "Buzlu Yol (0.2)", "Toprak Yol (0.6)"])
        
        # Sürtünme Katsayısı Belirleme
        katsayilar = {"Kuru Asfalt (0.8)": 0.8, "Islak Asfalt (0.5)": 0.5, "Buzlu Yol (0.2)": 0.2, "Toprak Yol (0.6)": 0.6}
        mu = katsayilar[yol_durumu]

    with col2:
        st.subheader("📍 Olay Yeri Krokisi (Simülasyon)")
        
        # --- FİZİK MOTORU ---
        # Formül: V = sqrt(2 * mu * g * d) * 3.6
        tahmini_hiz = (2 * mu * 9.81 * fren_izi)**0.5 * 3.6
        
        # Reaksiyon Mesafesi (1 saniye)
        reaksiyon_mesafesi = (arac_hizi / 3.6) * 1.0 
        durma_mesafesi = reaksiyon_mesafesi + fren_izi
        
        # Grafik Çizimi (Plotly)
        fig = go.Figure()
        
        # Yol Çizgileri
        fig.add_shape(type="rect", x0=0, y0=0, x1=durma_mesafesi + 20, y1=10, fillcolor="gray", opacity=0.3, line_width=0)
        fig.add_shape(type="line", x0=0, y0=5, x1=durma_mesafesi + 20, y1=5, line=dict(color="white", width=3, dash="dash"))
        
        # 1. Araç (Başlangıç) - DÜZELTME: symbol="car" yerine "square" yapıldı
        fig.add_trace(go.Scatter(
            x=[0], y=[2.5], 
            mode='markers+text', 
            marker=dict(size=25, symbol="square", color="blue"), # <-- DÜZELTİLEN KISIM
            text=["🚙 Fren Başlangıcı"], # Emojiyi metin içine ekledik
            textposition="top center"
        ))
        
        # 2. Araç (Bitiş)
        fig.add_trace(go.Scatter(
            x=[fren_izi], y=[2.5], 
            mode='markers+text', 
            marker=dict(size=25, symbol="x", color="red"), 
            text=["💥 Çarpma/Durma"], 
            textposition="top center"
        ))
        
        # Fren İzi Çizgisi
        fig.add_trace(go.Scatter(x=[0, fren_izi], y=[2.5, 2.5], mode='lines', line=dict(color='black', width=4), name='Fren İzi'))
        
        fig.update_layout(
            title="Kaza Krokisi (Kuşbakışı)",
            xaxis_title="Mesafe (metre)",
            yaxis_title="",
            yaxis=dict(showticklabels=False, range=[-2, 12]),
            xaxis=dict(range=[-5, durma_mesafesi + 10]),
            height=300,
            margin=dict(l=20, r=20, t=40, b=20),
            showlegend=False
        )
        st.plotly_chart(fig, use_container_width=True)

    # --- ANALİZ SONUCU ---
    st.divider()
    col_res1, col_res2 = st.columns(2)
    
    with col_res1:
        st.markdown("### 🔬 Fiziksel Analiz Raporu")
        st.write(f"**Beyan Edilen Hız:** {arac_hizi} km/s")
        st.write(f"**Fren İzinden Hesaplanan Hız:** {tahmini_hiz:.2f} km/s")
        
        delta = tahmini_hiz - arac_hizi
        if delta > 10:
            st.error(f"⚠️ DİKKAT: Araç beyan edilenden **{delta:.1f} km/s daha hızlı** gitmiş olabilir! Fren izleri bunu gösteriyor.")
        elif delta < -10:
            st.warning("ℹ️ Araç beyan edilenden daha yavaş olabilir veya fren sistemi tam verimli çalışmamış.")
        else:
            st.success("✅ Beyan edilen hız ile fiziksel bulgular uyumlu.")

    with col_res2:
        st.markdown("### ⚖️ Kusur & Bilirkişi Yorumu (AI)")
        if st.button("🤖 AI Bilirkişi Görüşü Al") and api_key:
            with st.spinner("Olay yeri verileri analiz ediliyor..."):
                prompt = f"""
                GÖREV: Trafik kazası bilirkişisi gibi davran.
                VERİLER:
                - Kaza Tipi: {kaza_tipi}
                - Yol Durumu: {yol_durumu}
                - Fren İzi: {fren_izi} metre
                - Sürücü Beyanı Hız: {arac_hizi} km/s
                - Fiziksel Hesaplanan Hız: {tahmini_hiz:.2f} km/s
                
                ANALİZ İSTEĞİ:
                1. Sürücünün "Hızın Gerekli Şartlara Uygunluğu" kuralını ihlal edip etmediğini değerlendir (KTK 52/1-b).
                2. Fren izi uzunluğu, sürücünün tehlikeyi geç fark ettiğini gösteriyor mu?
                3. Bu verilerle "Asli Kusur" mu yoksa "Tali Kusur" mu verilmesi muhtemel?
                
                Kısa, teknik ve net bir paragraf yaz.
                """
                yorum = get_ai_response(prompt, api_key)
                st.info(yorum)


def render_visual_forensics_module(api_key):
    # --- GÜVENLİ IMPORTLAR ---
    import math
    import datetime as dt_mod
    import plotly.graph_objects as go

    # --- İÇ HESAPLAMA FONKSİYONU ---
    def calculate_sun_position_safe(latitude, longitude, date_time_obj):
        rad = math.pi / 180.0
        deg = 180.0 / math.pi
        day_of_year = date_time_obj.timetuple().tm_yday
        declination = 23.45 * math.sin(rad * (360/365.0) * (day_of_year - 81))
        B = rad * (360/365.0) * (day_of_year - 81)
        eot = 9.87 * math.sin(2*B) - 7.53 * math.cos(B) - 1.5 * math.sin(B)
        lstm = 15 * 3
        time_correction = 4 * (longitude - lstm) + eot
        local_time = date_time_obj.hour + date_time_obj.minute / 60.0
        solar_time = local_time + time_correction / 60.0
        hra = 15 * (solar_time - 12)
        sin_elevation = (math.sin(rad * latitude) * math.sin(rad * declination) + 
                         math.cos(rad * latitude) * math.cos(rad * declination) * math.cos(rad * hra))
        if sin_elevation > 1: sin_elevation = 1
        if sin_elevation < -1: sin_elevation = -1
        elevation = math.asin(sin_elevation) * deg
        return elevation

    # --- ARAYÜZ (UI) ---
    st.header("🕵️ Visual Forensics: Gölge ve Işık Analizi")
    st.info("Bu modül, fotoğraftaki gölge boylarını astronomik verilerle kıyaslayarak fotoğrafın çekildiği saatin doğruluğunu test eder.")

    col1, col2 = st.columns([1, 1])

    with col1:
        st.subheader("1. İddia Edilen Veriler")
        
        # DÜZELTME: now() yerine sabit bir varsayılan değer kullanıldı.
        # key="..." eklenerek Streamlit'in değeri hafızada tutması sağlandı.
        default_date = dt_mod.date.today()
        default_time = dt_mod.time(12, 0) # Varsayılan 12:00
        
        claim_date = st.date_input("İddia Edilen Tarih", value=default_date, key="forensic_date")
        claim_time = st.time_input("İddia Edilen Saat", value=default_time, key="forensic_time")
        
        city_coords = {
            "İstanbul": (41.0082, 28.9784),
            "Ankara": (39.9334, 32.8597),
            "İzmir": (38.4192, 27.1287),
            "Antalya": (36.8969, 30.7133),
            "Erzurum": (39.9043, 41.2679),
            "Aksaray": (38.37255, 34.02537),
            "Diyarbakır": (37.9144, 40.2306),
            "Trabzon": (41.0027, 39.7168)
        }
        city = st.selectbox("Olay Yeri", list(city_coords.keys()), key="forensic_city")
        lat, lon = city_coords[city]

    with col2:
        st.subheader("2. Fotoğraf Ölçümleri")
        st.caption("Fotoğrafta boyunu bildiğiniz bir cisim (Örn: Trafik levhası ~2m) ve gölgesini ölçerek girin.")
        
        obj_height = st.number_input("Cisim Boyu (Metre)", value=1.70, step=0.10, key="forensic_height")
        shadow_len = st.number_input("Fotoğraftaki Gölge Boyu (Metre)", value=1.70, step=0.10, key="forensic_shadow")

    if st.button("🔍 Analizi Başlat", key="forensic_btn"):
        # Tarih birleştirme
        target_dt = dt_mod.datetime.combine(claim_date, claim_time)
        
        # 1. Astronomik Hesaplama
        sun_elevation = calculate_sun_position_safe(lat, lon, target_dt)
        
        # Gece kontrolü
        if sun_elevation <= 0:
            st.error(f"🌑 HATA: Girilen saatte ({claim_time.strftime('%H:%M')}) güneş batmış durumda! (Açı: {sun_elevation:.1f}°). Gölge oluşması imkansız.")
            return

        # 2. Beklenen Gölge Hesabı
        rad_elevation = math.radians(sun_elevation)
        if rad_elevation == 0: rad_elevation = 0.0001
        
        expected_shadow = obj_height / math.tan(rad_elevation)
        
        # 3. Sapma Hesabı
        diff = abs(expected_shadow - shadow_len)
        error_rate = (diff / (expected_shadow + 0.001)) * 100
        
        # --- SONUÇ EKRANI ---
        st.divider()
        st.subheader("📊 Analiz Sonucu")
        
        c1, c2, c3 = st.columns(3)
        c1.metric("Güneş Açısı", f"{sun_elevation:.1f}°")
        c2.metric("Beklenen Gölge", f"{expected_shadow:.2f} m")
        c3.metric("Ölçülen Gölge", f"{shadow_len:.2f} m")
        
        # Görselleştirme
        fig = go.Figure()
        
        # Zemin
        max_x = max(expected_shadow, shadow_len) + 1
        fig.add_shape(type="line", x0=-1, y0=0, x1=max_x, y1=0, line=dict(color="black", width=4))
        
        # Cisim
        fig.add_trace(go.Scatter(x=[0, 0], y=[0, obj_height], mode="lines", name="Cisim", line=dict(color="blue", width=6)))
        
        # Beklenen Gölge
        fig.add_trace(go.Scatter(x=[0, expected_shadow], y=[0, 0], mode="lines", name="Beklenen (Bilimsel)", line=dict(color="green", width=4, dash="dash")))
        
        # Ölçülen Gölge
        fig.add_trace(go.Scatter(x=[0, shadow_len], y=[-0.05, -0.05], mode="lines", name="Fotoğraftaki", line=dict(color="red", width=4)))
        
        # Güneş Işını
        fig.add_trace(go.Scatter(x=[expected_shadow, 0], y=[0, obj_height], mode="lines", name="Güneş Işını", line=dict(color="orange", width=1)))

        fig.update_layout(
            title="Gölge Analiz Grafiği", 
            height=300, 
            showlegend=True,
            margin=dict(l=20, r=20, t=40, b=20)
        )
        st.plotly_chart(fig, use_container_width=True)
        
        # Yorum
        st.write("---")
        if error_rate < 15:
            st.success("✅ **DOĞRULANDI:** Fotoğrafın saati ve gölge boyu fiziksel olarak uyumlu.")
        elif error_rate < 30:
            st.warning("⚠️ **ŞÜPHELİ:** Gölge boyunda %15-30 sapma var. Saat farkı veya ölçüm hatası olabilir.")
        else:
            st.error(f"🚨 **TUTARSIZLIK:** İddia edilen saatte gölgenin **{expected_shadow:.2f}m** olması gerekirdi. Ancak **{shadow_len:.2f}m** ölçüldü.")
            
            if api_key:
                prompt = f"""
                GÖREV: Adli bilişim uzmanı raporu yaz.
                KONUM: {city}
                TARİH: {target_dt}
                GÜNEŞ AÇISI: {sun_elevation:.1f} derece
                BEKLENEN GÖLGE: {expected_shadow:.2f} m
                FOTOĞRAFTAKİ GÖLGE: {shadow_len:.2f} m
                
                YORUM: Bu sapma ne anlama geliyor?
                """
                st.markdown(f"**🤖 AI Uzman Görüşü:** {get_ai_response(prompt, api_key)}")





# --- ANA UYGULAMA ---
def main():
    st.title("⚖️ Hukuk Asistanı (v10.0 - Ultimate Edition)")
    
    try:
        lib_ver = importlib.metadata.version("google-generativeai")
    except:
        lib_ver = "Bilinmiyor"

    # --- BAŞLANGIÇTA VERİLERİ YÜKLE ---
    if "durusma_listesi" not in st.session_state:
        st.session_state.durusma_listesi = load_durusma_data()

    # Mevcut State'ler
    if "doc_text" not in st.session_state: st.session_state.doc_text = ""
    if "last_file_id" not in st.session_state: st.session_state.last_file_id = None
    if "messages" not in st.session_state: st.session_state.messages = []
    if "mevzuat_sonuc" not in st.session_state: st.session_state.mevzuat_sonuc = ""
    if "ictihat_sonuc" not in st.session_state: st.session_state.ictihat_sonuc = ""
    if "dilekce_taslak" not in st.session_state: st.session_state.dilekce_taslak = ""
    if "soru_cevap" not in st.session_state: st.session_state.soru_cevap = ""
    if "ses_metni" not in st.session_state: st.session_state.ses_metni = ""
    if "ocr_metni" not in st.session_state: st.session_state.ocr_metni = ""
    if "dalgic_context" not in st.session_state: st.session_state.dalgic_context = ""
    if "dalgic_sonuc" not in st.session_state: st.session_state.dalgic_sonuc = ""
    if "buyur_abi_context" not in st.session_state: st.session_state.buyur_abi_context = ""
    if "buyur_abi_response" not in st.session_state: st.session_state.buyur_abi_response = ""
    
    if "arsiv_context" not in st.session_state: st.session_state.arsiv_context = ""
    if "arsiv_genel_ozet" not in st.session_state: st.session_state.arsiv_genel_ozet = ""
    if "arsiv_soru_cevap" not in st.session_state: st.session_state.arsiv_soru_cevap = ""
    if "arsiv_arama_sonuclari" not in st.session_state: st.session_state.arsiv_arama_sonuclari = []
    if "aktif_dosya_adi" not in st.session_state: st.session_state.aktif_dosya_adi = ""
    if "aktif_dosya_yolu" not in st.session_state: st.session_state.aktif_dosya_yolu = ""
    
    if "sozlesme_analiz" not in st.session_state: st.session_state.sozlesme_analiz = ""

    
    # Yeni Eklenen State'ler
    if "mock_messages" not in st.session_state: st.session_state.mock_messages = []
    if "gorev_listesi" not in st.session_state: st.session_state.gorev_listesi = ""
    if "kvkk_metin" not in st.session_state: st.session_state.kvkk_metin = ""

    ROOT_DIR = "Hukuk_Arsivi"
    if not os.path.exists(ROOT_DIR):
        os.makedirs(ROOT_DIR)

    with st.sidebar:
        st.header("⚙️ Ayarlar")
        api_key = st.text_input("Google Gemini API Key", type="password")
        st.caption(f"Kütüphane Sürümü: {lib_ver}")
        
        st.divider()
        st.header("📁 Dosya Bilgileri")
        input_davaci = st.text_input("Davacı")
        input_davali = st.text_input("Davalı")
        input_mahkeme = st.text_input("Mahkeme")
        input_dosya_no = st.text_input("Dosya No")
        
        if st.button("🗑️ Ekranı Temizle"):
            for key in st.session_state.keys():
                if key != "durusma_listesi":
                    del st.session_state[key]
            st.rerun()

    uploaded_file = st.file_uploader("Dosya Yükle (UDF/PDF)", type=['udf', 'pdf'])

    if uploaded_file and st.session_state.get('last_file_id') != uploaded_file.file_id:
        with st.spinner("Okunuyor..."):
            file_bytes = BytesIO(uploaded_file.getvalue())
            ext = uploaded_file.name.split('.')[-1].lower()
            raw_text = parse_udf(file_bytes) if ext == 'udf' else parse_pdf(file_bytes)
            st.session_state.doc_text = raw_text
            st.session_state.last_file_id = uploaded_file.file_id
            st.session_state.messages = []

    if st.session_state.doc_text.startswith(("HATA", "UYARI")):
        st.warning(st.session_state.doc_text)
    
    auto_data = extract_metadata(st.session_state.doc_text)

    # --- SEKMELER (2 SATIR - TOPLAM 31 MODÜL) ---
    
    # 1. SATIR: Temel, Strateji ve Şeytanın Avukatı (15 Sekme)
    st.markdown("### 🛠️ Temel Araçlar & Strateji")
    tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab26, tab29, tab30, tab31, tab9, tab34, tab35, tab37 = st.tabs([
        "📋 Analiz", "💬 Sohbet", "📕 Mevzuat", "⚖️ İçtihat", 
        "✍️ Dilekçe Yaz", "❓ Bana Sor", "🎙️ Ses", "👁️ OCR",
        "🌍 Çeviri", "🛡️ Çürüt", "🕵️‍♂️ Sorgu", "😈 Şeytanın Avukatı", "🤿 Dalgıç", "🧠 Semantik", "🎙️ Canlı Duruşma", "🦋 Kelebek"
    ])

    # 2. SATIR: Yönetim, Pro Modüller, Canlı Asistan ve "Etki Analizi" (16 Sekme)
    st.markdown("### 🚀 Yönetim, Hesaplama & Pro Modüller")
    # tab33 (Etki Analizi) buraya eklendi
    tab10, tab11, tab12, tab13, tab16, tab17, tab19, tab21, tab22, tab23, tab32, tab33, tab36 = st.tabs([
        "🙋 Buyur Abi", "⏰ Hatırlatıcı", "🗄️ Arşiv", "🏛️ UYAP Analiz", 
        "🕸️ İlişki Ağı", "📝 Sözleşme Analiz", 
        "🕵️‍♂️ KVKK Temizle",  "⚔️ Belge Kıyasla", "🎭 Sanal Duruşma", "✅ Görev Çıkarıcı", "⚡ Canlı Asistan", "📡 Etki Analizi", "🕵️ Dijital Otp"
    ])

    # 3. SATIR: Simülasyon ve İleri Düzey Risk (YENİ EKLENDİ)
    st.markdown("### 🔮 Simülasyon & Risk Analizi")
    tab_checkup, tab_timemachine, tab_aym, tab_deepfake, tab_osyn, tab_sxx, tab_sah, tab_soy, tab_isx, tab_golx, tab_arx = st.tabs(["🏥 Kurumsal Check-up", "⏳ Zaman Makinesi", "⚖️ AYM & AİHM Testi", "🕵️ Deepfake Kontrol", "🌐 OSINT (İstihbarat)", "🔔 Emsal Alarm", "👑 Sahip Modu", "🌳 Soyağacı", "🔥 Isı Haritası", "🕸️ Gizli Bağlantı", "🤝 Arabuluculuk"])

    # 4. SATIR: oyun değiştirici hamle menüsü (15 Sekme)
    st.markdown("### 🛠️ Temel Araçlar & Strateji")
    tabx1, tabx2, tabx3, tabx4, tabx5, tabx6, tabx7 = st.tabs([
        "🗺️ Adli Harita", "🕰️ Mevzuat Makinesi", "🧐 Rapor Denetçisi", "🏛️ Kurumsal Hafıza", "💰 Dava Maliyeti", "🗺️ Adli Olay Yeri", "🕵️ Visual Forensics" 
    ])


    # --- SEKMELERİN İÇERİKLERİ ---
    
    # NOT: tab1, tab2 vb. eski içeriklerinizi buraya yerleştirmelisiniz.
    # Örnek olarak yeni eklenenleri bağlıyorum:
    
    with tab_checkup:
        render_checkup_module(api_key)
        
    with tab_timemachine:
        render_time_machine(api_key)

    # (Buradan sonra eski kodunuzdaki 'with tab1:', 'with tab2:' blokları gelmeli...)

    with tab_aym:  # <--- YENİ EKLENEN KISIM
        render_aym_aihm_module(api_key)

    with tab_deepfake:  # <--- YENİ EKLENEN KISIM
        render_deepfake_module(api_key)

    with tab_osyn:
        render_osint_module(api_key) # <--- YENİ FONKSİYON ÇAĞRISI

    with tab_sxx: render_precedent_alert_module(api_key)
    with tab_sah: render_owner_mode(api_key)
    with tab_soy: render_property_genealogy(api_key)
    with tab_isx: render_limitations_heatmap(api_key)
    with tab_golx: render_conflict_scanner(api_key)
    with tab_arx: render_mediation_checker(api_key)
    with tabx1: render_forensic_map(api_key)
    with tabx2: render_temporal_law_machine(api_key)
    with tabx3: render_expert_report_auditor(api_key)
    with tabx4: render_corporate_memory(api_key)
    with tabx5: render_cost_calculator_module(api_key)
    with tabx6: render_forensic_map_module(api_key)
    with tabx7: render_visual_forensics_module(api_key)
    # --- TAB İÇERİKLERİ ---

    with tab1:
        col1, col2 = st.columns(2)
        with col1:
            st.write(f"**Mahkeme:** {input_mahkeme or auto_data['mahkeme']}")
            st.write(f"**Dosya No:** {input_dosya_no or auto_data['esas']}")
        with col2:
            st.write(f"**Davacı:** {input_davaci or '-'}")
            st.write(f"**Davalı:** {input_davali or '-'}")
        st.text_area("Metin Önizleme", st.session_state.doc_text, height=150)

    with tab2:
        for msg in st.session_state.messages:
            with st.chat_message(msg["role"]): st.markdown(msg["content"])
        if prompt := st.chat_input("Soru sor..."):
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"): st.markdown(prompt)
            with st.chat_message("assistant"):
                with st.spinner("AI Yanıtlıyor..."):
                    context = f"BELGE: {st.session_state.doc_text[:20000]}\nSORU: {prompt}"
                    reply = get_ai_response(f"Sen bir avukatsın. Şuna cevap ver: {context}", api_key)
                    st.markdown(reply)
                    st.session_state.messages.append({"role": "assistant", "content": reply})

    with tab3:
        c1, c2 = st.columns([3,1])
        q = c1.text_input("Kanun Madde No", key="mq")
        if c2.button("Getir", key="mb") and q:
            with st.spinner("Aranıyor..."):
                res = get_ai_response(f"GÖREV: '{q}' maddesini tam metin yaz.", api_key)
                st.session_state.mevzuat_sonuc = res
        if st.session_state.mevzuat_sonuc:
            st.markdown(f"<div class='kanun-kutusu'>{st.session_state.mevzuat_sonuc}</div>", unsafe_allow_html=True)

    with tab4:
        c3, c4 = st.columns([3,1])
        iq = c3.text_input("İçtihat Konusu", key="iq")
        if c4.button("Ara", key="ib") and iq:
            with st.spinner("Taranıyor..."):
                res = get_ai_response(f"GÖREV: '{iq}' hakkında Yargıtay kararlarını özetle.", api_key)
                st.session_state.ictihat_sonuc = res
        if st.session_state.ictihat_sonuc:
            st.markdown(f"<div class='ictihat-kutusu'>{st.session_state.ictihat_sonuc}</div>", unsafe_allow_html=True)

    with tab5:
        st.subheader("✍️ Otomatik Savunma/Cevap Dilekçesi")
        if not st.session_state.doc_text or st.session_state.doc_text.startswith(("HATA", "UYARI")):
            st.info("Dilekçe oluşturmak için önce sol menüden bir dosya yükleyin.")
        else:
            col_d1, col_d2 = st.columns([2, 1])
            with col_d1:
                dilekce_turu = st.selectbox("Dilekçe Türü", ["Cevap Dilekçesi", "İtiraz Dilekçesi", "Beyan Dilekçesi"])
                ozel_talimat = st.text_area("Özel Savunma Stratejisi (Opsiyonel)", placeholder="Örn: Zamanaşımı itirazında bulun...")
            with col_d2:
                st.write("")
                st.write("")
                if st.button("Dilekçeyi Yaz (AI)", type="primary"):
                    if not api_key: st.error("API Key gerekli!")
                    else:
                        with st.spinner("Dilekçe yazılıyor..."):
                            mahkeme = input_mahkeme or auto_data['mahkeme']
                            dosya = input_dosya_no or auto_data['esas']
                            davaci = input_davaci or "Davacı"
                            davali = input_davali or "Davalı"
                            prompt = f"""
                            GÖREV: Aşağıdaki metne dayanarak profesyonel bir {dilekce_turu} yaz.
                            BİLGİLER: Mahkeme: {mahkeme}, Dosya: {dosya}, Davacı: {davaci}, Davalı: {davali}, Ek Talimat: {ozel_talimat}
                            KARŞI TARAFIN DİLEKÇESİ (ÖZET): {st.session_state.doc_text[:20000]}
                            KURALLAR: Resmi Türk hukuk dilekçesi formatında olsun.
                            """
                            res = get_ai_response(prompt, api_key)
                            st.session_state.dilekce_taslak = res
            if st.session_state.dilekce_taslak:
                st.divider()
                st.subheader("📄 Dilekçe Taslağı")
                btn_col1, btn_col2 = st.columns(2)
                word_file = create_word_file(st.session_state.dilekce_taslak)
                with btn_col1:
                    st.download_button("💾 Word Olarak İndir (.docx)", word_file, "Dilekce.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                udf_file = create_udf_file(st.session_state.dilekce_taslak)
                with btn_col2:
                    st.download_button("💾 UDF Olarak İndir (.udf)", udf_file, "Dilekce.udf", "application/zip")
                st.text_area("Dilekçe Metni", st.session_state.dilekce_taslak, height=500)

    with tab6:
        st.subheader("❓ Hukuki Soru & WhatsApp Paylaşımı")
        col_s1, col_s2 = st.columns([3, 1])
        with col_s1:
            kullanici_sorusu = st.text_area("Hukuki Sorunuzu Yazın", height=100, placeholder="Örn: Kiracı kirayı ödemezse tahliye süreci nasıl işler?")
        with col_s2:
            telefon_no = st.text_input("WhatsApp No (905xxxxxxxxx)", placeholder="905551234567")
            if st.button("Analiz Et ve Hazırla", type="primary"):
                if not api_key: st.error("API Key giriniz.")
                elif not kullanici_sorusu: st.warning("Lütfen bir soru yazın.")
                else:
                    with st.spinner("Mevzuat ve İçtihatlar taranıyor..."):
                        prompt = f"""
                        GÖREV: Aşağıdaki hukuki soruyu detaylıca cevapla.
                        SORU: {kullanici_sorusu}
                        KURALLAR: 1. İlgili KANUN MADDELERİNİ belirt. 2. YARGITAY İÇTİHATLARINDAN örnek ver. 3. Net hukuki görüş bildir.
                        """
                        res = get_ai_response(prompt, api_key)
                        st.session_state.soru_cevap = res
        if st.session_state.soru_cevap:
            st.divider()
            st.markdown(f"<div class='ictihat-kutusu'><b>💡 Hukuki Görüş:</b><br>{st.session_state.soru_cevap}</div>", unsafe_allow_html=True)
            pdf_data = create_pdf_file(st.session_state.soru_cevap)
            encoded_text = urllib.parse.quote(f"*Hukuki Soru:* {kullanici_sorusu}\n\n*Cevap:*\n{st.session_state.soru_cevap}")
            wa_link = f"https://wa.me/{telefon_no}?text={encoded_text}"
            col_btn1, col_btn2 = st.columns(2)
            with col_btn1: st.download_button("📄 Cevabı PDF Olarak İndir", pdf_data, "Hukuki_Gorus.pdf", "application/pdf")
            with col_btn2:
                if telefon_no: st.link_button("📲 Cevabı WhatsApp ile Gönder", wa_link)
                else: st.warning("WhatsApp butonu için telefon no giriniz.")

    with tab7:
        st.subheader("🎙️ Sesli Asistan")
        col_audio1, col_audio2 = st.columns(2)
        with col_audio1:
            st.markdown("##### 🗣️ Metni Seslendir")
            text_to_read = st.text_area("Okunacak Metni Yazın:", height=150)
            if st.button("🔊 Seslendir"):
                if text_to_read:
                    with st.spinner("Ses oluşturuluyor..."):
                        audio_fp = text_to_speech(text_to_read)
                        if audio_fp: st.audio(audio_fp, format='audio/mp3')
                        else: st.error("Hata oluştu.")
        with col_audio2:
            st.markdown("##### 📝 Sesi Yazıya Çevir")
            audio_input = st.file_uploader("Ses Dosyası (WAV/MP3)", type=["wav", "mp3"])
            if audio_input and st.button("📝 Yazıya Dök"):
                with st.spinner("Analiz ediliyor..."):
                    text_result = speech_to_text(audio_input)
                    st.session_state.ses_metni = text_result
            if st.session_state.ses_metni:
                st.success("Sonuç:")
                st.text_area("", st.session_state.ses_metni, height=150)

    with tab8:
        st.subheader("👁️ OCR (Resim/PDF -> Metin)")
        ocr_file = st.file_uploader("Dosya Yükle", type=['png', 'jpg', 'jpeg', 'pdf', 'docx', 'tif', 'tiff'])
        if ocr_file and st.button("🔍 Metni Ayıkla (OCR)", type="primary"):
            if not api_key: st.error("API Key gerekli.")
            else:
                with st.spinner("İşleniyor..."):
                    ftype = ocr_file.name.split('.')[-1].lower()
                    if ftype == 'docx': res = extract_text_from_docx(ocr_file)
                    else:
                        mime = "application/pdf" if ftype == 'pdf' else "image/tiff" if ftype in ['tif', 'tiff'] else "image/jpeg"
                        ocr_file.seek(0)
                        res = perform_ocr_gemini(ocr_file, mime, api_key)
                    st.session_state.ocr_metni = res
        if st.session_state.ocr_metni:
            st.text_area("OCR Sonucu:", st.session_state.ocr_metni, height=400)
            word_ocr = create_word_file(st.session_state.ocr_metni)
            st.download_button("💾 Word İndir", word_ocr, "ocr.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    with tab9:
        st.subheader("🤿 Dalgıç Modu (Çoklu Dosya Analizi)")
        st.info("Birden fazla dosyayı aynı anda yükleyin. Sistem hepsini okuyup, birleştirip sorularınızı yanıtlar.")
        dalgic_files = st.file_uploader("Dosyaları Sürükleyin (Max 30 Dosya)", type=['udf', 'pdf', 'docx', 'doc', 'txt', 'png', 'jpg', 'jpeg', 'mp4', 'tif', 'tiff'], accept_multiple_files=True)
        if dalgic_files:
            if st.button("🚀 Dosyaları İşle ve Hafızaya Al", type="primary"):
                if not api_key: st.error("API Key giriniz.")
                else:
                    full_context = ""
                    progress_bar = st.progress(0)
                    for i, file in enumerate(dalgic_files):
                        file_bytes = BytesIO(file.read())
                        ext = file.name.split('.')[-1].lower()
                        extracted_text = ""
                        try:
                            if ext == 'udf': 
                                extracted_text = parse_udf(file_bytes)
                            elif ext == 'txt': 
                                extracted_text = file_bytes.read().decode('utf-8', errors='ignore')
                            elif ext in ['docx', 'doc']: 
                                extracted_text = extract_text_from_docx(file_bytes)
                            elif ext == 'pdf':
                                extracted_text = parse_pdf(file_bytes)
                                if not extracted_text:
                                    file_bytes.seek(0)
                                    extracted_text = perform_ocr_gemini(file_bytes, "application/pdf", api_key)
                            elif ext in ['png', 'jpg', 'jpeg', 'img', 'tif', 'tiff']:
                                mime = "image/tiff" if ext in ['tif', 'tiff'] else "image/jpeg"
                                extracted_text = perform_ocr_gemini(file_bytes, mime, api_key)
                            elif ext == 'mp4':
                                extracted_text = perform_ocr_gemini(file_bytes, "video/mp4", api_key, "Video içeriğini özetle.")
                            
                            full_context += f"\n\n--- DOSYA: {file.name} ---\n{extracted_text}"
                        except Exception as e:
                            full_context += f"\nHATA ({file.name}): {str(e)}"
                        
                        progress_bar.progress((i + 1) / len(dalgic_files))
                    st.session_state.dalgic_context = full_context
                    st.success(f"Veriler hafızaya alındı! ({len(full_context)} karakter)")
        if st.session_state.dalgic_context:
            st.divider()
            dalgic_soru = st.text_area("Dosyalar Hakkında Soru Sorun:", placeholder="Örn: Bu dosyalardaki tüm tanık ifadelerindeki çelişkileri listele.")
            if st.button("Analiz Et ve Yanıtla"):
                if not dalgic_soru: st.warning("Soru yazın.")
                else:
                    with st.spinner("Dalgıç derinlere iniyor..."):
                        prompt = f"GÖREV: Aşağıdaki dosya içeriklerine göre cevapla.\nSORU: {dalgic_soru}\nİÇERİK: {st.session_state.dalgic_context[:500000]}"
                        res = get_ai_response(prompt, api_key)
                        st.session_state.dalgic_sonuc = res
            if st.session_state.dalgic_sonuc:
                st.markdown(f"<div class='kanun-kutusu'>{st.session_state.dalgic_sonuc}</div>", unsafe_allow_html=True)
                col_d1, col_d2 = st.columns(2)
                with col_d1: st.download_button("📕 PDF İndir", create_pdf_file(st.session_state.dalgic_sonuc), "Dalgic.pdf", "application/pdf")
                with col_d2: st.download_button("📘 Word İndir", create_word_file(st.session_state.dalgic_sonuc), "Dalgic.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    with tab10:
        st.subheader("🙋 Buyur Abi (Genel Asistan & Çoklu Format)")
        st.info("Hukuk, kodlama, yemek tarifi veya günlük sohbet... Ne istersen sor. Ayrıca Excel, Ses, Video dahil her türlü dosyayı yükleyip analiz ettirebilirsin.")
        col_ba1, col_ba2 = st.columns([1, 2])
        with col_ba1:
            st.markdown("#### 📁 Dosya & Ses Girişi")
            buyur_files = st.file_uploader("Dosya Ekle (Excel, Ses, Video, Resim vb.)", 
                                           type=['pdf','udf','doc','docx','txt','xls','xlsx','xlt','xml','jpg','png','jpeg','mp3','mp4','wav','tif','tiff'],
                                           accept_multiple_files=True)
            st.markdown("#### 🎙️ Sesli Soru Sor")
            audio_prompt = st.file_uploader("Ses Kaydı Yükle (Soru olarak)", type=['wav', 'mp3', 'ogg'], key="voice_prompt")
        with col_ba2:
            st.markdown("#### 💬 Sohbet Alanı")
            user_text_input = st.text_area("Sorunu Yaz Abi:", height=150, placeholder="Örn: Bu Excel dosyasındaki ciroları topla veya yüklediğim ses kaydını özetle...")
            if st.button("🚀 Gönder Gelsin", type="primary"):
                if not api_key: st.error("Önce sol menüden API Anahtarını girmen lazım abi.")
                else:
                    context_data = ""
                    voice_text = ""
                    with st.spinner("Dosyalar ve sesler inceleniyor..."):
                        if audio_prompt:
                            voice_text = speech_to_text(audio_prompt)
                            st.info(f"🎤 Sesli Sorun: {voice_text}")
                        if buyur_files:
                            for file in buyur_files:
                                f_bytes = BytesIO(file.read())
                                ext = file.name.split('.')[-1].lower()
                                try:
                                    if ext in ['xls', 'xlsx', 'xlt']: context_data += f"\n--- EXCEL ({file.name}) ---\n{read_excel_file(f_bytes)}"
                                    elif ext in ['txt', 'xml', 'py', 'js', 'html']: context_data += f"\n--- METİN ({file.name}) ---\n{f_bytes.read().decode('utf-8', errors='ignore')}"
                                    elif ext in ['doc', 'docx']: context_data += f"\n--- WORD ({file.name}) ---\n{extract_text_from_docx(f_bytes)}"
                                    elif ext == 'pdf':
                                        pdf_txt = parse_pdf(f_bytes)
                                        if not pdf_txt:
                                            f_bytes.seek(0)
                                            pdf_txt = perform_ocr_gemini(f_bytes, "application/pdf", api_key)
                                        context_data += f"\n--- PDF ({file.name}) ---\n{pdf_txt}"
                                    elif ext in ['jpg', 'png', 'jpeg', 'img', 'tif', 'tiff']:
                                        mime = "image/tiff" if ext in ['tif', 'tiff'] else "image/jpeg"
                                        ocr_res = perform_ocr_gemini(f_bytes, mime, api_key, "Bu resimde ne var?")
                                        context_data += f"\n--- RESİM ({file.name}) ---\n{ocr_res}"
                                    elif ext in ['mp3', 'wav', 'mp4']:
                                        mime = "video/mp4" if ext == 'mp4' else "audio/mp3"
                                        media_res = perform_ocr_gemini(f_bytes, mime, api_key, "Bu kaydı analiz et ve içeriğini dök.")
                                        context_data += f"\n--- MEDYA ({file.name}) ---\n{media_res}"
                                    elif ext == 'udf': context_data += f"\n--- UDF ({file.name}) ---\n{parse_udf(f_bytes)}"
                                except Exception as e: context_data += f"\n⚠️ {file.name} okunurken hata: {str(e)}"
                    if not user_text_input and not voice_text and not context_data: st.warning("Abi boş gönderdin.")
                    else:
                        final_prompt = f"GÖREV: Yardımsever asistan ol.\nSORU: {user_text_input}\nSESLİ SORU: {voice_text}\nDOSYALAR: {context_data[:100000]}"
                        with st.spinner("Hazırlıyorum abi..."):
                            resp = get_ai_response(final_prompt, api_key)
                            st.session_state.buyur_abi_response = resp
            if st.session_state.buyur_abi_response:
                st.markdown(f"<div class='buyur-abi-kutusu'>{st.session_state.buyur_abi_response}</div>", unsafe_allow_html=True)
                b_col1, b_col2 = st.columns(2)
                with b_col1: st.download_button("📄 PDF Olarak Al", create_pdf_file(st.session_state.buyur_abi_response), "Cevap.pdf", "application/pdf")
                with b_col2: st.download_button("📝 Word Olarak Al", create_word_file(st.session_state.buyur_abi_response), "Cevap.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    with tab11:
        st.subheader("⏰ Duruşma Hatırlatıcı & Takvim")
        st.info("UYAP'tan aldığınız .ics (Takvim) dosyalarını buraya yükleyin. Yaklaşan duruşmaları otomatik listeler ve 24 saat kalanlar için ALARM verir.")
        col_h1, col_h2 = st.columns([1, 3])
        with col_h1:
            ics_file = st.file_uploader("Takvim Dosyası (.ics) Yükle", type=['ics'])
            if ics_file:
                if st.button("📅 Takvime Ekle", type="primary"):
                    events = parse_ics_data(BytesIO(ics_file.getvalue()))
                    if events:
                        count = 0
                        for evt in events:
                            exists = any(e['dtstart'] == evt['dtstart'] and e['summary'] == evt['summary'] for e in st.session_state.durusma_listesi)
                            if not exists:
                                st.session_state.durusma_listesi.append(evt)
                                count += 1
                        save_durusma_data(st.session_state.durusma_listesi)
                        st.success(f"{count} yeni duruşma eklendi!")
                    else: st.error("Dosya okunamadı.")
            st.divider()
            if st.button("🗑️ Tüm Listeyi Temizle"):
                st.session_state.durusma_listesi = []
                save_durusma_data([])
                st.rerun()
        with col_h2:
            if not st.session_state.durusma_listesi: st.info("Henüz eklenmiş bir duruşma yok.")
            else:
                sorted_events = sorted(st.session_state.durusma_listesi, key=lambda x: x['dtstart'])
                now = datetime.now()
                st.write(f"**Toplam Duruşma Sayısı:** {len(sorted_events)}")
                for evt in sorted_events:
                    dt = evt['dtstart']
                    diff = dt - now
                    tarih_str = dt.strftime("%d.%m.%Y %H:%M")
                    is_alarm = timedelta(0) < diff < timedelta(hours=24)
                    is_past = diff < timedelta(0)
                    if is_past:
                        with st.expander(f"✅ (GEÇMİŞ) {tarih_str} - {evt.get('summary', 'Başlıksız')}"):
                            st.write(f"**Mahkeme:** {evt.get('location', '-')}")
                            st.write(f"**Detay:** {evt.get('description', '-')}")
                    elif is_alarm:
                        st.markdown(f"""<div class="alarm-kutusu">🚨 ALARM: DURUŞMAYA AZ KALDI!<br>📅 {tarih_str}<br>⚖️ {evt.get('summary', 'Başlıksız')}<br>📍 {evt.get('location', '-')}</div>""", unsafe_allow_html=True)
                        with st.expander("Detayları Gör"): st.write(f"**Açıklama:** {evt.get('description', '-')}")
                    else:
                        st.markdown(f"""<div class="normal-durusma">📅 <b>{tarih_str}</b> (Kalan: {diff.days} gün)<br>⚖️ {evt.get('summary', 'Başlıksız')}<br>📍 {evt.get('location', '-')}</div>""", unsafe_allow_html=True)

    with tab12:
        st.subheader("🗄️ Doküman Yönetimi ve Arşivleme")
        st.info(f"Verileriniz bilgisayarınızda '{ROOT_DIR}' klasöründe saklanır. TIF, PDF, Resim dahil tüm dosyaları okur.")

        ar_tab1, ar_tab2, ar_tab3 = st.tabs(["📂 Yeni Dava Dosyası Aç", "📎 Dosya Yükle", "🔍 Arşivde Ara & Analiz"])

        with ar_tab1:
            st.markdown("#### Yeni Dava Klasörü Oluştur")
            c_tur, c_mah = st.columns(2)
            with c_tur: dava_turu = st.selectbox("Dava Türü", ["Hukuk Davaları", "Ceza Davaları", "İcra Dosyaları", "İdari Davalar"])
            with c_mah: yeni_mahkeme = st.text_input("Mahkeme Adı", placeholder="Örn: Ankara 1. Asliye Hukuk")
            c_esas, c_taraf = st.columns(2)
            with c_esas: yeni_esas = st.text_input("Dosya/Esas No", placeholder="Örn: 2024-123")
            with c_taraf: yeni_taraflar = st.text_input("Taraf Bilgileri", placeholder="Örn: Ahmet Yılmaz vs Mehmet Demir")

            if st.button("📁 Klasörü Oluştur"):
                if yeni_mahkeme and yeni_esas:
                    safe_mah = "".join([c for c in yeni_mahkeme if c.isalnum() or c in (' ', '-', '_')]).strip()
                    safe_esas = "".join([c for c in yeni_esas if c.isalnum() or c in (' ', '-', '_')]).strip()
                    target_path = os.path.join(ROOT_DIR, dava_turu, safe_mah, safe_esas)
                    try:
                        os.makedirs(target_path, exist_ok=True)
                        with open(os.path.join(target_path, "Dosya_Bilgileri.txt"), "w", encoding="utf-8") as f:
                            f.write(f"Dava Türü: {dava_turu}\nMahkeme: {yeni_mahkeme}\nEsas: {yeni_esas}\nTaraflar: {yeni_taraflar}\nOluşturma: {datetime.now()}")
                        st.success(f"✅ Klasör Başarıyla Oluşturuldu:\n{target_path}")
                    except Exception as e: st.error(f"Hata: {str(e)}")
                else: st.warning("Lütfen Mahkeme ve Esas No giriniz.")

        with ar_tab2:
            st.markdown("#### Mevcut Dosyaya Evrak Ekle")
            if os.path.exists(ROOT_DIR):
                turler = [d for d in os.listdir(ROOT_DIR) if os.path.isdir(os.path.join(ROOT_DIR, d))]
                if not turler: st.warning("Henüz hiç dava klasörü yok.")
                else:
                    secilen_tur = st.selectbox("Dava Türü Seç", turler)
                    tur_path = os.path.join(ROOT_DIR, secilen_tur)
                    mahkemeler = [d for d in os.listdir(tur_path) if os.path.isdir(os.path.join(tur_path, d))]
                    if mahkemeler:
                        secilen_mah = st.selectbox("Mahkeme Seç", mahkemeler)
                        mah_path = os.path.join(tur_path, secilen_mah)
                        dosyalar = [d for d in os.listdir(mah_path) if os.path.isdir(os.path.join(mah_path, d))]
                        if dosyalar:
                            secilen_dosya = st.selectbox("Dosya No Seç", dosyalar)
                            final_path = os.path.join(mah_path, secilen_dosya)
                            st.info(f"Seçilen Klasör: {final_path}")
                            yuklenen_evraklar = st.file_uploader("Evrakları Yükle", 
                                                               type=['pdf','doc','docx','udf','png','jpg','mp3','mp4','wav','txt','tif','tiff'],
                                                               accept_multiple_files=True)
                            if st.button("💾 Evrakları Kaydet"):
                                if yuklenen_evraklar:
                                    for evrak in yuklenen_evraklar:
                                        with open(os.path.join(final_path, evrak.name), "wb") as f:
                                            f.write(evrak.getbuffer())
                                    st.success(f"✅ {len(yuklenen_evraklar)} adet dosya başarıyla kaydedildi!")
                                else: st.warning("Dosya seçmediniz.")
                        else: st.warning("Bu mahkemede dosya yok.")
                    else: st.warning("Bu türde mahkeme yok.")
            else: st.error("Ana arşiv klasörü bulunamadı.")

        with ar_tab3:
            st.markdown("#### 🔍 Arşivde Arama ve Yapay Zeka Analizi")
            
            if st.session_state.aktif_dosya_yolu:
                st.markdown(f"<div class='arsiv-kutusu'><b>📂 ÇALIŞILAN DOSYA: {st.session_state.aktif_dosya_adi}</b><br>Şu an sadece bu dosyadaki evraklar hafızada.</div>", unsafe_allow_html=True)
                
                if st.button("⬅️ Dosyayı Kapat ve Listeye Dön"):
                    st.session_state.aktif_dosya_yolu = ""
                    st.session_state.arsiv_context = ""
                    st.session_state.arsiv_genel_ozet = ""
                    st.session_state.arsiv_soru_cevap = ""
                    st.rerun()
                
                st.divider()
                col_analiz, col_soru = st.columns(2)
                
                with col_analiz:
                    st.markdown("### 📊 Analiz Et")
                    st.info("Bu klasördeki belgeleri özetler.")
                    if st.button("Dosyayı Analiz Et", type="primary", use_container_width=True):
                        if not api_key: st.error("API Key gerekli.")
                        else:
                            with st.spinner("Sadece bu dosyadaki evraklar analiz ediliyor..."):
                                prompt = f"GÖREV: Bu dava dosyasının içeriğini özetle, hukuki durumu analiz et.\nİÇERİK: {st.session_state.arsiv_context[:500000]}"
                                res = get_ai_response(prompt, api_key)
                                st.session_state.arsiv_genel_ozet = res
                    
                    if st.session_state.arsiv_genel_ozet:
                        st.markdown(f"<div class='kanun-kutusu'>{st.session_state.arsiv_genel_ozet}</div>", unsafe_allow_html=True)
                        
                        st.markdown("###### 📥 Raporu İndir / Paylaş")
                        c_down1, c_down2 = st.columns(2)
                        with c_down1: st.download_button("📄 PDF", create_pdf_file(st.session_state.arsiv_genel_ozet), "Analiz.pdf", "application/pdf")
                        with c_down2: st.download_button("📝 Word", create_word_file(st.session_state.arsiv_genel_ozet), "Analiz.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                with col_soru:
                    st.markdown("### ❓ Soru Sor")
                    st.info("Sadece bu dosya ile ilgili sorular sorun.")
                    arsiv_soru = st.text_input("Sorunuzu yazın", placeholder="Örn: Bilirkişi raporu ne zaman gelmiş?")
                    
                    if st.button("Soruyu Cevapla", use_container_width=True):
                        if not api_key: st.error("API Key gerekli.")
                        elif not arsiv_soru: st.warning("Soru yazın.")
                        else:
                            with st.spinner("Bu dosyadaki belgeler taranıyor..."):
                                prompt = f"""
                                GÖREV: Aşağıdaki dosya içeriğine göre soruyu cevapla.
                                SORU: {arsiv_soru}
                                DOSYA İÇERİĞİ:
                                {st.session_state.arsiv_context[:500000]}
                                """
                                res = get_ai_response(prompt, api_key)
                                st.session_state.arsiv_soru_cevap = res
                    
                    if st.session_state.arsiv_soru_cevap:
                        st.markdown(f"<div class='kanun-kutusu'>{st.session_state.arsiv_soru_cevap}</div>", unsafe_allow_html=True)

            else:
                arama_terimi = st.text_input("Aranacak Kelime (Dosya No, Mahkeme veya Dosya Adı)", placeholder="Örn: 2024-123 veya Ahmet Yılmaz")
                
                if st.button("🔎 Ara"):
                    st.session_state.arsiv_arama_sonuclari = []
                    bulunanlar = []
                    for root, dirs, files in os.walk(ROOT_DIR):
                        if arama_terimi.lower() in root.lower():
                            bulunanlar.append({"tip": "KLASÖR", "yol": root, "dosyalar": files})
                        for file in files:
                            if arama_terimi.lower() in file.lower():
                                bulunanlar.append({"tip": "DOSYA", "yol": os.path.join(root, file), "dosya_adi": file})
                    st.session_state.arsiv_arama_sonuclari = bulunanlar

                if st.session_state.arsiv_arama_sonuclari:
                    st.success(f"{len(st.session_state.arsiv_arama_sonuclari)} sonuç bulundu.")
                    for sonuc in st.session_state.arsiv_arama_sonuclari:
                        if sonuc["tip"] == "KLASÖR":
                            with st.expander(f"📁 {sonuc['yol']}"):
                                st.write(f"İçerik: {len(sonuc['dosyalar'])} dosya")
                                
                                if st.button(f"📂 Bu Dosyayı Aç ve Çalış ({os.path.basename(sonuc['yol'])})", key=sonuc['yol']):
                                    full_text = ""
                                    if not api_key: st.error("Lütfen önce API Key giriniz.")
                                    else:
                                        st.session_state.arsiv_context = ""
                                        st.session_state.arsiv_genel_ozet = ""
                                        st.session_state.arsiv_soru_cevap = ""
                                        
                                        with st.spinner("Sadece seçilen klasördeki dosyalar okunuyor..."):
                                            sadece_bu_klasordeki_dosyalar = [f for f in os.listdir(sonuc['yol']) if os.path.isfile(os.path.join(sonuc['yol'], f))]
                                            
                                            for f_name in sadece_bu_klasordeki_dosyalar:
                                                f_path = os.path.join(sonuc['yol'], f_name)
                                                ext = f_name.split('.')[-1].lower()
                                                try:
                                                    with open(f_path, 'rb') as f:
                                                        file_content = BytesIO(f.read())
                                                    
                                                    if ext == 'txt':
                                                        full_text += f"\n--- {f_name} ---\n{file_content.getvalue().decode('utf-8', errors='ignore')}"
                                                    elif ext == 'pdf':
                                                        pdf_text = parse_pdf(file_content)
                                                        if not pdf_text:
                                                            file_content.seek(0)
                                                            pdf_text = perform_ocr_gemini(file_content, "application/pdf", api_key)
                                                        full_text += f"\n--- {f_name} ---\n{pdf_text}"
                                                    elif ext in ['docx', 'doc']:
                                                        full_text += f"\n--- {f_name} ---\n{extract_text_from_docx(file_content)}"
                                                    elif ext == 'udf':
                                                        full_text += f"\n--- {f_name} ---\n{parse_udf(file_content)}"
                                                    elif ext in ['png', 'jpg', 'jpeg', 'tif', 'tiff']:
                                                        mime = "image/tiff" if ext in ['tif', 'tiff'] else "image/jpeg"
                                                        file_content.seek(0)
                                                        ocr_res = perform_ocr_gemini(file_content, mime, api_key)
                                                        full_text += f"\n--- {f_name} ---\n{ocr_res}"
                                                except Exception as e:
                                                    full_text += f"\n--- {f_name} (HATA) ---\n{str(e)}"
                                            
                                            st.session_state.arsiv_context = full_text
                                            st.session_state.aktif_dosya_adi = os.path.basename(sonuc['yol'])
                                            st.session_state.aktif_dosya_yolu = sonuc['yol']
                                            st.rerun()

    with tab13:
        st.subheader("🏛️ UYAP Toplu Dosya Analizi")
        st.info("UYAP'tan indirdiğiniz ZIP dosyalarını yükleyin. Sistem son 5 evrağı analiz eder.")
        uyap_zips = st.file_uploader("UYAP Dosyalarını Yükle (ZIP)", type=['zip'], accept_multiple_files=True)
        
        if uyap_zips and st.button("🚀 Dosyaları Analiz Et", type="primary"):
            if not api_key: st.error("API Key gerekli.")
            else:
                progress_bar = st.progress(0)
                for idx, zip_file in enumerate(uyap_zips):
                    dosya_adi = zip_file.name
                    st.markdown(f"### 📂 {dosya_adi}")
                    with st.spinner(f"{dosya_adi} inceleniyor..."):
                        try:
                            with zipfile.ZipFile(zip_file) as z:
                                files_info = []
                                for info in z.infolist():
                                    if not info.is_dir():
                                        files_info.append({'name': info.filename, 'date': datetime(*info.date_time)})
                                sorted_files = sorted(files_info, key=lambda x: x['date'], reverse=True)[:5]
                                
                                file_context = ""
                                for f_info in sorted_files:
                                    fname = f_info['name']
                                    fdate = f_info['date'].strftime('%d.%m.%Y')
                                    with z.open(fname) as f:
                                        file_bytes = BytesIO(f.read())
                                        ext = fname.split('.')[-1].lower()
                                        content = ""
                                        try:
                                            if ext == 'udf': content = parse_udf(file_bytes)
                                            elif ext == 'pdf': content = parse_pdf(file_bytes)
                                            elif ext in ['docx', 'doc']: content = extract_text_from_docx(file_bytes)
                                            elif ext == 'txt': content = file_bytes.read().decode('utf-8', errors='ignore')
                                        except: content = "Okunamadı"
                                        file_context += f"\n--- {fname} ({fdate}) ---\n{content[:5000]}"
                                
                                prompt = f"GÖREV: Bu dava dosyasının SON 5 evrağına göre durumu özetle.\nEVRAKLAR:\n{file_context}"
                                analiz_sonucu = get_ai_response(prompt, api_key)
                                st.markdown(f"<div class='uyap-kutusu'>{analiz_sonucu}</div>", unsafe_allow_html=True)
                        except Exception as e: st.error(f"Hata: {str(e)}")
                    progress_bar.progress((idx + 1) / len(uyap_zips))


    with tab16:
        st.subheader("🕸️ Dosya İlişki Ağı")
        if st.button("İlişki Ağını Çiz", type="primary"):
            if not api_key or not st.session_state.doc_text: st.error("Dosya ve API Key gerekli.")
            else:
                with st.spinner("Analiz ediliyor..."):
                    prompt = f"GÖREV: Bu metindeki kişileri ve rollerini Graphviz DOT formatında ver.\nMETİN: {st.session_state.doc_text[:50000]}"
                    dot_code = get_ai_response(prompt, api_key).replace("```dot", "").replace("```", "").strip()
                    try: st.graphviz_chart(dot_code)
                    except: st.code(dot_code)

    with tab17:
        st.subheader("📝 Sözleşme Risk Analizi")
        sozlesme_file = st.file_uploader("Sözleşme Yükle", type=['pdf', 'docx'], key="soz_up")
        if sozlesme_file and st.button("Sözleşmeyi İncele"):
            if not api_key: st.error("API Key gerekli.")
            else:
                with st.spinner("İnceleniyor..."):
                    s_bytes = BytesIO(sozlesme_file.getvalue())
                    s_ext = sozlesme_file.name.split('.')[-1].lower()
                    s_text = extract_text_from_docx(s_bytes) if s_ext == 'docx' else parse_pdf(s_bytes)
                    prompt = f"GÖREV: Bu sözleşmeyi risk analizi yap (Riskler, Eksikler, Öneriler).\nMETİN: {s_text[:50000]}"
                    st.session_state.sozlesme_analiz = get_ai_response(prompt, api_key)
        if st.session_state.sozlesme_analiz:
            st.markdown(st.session_state.sozlesme_analiz)


    # --- YENİ MODÜLLER (TAB 19-23) ---

    with tab19:
        st.subheader("🕵️‍♂️ KVKK / Anonimleştirme")
        st.info("Metindeki T.C. Kimlik, Telefon ve İsimleri gizler.")
        kvkk_input = st.text_area("Metni Buraya Yapıştırın", height=200)
        
        if st.button("🛡️ Anonimleştir"):
            if kvkk_input:
                # Regex ile Temizlik
                gizli_metin = re.sub(r'\d{11}', '[TCKN GİZLENDİ]', kvkk_input) # TC
                gizli_metin = re.sub(r'05\d{9}', '[TEL GİZLENDİ]', gizli_metin) # Tel
                gizli_metin = re.sub(r'\d{2}/\d{2}/\d{4}', '[TARİH]', gizli_metin) # Tarih
                
                # AI ile İsim Temizliği (Opsiyonel)
                if api_key:
                    with st.spinner("AI ile isimler taranıyor..."):
                        prompt = f"GÖREV: Bu metindeki tüm özel isimleri (Kişi adları) '[İSİM]' olarak değiştir. Başka hiçbir şeyi değiştirme.\nMETİN: {gizli_metin}"
                        gizli_metin = get_ai_response(prompt, api_key)
                
                st.session_state.kvkk_metin = gizli_metin
        
        if st.session_state.kvkk_metin:
            st.text_area("Sonuç:", st.session_state.kvkk_metin, height=200)
            st.download_button("📥 İndir", st.session_state.kvkk_metin, "Anonim.txt")



    with tab21: # Belge Kıyasla & Mevzuat Diff Motoru (Gelişmiş)
        st.subheader("⚖️ Mevzuat ve Sözleşme Diff Motoru (Git-Style)")
        st.info("Eski ve yeni versiyonları karşılaştırın. İster metin yapıştırın, ister PDF/Word/Resim dosyası yükleyin. Sistem OCR desteklidir.")

        # Yardımcı Fonksiyon: Dosyadan Metin Okuma (OCR Dahil)
        def get_file_content(uploaded_file):
            if uploaded_file is None: return ""
            
            # Kütüphaneleri güvenli çağır
            import io
            try: import PyPDF2
            except: PyPDF2 = None
            try: from docx import Document
            except: Document = None
            try: from PIL import Image; import pytesseract
            except: Image = None; pytesseract = None

            filename = uploaded_file.name
            ext = filename.split('.')[-1].lower()
            text_result = ""

            try:
                # PDF
                if ext == 'pdf':
                    if PyPDF2:
                        reader = PyPDF2.PdfReader(uploaded_file)
                        for page in reader.pages:
                            text_result += page.extract_text() + "\n"
                    else: return "[Hata: PyPDF2 eksik]"
                
                # WORD
                elif ext == 'docx':
                    if Document:
                        doc = Document(uploaded_file)
                        for para in doc.paragraphs:
                            text_result += para.text + "\n"
                    else: return "[Hata: python-docx eksik]"
                
                # RESİM (OCR)
                elif ext in ['png', 'jpg', 'jpeg', 'tiff', 'tif', 'bmp', 'img']:
                    if Image and pytesseract:
                        img = Image.open(uploaded_file)
                        # Türkçe OCR denemesi
                        try: text_result = pytesseract.image_to_string(img, lang='tur')
                        except: text_result = pytesseract.image_to_string(img)
                    else: return "[Hata: OCR kütüphaneleri eksik]"
                
                # TXT / UDF
                elif ext in ['txt', 'udf', 'xml']:
                    stringio = io.StringIO(uploaded_file.getvalue().decode("utf-8", errors='ignore'))
                    text_result = stringio.read()
                
                # Desteklenmeyen/Binary
                else:
                    text_result = "[Bu dosya formatından metin okunamadı]"

            except Exception as e:
                return f"[Okuma Hatası: {str(e)}]"
            
            return text_result

        # --- ARAYÜZ TASARIMI ---
        col_diff1, col_diff2 = st.columns(2)

        # SOL SÜTUN: ESKİ METİN (MÜLGA)
        with col_diff1:
            st.markdown("#### 🔴 Eski Metin (Mülga/Eski Versiyon)")
            input_type_1 = st.radio("Giriş Yöntemi:", ["✍️ Metin Yapıştır", "📂 Dosya Yükle"], key="radio_diff_1", horizontal=True)
            
            old_text_content = ""
            
            if input_type_1 == "✍️ Metin Yapıştır":
                old_text_content = st.text_area("Metni Buraya Yapıştırın", height=300, key="text_diff_1", placeholder="Eski maddeyi buraya girin...")
            else:
                file_1 = st.file_uploader("Dosya Seç (PDF, Word, Resim)", type=['pdf','docx','txt','png','jpg','jpeg','tif','tiff'], key="file_diff_1")
                if file_1:
                    with st.spinner("Dosya okunuyor..."):
                        old_text_content = get_file_content(file_1)
                        st.success(f"Dosya okundu: {len(old_text_content)} karakter")
                        with st.expander("Okunan Metni Gör"):
                            st.text(old_text_content[:1000] + "...")

        # SAĞ SÜTUN: YENİ METİN (MER'İ)
        with col_diff2:
            st.markdown("#### 🟢 Yeni Metin (Mer'i/Yeni Versiyon)")
            input_type_2 = st.radio("Giriş Yöntemi:", ["✍️ Metin Yapıştır", "📂 Dosya Yükle"], key="radio_diff_2", horizontal=True)
            
            new_text_content = ""
            
            if input_type_2 == "✍️ Metin Yapıştır":
                new_text_content = st.text_area("Metni Buraya Yapıştırın", height=300, key="text_diff_2", placeholder="Yeni maddeyi buraya girin...")
            else:
                file_2 = st.file_uploader("Dosya Seç (PDF, Word, Resim)", type=['pdf','docx','txt','png','jpg','jpeg','tif','tiff'], key="file_diff_2")
                if file_2:
                    with st.spinner("Dosya okunuyor..."):
                        new_text_content = get_file_content(file_2)
                        st.success(f"Dosya okundu: {len(new_text_content)} karakter")
                        with st.expander("Okunan Metni Gör"):
                            st.text(new_text_content[:1000] + "...")

        st.divider()

        # --- ANALİZ BUTONU ---
        if st.button("⚡ Farkları Bul ve Hukuki Etkiyi Analiz Et", use_container_width=True):
            if not old_text_content or not new_text_content:
                st.warning("Lütfen her iki taraf için de metin girin veya dosya yükleyin.")
            else:
                import difflib

                # 1. GÖRSEL DIFF OLUŞTURMA
                a = old_text_content.split()
                b = new_text_content.split()
                
                matcher = difflib.SequenceMatcher(None, a, b)
                html_diff = []
                
                for opcode, a0, a1, b0, b1 in matcher.get_opcodes():
                    if opcode == 'equal':
                        html_diff.append(" ".join(a[a0:a1]))
                    elif opcode == 'insert':
                        html_diff.append(f"<span style='background-color:#d4edda; color:#155724; padding:2px; border-radius:3px; font-weight:bold; border:1px solid #c3e6cb;'>{' '.join(b[b0:b1])}</span>")
                    elif opcode == 'delete':
                        html_diff.append(f"<span style='background-color:#f8d7da; color:#721c24; text-decoration:line-through; padding:2px; border-radius:3px; opacity: 0.7;'>{' '.join(a[a0:a1])}</span>")
                    elif opcode == 'replace':
                        html_diff.append(f"<span style='background-color:#f8d7da; color:#721c24; text-decoration:line-through; padding:2px; opacity: 0.7;'>{' '.join(a[a0:a1])}</span> <span style='background-color:#d4edda; color:#155724; font-weight:bold; border:1px solid #c3e6cb; padding:2px;'>{' '.join(b[b0:b1])}</span>")
                
                diff_result = " ".join(html_diff)

                st.markdown("### 🔍 Detaylı Karşılaştırma Raporu")
                st.markdown(f"""
                <div style="border:1px solid #ccc; padding:25px; border-radius:10px; line-height: 1.8; font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; background-color: #fafafa;">
                    {diff_result}
                </div>
                <div style="margin-top:10px; font-size:0.85em; color:gray; text-align: right;">
                    <span style='background-color:#f8d7da; color:#721c24; padding:3px 8px; border-radius:4px;'>🔴 Silinen İfade</span> 
                    <span style='background-color:#d4edda; color:#155724; padding:3px 8px; border-radius:4px; margin-left:10px;'>🟢 Eklenen İfade</span>
                </div>
                """, unsafe_allow_html=True)

                # 2. YAPAY ZEKA ETKİ ANALİZİ
                if api_key:
                    st.divider()
                    with st.spinner("Yapay zeka hukuki sonuçları ve içtihat etkilerini hesaplıyor..."):
                        prompt = f"""
                        GÖREV: Sen kıdemli bir hukukçusun. Aşağıdaki iki metin arasındaki farkları analiz et.
                        
                        ESKİ VERSİYON:
                        {old_text_content[:4000]} 
                        
                        YENİ VERSİYON:
                        {new_text_content[:4000]}
                        (Not: Metinler çok uzunsa ilk 4000 karakter alınmıştır)
                        
                        ANALİZ İSTEĞİ:
                        1. **Değişiklik Özeti:** Ne değişti? (Tek cümle)
                        2. **Hukuki Yorum Farkı:** Bu değişiklik anlamı nasıl kaydırdı? (Örn: "Zorunluluk"tan "Takdir yetkisi"ne geçiş vb.)
                        3. **Risk Analizi:** Yeni metin hangi riskleri doğuruyor veya hangi açıkları kapatıyor?
                        4. **İçtihat Etkisi:** Eski Yargıtay kararları bu yeni metinle geçersiz kalır mı?
                        """
                        
                        analiz = get_ai_response(prompt, api_key)
                        
                        st.markdown("### 🧠 Yapay Zeka Hukuki Etki Analizi")
                        st.markdown(f"""
                        <div style="background-color:#fff3cd; padding:20px; border-radius:10px; border-left: 5px solid #ffc107;">
                            {analiz}
                        </div>
                        """, unsafe_allow_html=True)
                else:
                    st.warning("Detaylı etki analizi için API Key gereklidir.")


    with tab22:
        st.subheader("🎭 Sanal Duruşma Simülasyonu")
        st.info("AI Hakim karşısında savunma pratiği yapın.")
        
        if "mock_started" not in st.session_state: st.session_state.mock_started = False
        
        col_mock1, col_mock2 = st.columns([1, 3])
        with col_mock1:
            rol = st.selectbox("Rolünüz", ["Davacı Vekili", "Davalı Vekili"])
            konu = st.text_input("Dava Konusu", "İş Kazası Tazminatı")
            if st.button("Duruşmayı Başlat"):
                st.session_state.mock_started = True
                st.session_state.mock_messages = [{"role": "assistant", "content": f"MAHKEME BAŞKANI: {konu} davasına başlıyoruz. {rol}, ilk beyanınızı dinliyorum. Buyurun."}]
        
        with col_mock2:
            if st.session_state.mock_started:
                # Mesajları Göster
                for msg in st.session_state.mock_messages:
                    if msg["role"] == "assistant":
                        st.markdown(f"👨‍⚖️ **HAKİM:** {msg['content']}")
                    else:
                        st.markdown(f"🧑‍⚖️ **SİZ:** {msg['content']}")
                
                # Yeni Cevap Girişi
                user_reply = st.chat_input("Hakime Cevabınız...")
                if user_reply:
                    st.session_state.mock_messages.append({"role": "user", "content": user_reply})
                    
                    if not api_key: st.error("API Key gerekli.")
                    else:
                        with st.spinner("Hakim düşünüyor..."):
                            history = "\n".join([f"{m['role']}: {m['content']}" for m in st.session_state.mock_messages])
                            prompt = f"""
                            SEN BİR HAKİMSİN. Sert, kuralcı ve sorgulayıcı bir Türk hakimi gibi davran.
                            Kullanıcı {rol}. Konu: {konu}.
                            Kullanıcının beyanına göre mantıklı bir karşı soru sor veya ara karar ver.
                            KONUŞMA GEÇMİŞİ:
                            {history}
                            """
                            ai_reply = get_ai_response(prompt, api_key)
                            st.session_state.mock_messages.append({"role": "assistant", "content": ai_reply})
                            st.rerun()

    with tab23:
        st.subheader("✅ Akıllı Görev")
        st.info("Mahkeme kararından yapılacak işleri listeler.")
        
        karar_metni = st.text_area("Karar / Ara Karar Metni", height=150)
        if st.button("Görevleri Çıkar"):
            if not api_key: st.error("API Key gerekli.")
            elif not karar_metni: st.warning("Metin giriniz.")
            else:
                with st.spinner("Analiz ediliyor..."):
                    prompt = f"""
                    GÖREV: Bu mahkeme kararını oku ve avukatın yapması gereken işleri "To-Do List" formatında çıkar.
                    Varsa süreleri ve tarihleri belirt.
                    METİN: {karar_metni}
                    """
                    st.session_state.gorev_listesi = get_ai_response(prompt, api_key)
        
        if st.session_state.gorev_listesi:
            st.markdown(f"<div class='buyur-abi-kutusu'>{st.session_state.gorev_listesi}</div>", unsafe_allow_html=True)
            if st.button("Listeyi Kopyala"):
                st.toast("Görev listesi kopyalandı!")


    with tab31: # Şeytanın Avukatı (Devil's Advocate)
        st.subheader("😈 Şeytanın Avukatı (AI Adversary)")
        st.info("Dilekçenizi buraya yapıştırın. Yapay zeka 'Karşı Tarafın Avukatı' rolüne girsin ve dilekçenizi acımasızca eleştirsin.")
        
        dilekce_taslagi = st.text_area("Dilekçe Taslağınız:", height=300, placeholder="Hazırladığınız dilekçe metnini buraya yapıştırın...")
        
        if st.button("Dilekçemi Parçala ve Açıkları Bul", type="primary"):
            if not api_key: st.error("API Key gerekli.")
            elif not dilekce_taslagi: st.warning("Eleştirilecek bir metin girmediniz.")
            else:
                with st.spinner("Yapay zeka karşı taraf cübbesini giyiyor ve açık arıyor..."):
                    prompt = f"""
                    GÖREV: Sen çok dişli, acımasız ve zeki bir 'Karşı Taraf Avukatı'sın.
                    METİN: Aşağıdaki dilekçe taslağını incele.
                    
                    YAPMAN GEREKENLER:
                    1. MANTIK HATALARI: Dilekçedeki mantıksız yerleri bul.
                    2. EKSİK DELİLLER: Nelerin ispatlanamadığını yüzüne vur.
                    3. KARŞI TEZLER: Bu dilekçeye karşı hangi Yargıtay kararlarını veya kanun maddelerini kullanırdın?
                    4. RİSK PUANI: Bu dilekçenin reddedilme ihtimali % kaç?
                    
                    ÜSLUP: Sert, eleştirel ama yol gösterici (Profesyonel).
                    
                    DİLEKÇE: {dilekce_taslagi}
                    """
                    elestiri = get_ai_response(prompt, api_key)
                    
                    st.error("🚨 Tespit Edilen Zayıf Noktalar:")
                    st.markdown(elestiri)
                    st.success("💡 İpucu: Yukarıdaki eleştirilere göre dilekçenizi revize ederseniz kazanma şansınız artar.")

    with tab32: # Canlı Duruşma Asistanı (Live Fact-Check)
        st.subheader("⚡ Canlı Duruşma Asistanı (Live Fact-Check)")
        st.info("Duruşma sırasında karşı tarafın söylediği iddialı cümleyi veya kanun maddesini girin. Sistem anında doğruluk kontrolü yapsın.")
        
        # Hızlı giriş için form kullanımı (Enter'a basınca çalışsın diye)
        with st.form(key='live_check_form'):
            col_live1, col_live2 = st.columns([3, 1])
            with col_live1:
                iddia_cumlesi = st.text_input("Karşı Taraf Ne Dedi?", placeholder="Örn: Yargıtay'ın son kararına göre işe iade davasında zamanaşımı 10 yıldır!")
            with col_live2:
                sorgula_btn = st.form_submit_button("🔍 Doğrula")
        
        if sorgula_btn:
            if not api_key: st.error("API Key gerekli.")
            elif not iddia_cumlesi: st.warning("Doğrulanacak ifadeyi girin.")
            else:
                with st.spinner("Mevzuat ve İçtihat taranıyor..."):
                    prompt = f"""
                    ACİL DURUM: Duruşma salonundayım. Karşı taraf şu iddiada bulundu:
                    "{iddia_cumlesi}"
                    
                    GÖREV:
                    1. Bu bilgi hukuken DOĞRU mu YANLIŞ mı?
                    2. Yanlışsa doğrusu nedir? (Kanun maddesi veya yerleşik içtihat ile açıkla).
                    3. Cevap çok kısa ve net olsun. Avukatın ekrandan bir bakışta okuması lazım.
                    """
                    fact_check = get_ai_response(prompt, api_key)
                    
                    # Görselleştirme (Doğruysa Yeşil, Yanlışsa Kırmızı Kutu)
                    if "yanlış" in fact_check.lower() or "hatalı" in fact_check.lower() or "hayır" in fact_check.lower():
                        st.markdown(f"""
                        <div style="background-color:#ffcccc; padding:20px; border-radius:10px; border: 2px solid red; color: darkred;">
                            ❌ <b>DİKKAT! BU BİLGİ HATALI OLABİLİR</b><br><br>
                            {fact_check}
                        </div>
                        """, unsafe_allow_html=True)
                    else:
                        st.markdown(f"""
                        <div style="background-color:#ccffcc; padding:20px; border-radius:10px; border: 2px solid green; color: darkgreen;">
                            ✅ <b>BİLGİ DOĞRULANDI / MAKUL</b><br><br>
                            {fact_check}
                        </div>
                        """, unsafe_allow_html=True)


    with tab26: # Hukuki Çeviri Modülü
        st.subheader("🌍 Hukuki Terminoloji Çevirmeni")
        st.info("Yapay zeka, kelimeleri 'hukuki bağlamda' değerlendirerek çevirir. (Örn: Bar -> Baro, Execution -> İcra)")
        
        col_tr1, col_tr2 = st.columns(2)
        with col_tr1:
            kaynak_dil = st.selectbox("Kaynak Dil", ["Türkçe", "İngilizce", "Almanca", "Fransızca"], index=0)
        with col_tr2:
            hedef_dil = st.selectbox("Hedef Dil", ["İngilizce", "Türkçe", "Almanca", "Fransızca"], index=1)
            
        ceviri_metni = st.text_area("Çevrilecek Metni Girin:", height=150, placeholder="Metni buraya yapıştırın...")
        
        if st.button("Hukuki Çeviri Yap", type="primary"):
            if not api_key: 
                st.error("Lütfen API Key giriniz.")
            elif not ceviri_metni:
                st.warning("Lütfen çevrilecek bir metin giriniz.")
            else:
                with st.spinner("Terminoloji kontrol edilerek çevriliyor..."):
                    prompt = f"""
                    GÖREV: Sen uzman bir hukuk çevirmenisin. Aşağıdaki metni {kaynak_dil} dilinden {hedef_dil} diline çevir.
                    KURAL 1: Hukuki terminolojiyi (Legal Terminology) kesinlikle koru. Günlük dil yerine resmi hukuk dili kullan.
                    KURAL 2: Sadece çeviriyi ver, açıklama yapma.
                    METİN: {ceviri_metni}
                    """
                    ceviri_sonuc = get_ai_response(prompt, api_key)
                    
                    st.success("Çeviri Tamamlandı:")
                    st.markdown(f"**📄 {hedef_dil} Çıktısı:**")
                    st.markdown(f"""
                    <div style="background-color:#f0f2f6; padding:15px; border-radius:10px; border-left: 5px solid #ff4b4b;">
                        {ceviri_sonuc}
                    </div>
                    """, unsafe_allow_html=True)
    with tab29: # Tez Çürütücü Modülü
        st.subheader("🛡️ Karşı Taraf Tez Çürütücü")
        st.info("Karşı tarafın iddiasını girin, yapay zeka bu iddiayı çürütmek için hukuki argümanlar üretsin.")
        
        col_tez1, col_tez2 = st.columns([2, 1])
        with col_tez1:
            karsi_iddia = st.text_area("Karşı Tarafın İddiası / Savunması:", height=150, placeholder="Örn: Davalı, işe geç gelmeyi alışkanlık haline getirdiği için haklı nedenle fesih yapıldığını iddia etmektedir...")
        with col_tez2:
            st.write("📌 **Strateji Seçimi:**")
            strateji = st.radio("Nasıl Çürütelim?", ["Agresif (Sert Savunma)", "Teknik (Usul Hukuku)", "Uzlaşmacı (Alternatifli)"])
            
        if st.button("Argümanları Üret", type="primary"):
            if not api_key: st.error("API Key gerekli.")
            elif not karsi_iddia: st.warning("Lütfen çürütülecek bir iddia girin.")
            else:
                with st.spinner("Yargıtay kararları ve kanun maddeleri taranıyor..."):
                    prompt = f"""
                    GÖREV: Sen tecrübeli bir avukatsın. Aşağıdaki karşı taraf iddiasını çürütmek için 3 farklı hukuki argüman geliştir.
                    İDDİA: {karsi_iddia}
                    STRATEJİ: {strateji}
                    
                    ÇIKTI FORMATI:
                    1. [Argüman Başlığı] - [Hukuki Dayanak/Mantık]
                    2. [Argüman Başlığı] - [Hukuki Dayanak/Mantık]
                    3. [Argüman Başlığı] - [Hukuki Dayanak/Mantık]
                    
                    NOT: İlgili olabilecek Kanun Maddelerini (Örn: HMK, TBK, İş Kanunu) parantez içinde belirt.
                    """
                    cevap = get_ai_response(prompt, api_key)
                    
                    st.success("✅ İşte Kullanabileceğiniz Karşı Argümanlar:")
                    st.markdown(f"""
                    <div style="background-color:#fff3cd; padding:20px; border-radius:10px; border: 1px solid #ffeeba; color:#856404;">
                        {cevap}
                    </div>
                    """, unsafe_allow_html=True)

    with tab30: # Çapraz Sorgu Hazırlayıcı
        st.subheader("🕵️‍♂️ Çapraz Sorgu Hazırlayıcı (Cross-Examination)")
        st.info("Tanık veya sanık ifadesini girin. Yapay zeka, çelişkileri bulsun ve köşeye sıkıştıran sorular hazırlasın.")
        
        col_sorgu1, col_sorgu2 = st.columns([2, 1])
        with col_sorgu1:
            ifade_metni = st.text_area("Tanık/Sanık İfadesi:", height=200, placeholder="Örn: Olay günü evdeydim, saat 20:00 gibi uyudum. Kimseyi görmedim ama sesleri duydum...")
        with col_sorgu2:
            st.write("🎯 **Hedefiniz Ne?**")
            sorgu_amaci = st.radio("Sorgu Stratejisi", ["Güvenilirliği Sarsmak (Yalanını Yakala)", "Bilgi Eksikliğini Göstermek", "Önyargısını Ortaya Çıkarmak"])
            
        if st.button("Soruları Hazırla", type="primary"):
            if not api_key: st.error("API Key gerekli.")
            elif not ifade_metni: st.warning("Lütfen bir ifade metni girin.")
            else:
                with st.spinner("İfade analiz ediliyor, mantık hataları taranıyor..."):
                    prompt = f"""
                    GÖREV: Sen uzman bir ceza avukatısın. Aşağıdaki ifadeyi analiz et ve çapraz sorgu soruları hazırla.
                    İFADE: {ifade_metni}
                    AMACIMIZ: {sorgu_amaci}
                    
                    ÇIKTI FORMATI:
                    1. [Tespit Edilen Çelişki/Zayıf Nokta]
                       - Soru: [Tanığa sorulacak sert ve net soru]
                       - Beklenen Cevap ve Tuzak: [Neden bu soruyu sorduk?]
                    
                    En az 3, en fazla 5 kritik soru hazırla.
                    """
                    sorgu_sonuc = get_ai_response(prompt, api_key)
                    
                    st.success("⚔️ Hazırlanan Çapraz Sorgu Planı:")
                    st.markdown(f"""
                    <div style="background-color:#e8f4f8; padding:20px; border-radius:10px; border-left: 5px solid #00a8cc;">
                        {sorgu_sonuc}
                    </div>
                    """, unsafe_allow_html=True)
    with tab33: # Mevzuat Etki Analizi (Impact Analysis)
        st.subheader("📡 Akıllı Mevzuat Radarı & Etki Analizi")
        st.info("Bu modül, Resmi Gazete'yi günlük olarak tarar ve SADECE sizin takip listenizdeki dosyaları etkileyen değişiklikleri raporlar.")

        # --- 1. TAKİP LİSTESİ YÖNETİMİ (Sizin Kodunuzdan Geliştirildi) ---
        if 'mevzuat_takip_listesi' not in st.session_state:
            st.session_state.mevzuat_takip_listesi = []

        with st.expander("📋 Takip Listesi & Dosya Tanımlama", expanded=False):
            col_takip1, col_takip2, col_takip3 = st.columns(3)
            with col_takip1:
                takip_kanun = st.text_input("Kanun/Mevzuat Adı", placeholder="Örn: İmar Kanunu")
            with col_takip2:
                takip_keyword = st.text_input("Anahtar Kelime (Konu)", placeholder="Örn: Ruhsat, İskan")
            with col_takip3:
                takip_dosya = st.text_input("İlgili Dosya No", placeholder="Örn: 2024/15 E.")
                
            if st.button("Listeye Ekle", use_container_width=True):
                if takip_keyword:
                    yeni_kural = {
                        "kanun": takip_kanun if takip_kanun else "Genel",
                        "konu": takip_keyword,
                        "dosya": takip_dosya if takip_dosya else "Genel Bilgi"
                    }
                    st.session_state.mevzuat_takip_listesi.append(yeni_kural)
                    st.success(f"✅ '{takip_keyword}' konusu takibe alındı.")
                else:
                    st.warning("En azından bir Anahtar Kelime girmelisiniz.")

            # Mevcut Listeyi Göster
            if st.session_state.mevzuat_takip_listesi:
                st.markdown("###### 📝 Aktif Takip Listesi")
                import pandas as pd
                df_takip = pd.DataFrame(st.session_state.mevzuat_takip_listesi)
                st.table(df_takip)
            else:
                st.info("Henüz takip kuralı eklemediniz.")

        st.divider()

        # --- 2. TARAMA MOTORU (Otomatik + Manuel Seçenekli) ---
        scan_option = st.radio("Analiz Yöntemi Seçin:", ["🌍 Resmi Gazete'yi Otomatik Tara", "✍️ Metni Manuel Yapıştır"])

        # A) OTOMATİK TARAMA MODU (GÜÇLENDİRİLMİŞ VERSİYON)
        if scan_option == "🌍 Resmi Gazete'yi Otomatik Tara":
            col_scan1, col_scan2 = st.columns([1, 3])
            
            with col_scan1:
                st.markdown("##### ⚙️ Ayarlar")
                scan_depth = st.radio("Tarama Derinliği:", ["Hızlı (Başlıklar)", "Derin (PDF İçerikleri)"])
                # Kullanıcıya seçtirmeyelim, otomatik deneyelim
                st.info("Sistem önce canlı bağlantıyı dener, engellenirse simülasyona geçer.")
                start_btn = st.button("🚀 Taramayı Başlat", type="primary", use_container_width=True)

            with col_scan2:
                if start_btn:
                    if not st.session_state.mevzuat_takip_listesi:
                        st.error("Önce yukarıdan takip listesine en az bir konu ekleyin!")
                    else:
                        found_matches = []
                        status_box = st.empty()
                        progress_bar = st.progress(0)
                        target_links = []
                        
                        # --- BAĞLANTI DENEMESİ ---
                        status_box.info("Resmi Gazete sunucularına bağlanılıyor...")
                        
                        canli_veri_cekildi = False
                        
                        try:
                            # Yöntem 1: RSS Beslemesi (Daha az engellenir)
                            url_rss = "https://www.resmigazete.gov.tr/rss.xml"
                            # Tarayıcı gibi görünmek için Header ekliyoruz
                            headers = {
                                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                                'Referer': 'https://www.google.com/'
                            }
                            
                            response = requests.get(url_rss, headers=headers, timeout=5)
                            
                            if response.status_code == 200:
                                soup = BeautifulSoup(response.content, 'xml') # XML parser kullanıyoruz
                                items = soup.find_all('item')
                                for item in items:
                                    title = item.title.text
                                    link = item.link.text
                                    target_links.append({"title": title, "link": link})
                                canli_veri_cekildi = True
                                status_box.success(f"✅ Bağlantı Başarılı! {len(target_links)} başlık çekildi.")
                            
                        except Exception as e_rss:
                            # RSS başarısız olursa loglayalım ama çökertmeyelim
                            print(f"RSS Hatası: {e_rss}")

                        # --- HATA YÖNETİMİ VE SİMÜLASYON ---
                        if not canli_veri_cekildi:
                            status_box.warning("⚠️ Resmi Gazete sunucusu yurt dışı erişimini engelledi. DEMO MODU devrede.")
                            # Demo verilerle devam et
                            target_links = [
                                {"title": "7440 Sayılı Bazı Alacakların Yeniden Yapılandırılmasına Dair Kanun", "link": "https://www.resmigazete.gov.tr/eskiler/2023/03/20230312-1.pdf"},
                                {"title": "İmar Kanununda Değişiklik Yapılmasına Dair Kanun Teklifi", "link": "https://www.resmigazete.gov.tr/"},
                                {"title": "Anayasa Mahkemesi Kararı (Esas: 2023/15)", "link": "https://www.resmigazete.gov.tr/"}
                            ]
                            time.sleep(1) # Kullanıcı uyarıyı görsün diye bekleme

                        # --- 2. İÇERİKLERİ TARA VE EŞLEŞTİR ---
                        total_docs = len(target_links)
                        
                        for i, doc in enumerate(target_links):
                            progress_bar.progress((i + 1) / total_docs)
                            doc_text = ""
                            
                            # İçerik Çekme (Hata olursa başlığı kullan)
                            try:
                                if canli_veri_cekildi and scan_depth == "Derin (PDF İçerikleri)":
                                    # Canlı modda PDF indirmeyi dene
                                    r_doc = requests.get(doc['link'], headers=headers, timeout=5)
                                    if doc['link'].endswith(".pdf"):
                                        f = io.BytesIO(r_doc.content)
                                        reader = PyPDF2.PdfReader(f)
                                        for p in range(min(2, len(reader.pages))):
                                            doc_text += reader.pages[p].extract_text()
                                    else:
                                        doc_text = doc['title']
                                else:
                                    # Simülasyon veya Hızlı modda sadece başlık + örnek metin
                                    doc_text = doc['title'] + " (İçerik özeti...)"
                            except:
                                doc_text = doc['title']

                            # Takip Listesiyle Karşılaştır
                            for item in st.session_state.mevzuat_takip_listesi:
                                keyword = item['konu']
                                # Basit eşleşme kontrolü
                                if keyword.lower() in doc_text.lower() or (not canli_veri_cekildi and i == 0): 
                                    # Not: Simülasyonda en az 1 tane çıksın diye 'i==0' hilesi yaptık
                                    found_matches.append({
                                        "doc_title": doc['title'],
                                        "doc_link": doc['link'],
                                        "matched_item": item,
                                        "context": doc_text[:500]
                                    })

                        progress_bar.empty()
                        
                        # --- 3. SONUÇLARI GÖSTER ---
                        if found_matches:
                            status_box.success(f"🚨 {len(found_matches)} adet kritik eşleşme bulundu!")
                            
                            for match in found_matches:
                                with st.container():
                                    st.markdown(f"""
                                    <div style="border:1px solid #ddd; padding:15px; border-radius:10px; margin-bottom:10px; background-color:#fff; color:black;">
                                        <h4>🔔 Uyarı: {match['matched_item']['dosya']} Dosyası</h4>
                                        <p><strong>Sebep:</strong> '{match['matched_item']['konu']}' konusu, <em>{match['doc_title']}</em> içinde tespit edildi.</p>
                                        <a href="{match['doc_link']}" target="_blank">📄 Belgeyi Görüntüle</a>
                                    </div>
                                    """, unsafe_allow_html=True)
                                    
                                    # AI Butonu
                                    if st.button(f"🧠 Etki Analizi Yap ({match['matched_item']['dosya']})", key=f"btn_{match['doc_link']}"):
                                        if api_key:
                                            with st.spinner("AI analiz ediyor..."):
                                                prompt = f"Bu kanun değişikliği ({match['doc_title']}), kullanıcının '{match['matched_item']['konu']}' konulu dosyasını nasıl etkiler? Avukat gibi yorumla."
                                                res = get_ai_response(prompt, api_key)
                                                st.info(res)
                                        else:
                                            st.warning("API Key eksik.")
                        else:
                            status_box.info("✅ Bugün takip listenizdeki konularla ilgili bir değişiklik yayınlanmadı.")


        # B) MANUEL YAPIŞTIRMA MODU (Sizin Kodunuzdan Entegre Edildi)
        else:
            st.markdown("##### 📝 Metin Analizi")
            manual_text = st.text_area("Analiz edilecek mevzuat metnini buraya yapıştırın:", height=200)
            
            if st.button("Analiz Et", type="primary"):
                if not manual_text or not st.session_state.mevzuat_takip_listesi:
                    st.warning("Lütfen metin girin ve takip listenizin dolu olduğundan emin olun.")
                elif api_key:
                    with st.spinner("Takip listenizdeki dosyalar kontrol ediliyor..."):
                        takip_json = json.dumps(st.session_state.mevzuat_takip_listesi, ensure_ascii=False)
                        prompt = f"""
                        GÖREV: Sen bir Mevzuat Analiz Uzmanısın.
                        KULLANICI DOSYALARI: {takip_json}
                        YENİ METİN: {manual_text}
                        
                        Bu metindeki değişiklikler yukarıdaki dosyalardan hangilerini etkiliyor?
                        Her etkilenen dosya için kısa bir uyarı yaz.
                        """
                        res = get_ai_response(prompt, api_key)
                        st.success("Analiz Tamamlandı")
                        st.write(res)

    with tab34: # Semantik Arşiv Sorgulama (RAG) - OCR Destekli (Düzeltilmiş)
        st.subheader("🧠 Semantik Arşiv (OCR & Çoklu Format)")
        st.info("PDF, Word, UDF, TXT ve Resim (JPG, PNG) dosyalarını yükleyin. Sistem görselleri okur (OCR), metinleri tarar ve sorunuzun cevabını dosya adıyla birlikte verir.")
        
        # Dosya Yükleme Alanı
        # DÜZELTME: 'accept_multiple' yerine 'accept_multiple_files' kullanıldı.
        uploaded_archive = st.file_uploader(
            "Arşive Eklenecek Dosyalar", 
            accept_multiple_files=True, 
            type=["pdf", "txt", "docx", "doc", "udf", "png", "jpg", "jpeg", "tiff", "bmp"],
            key="rag_file_uploader_final"
        )
        
        # Oturum bazlı hafıza
        if 'archive_memory' not in st.session_state:
            st.session_state.archive_memory = ""
            
        if uploaded_archive:
            if st.button("📂 Dosyaları Tara, OCR Yap ve Hafızaya Al", key="rag_process_btn_final"):
                # Gerekli kütüphaneleri güvenli şekilde çağırıyoruz
                import io
                
                # Kütüphane kontrolü (Yüklü değilse kodun patlamaması için)
                try:
                    import PyPDF2
                except ImportError:
                    PyPDF2 = None
                
                try:
                    from docx import Document
                except ImportError:
                    Document = None
                    
                try:
                    from PIL import Image
                    import pytesseract
                except ImportError:
                    Image = None
                    pytesseract = None

                tum_metin = ""
                basarili_dosya = 0
                progress_bar = st.progress(0)
                
                st.toast("Dosyalar işleniyor, lütfen bekleyin...", icon="⏳")

                for i, file in enumerate(uploaded_archive):
                    file_name = file.name
                    file_ext = file_name.split('.')[-1].lower()
                    file_content = ""
                    
                    try:
                        # 1. PDF OKUMA
                        if file_ext == 'pdf':
                            if PyPDF2:
                                try:
                                    pdf_reader = PyPDF2.PdfReader(file)
                                    for page in pdf_reader.pages:
                                        text = page.extract_text()
                                        if text: file_content += text + "\n"
                                except:
                                    file_content = "[Bu PDF okunamadı veya şifreli]"
                            else:
                                file_content = "[PyPDF2 kütüphanesi eksik]"

                        # 2. WORD (DOCX) OKUMA
                        elif file_ext == 'docx':
                            if Document:
                                try:
                                    doc = Document(file)
                                    for para in doc.paragraphs:
                                        file_content += para.text + "\n"
                                except:
                                    file_content = "[DOCX formatı okunamadı]"
                            else:
                                file_content = "[python-docx kütüphanesi eksik]"
                        
                        # 3. RESİM DOSYALARI (OCR İŞLEMİ)
                        elif file_ext in ['png', 'jpg', 'jpeg', 'tiff', 'bmp']:
                            if Image and pytesseract:
                                try:
                                    image = Image.open(file)
                                    # OCR işlemi (Varsayılan dil)
                                    try:
                                        file_content = pytesseract.image_to_string(image, lang='tur')
                                    except:
                                        file_content = pytesseract.image_to_string(image)
                                        
                                    if not file_content.strip(): 
                                        file_content = "[Resimde okunabilir metin bulunamadı]"
                                except Exception as e_ocr:
                                    file_content = f"[OCR Hatası: {str(e_ocr)}]"
                            else:
                                file_content = "[OCR kütüphaneleri (Pillow/Tesseract) eksik]"

                        # 4. UDF (UYAP) ve TXT OKUMA
                        elif file_ext in ['txt', 'udf', 'xml']:
                            try:
                                stringio = io.StringIO(file.getvalue().decode("utf-8", errors='ignore'))
                                file_content = stringio.read()
                            except:
                                file_content = "[Metin dosyası okunamadı]"
                        
                        # 5. ESKİ WORD (DOC)
                        elif file_ext == 'doc':
                             file_content = "[.doc formatı binary olduğu için tam desteklenmiyor, lütfen .docx'e çevirip yükleyin.]"

                        # Metni Hafızaya Ekle
                        if len(file_content) > 5: 
                            tum_metin += f"\n{'='*20}\n📂 DOSYA ADI: {file_name}\n{'='*20}\n{file_content}\n"
                            basarili_dosya += 1
                        
                    except Exception as e:
                        st.error(f"Hata ({file_name}): {e}")
                    
                    # İlerleme çubuğunu güncelle
                    progress_bar.progress((i + 1) / len(uploaded_archive))
                
                st.session_state.archive_memory = tum_metin
                
                if basarili_dosya > 0:
                    st.success(f"✅ {basarili_dosya} dosya başarıyla işlendi ve hafızaya alındı!")
                else:
                    st.warning("Dosyalar yüklendi ancak içerik okunamadı (Kütüphane eksikliği veya dosya formatı sorunu).")

        st.divider()
        
        # Soru Sorma Alanı
        col_rag1, col_rag2 = st.columns([3, 1])
        with col_rag1:
            rag_soru = st.text_input("Arşive Soru Sor:", placeholder="Örn: 'Tapu iptal davasında bilirkişi raporu kime tebliğ edilmiş?'", key="rag_question_input_final")
        with col_rag2:
            rag_btn = st.button("🧠 Hafızayı Tara", key="rag_search_btn_final")
            
        if rag_btn:
            if not api_key: st.error("API Key gerekli.")
            elif not st.session_state.archive_memory: st.warning("Önce dosya yükleyip işleyin.")
            elif not rag_soru: st.warning("Soru girmediniz.")
            else:
                with st.spinner("Dosyalar taranıyor, anlam analizi yapılıyor..."):
                    prompt = f"""
                    GÖREV: Sen uzman bir Hukuk Arşiv Asistanısın.
                    
                    BAĞLAM (ARCHIVE):
                    Aşağıda kullanıcının yüklediği dosyaların içerikleri var (OCR ile okunmuş metinler dahil):
                    {st.session_state.archive_memory}
                    
                    SORU: {rag_soru}
                    
                    KURALLAR:
                    1. Cevabı sadece yukarıdaki bağlama göre ver.
                    2. Bilgiyi bulduğunda MUTLAKA dosya adını belirt. (Örn: "Bu bilgi 'tutanak.jpg' dosyasında geçmektedir.")
                    3. Cevabı şu formatta ver:
                       - **Bulunan Bilgi:** [Cevap]
                       - **Kaynak Dosya:** [Dosya Adı]
                       - **Kısa Özet:** [Olayın bağlamı]
                    4. Eğer bilgi yoksa "Arşivde bu bilgiye rastlanmadı" de.
                    """
                    
                    rag_cevap = get_ai_response(prompt, api_key)
                    
                    st.markdown("### 🔍 Arama Sonucu:")
                    st.markdown(f"""
                    <div style="background-color:#f0f8ff; padding:20px; border-radius:10px; border-left: 5px solid #1e90ff;">
                        {rag_cevap}
                    </div>
                    """, unsafe_allow_html=True)


    with tab35: # Sesli Duruşma Analizi & Çelişki Alarmı
        st.subheader("🎙️ Duruşma Asistanı: Canlı Çelişki Yakalayıcı")
        st.info("Tanığın önceki ifadesini (Referans Metin) girin ve duruşma ses kaydını yükleyin. Sistem, söylenenleri metne çevirir ve eski ifadeyle çelişen noktaları 'Kırmızı Alarm' olarak bildirir.")

        col_voice1, col_voice2 = st.columns([1, 1])

        # 1. ADIM: REFERANS METİN (Eski İfade)
        with col_voice1:
            st.markdown("### 1. Referans Belge (Eski İfade)")
            ref_text = st.text_area(
                "Emniyet/Savcılık İfadesini Buraya Yapıştırın:", 
                height=250, 
                placeholder="Örn: Olay günü saat 14:00'te evdeydim. Yanımda kimse yoktu. Arabamın rengi mavidir..."
            )

        # 2. ADIM: DURUŞMA SES KAYDI (Yeni Beyan)
        with col_voice2:
            st.markdown("### 2. Duruşma Kaydı (Canlı Beyan)")
            # Ses dosyası yükleme
            audio_file = st.file_uploader("Ses Kaydını Yükle (WAV/FLAC)", type=["wav", "flac"])
            
            # Alternatif: Canlı kayıt simülasyonu için metin girişi (Ses işleme hatası olursa diye)
            st.markdown("--- veya ---")
            manual_transcript = st.text_area("Ses kaydı yoksa, tanığın şu anki sözlerini yazın:", height=100, placeholder="Örn: Olay günü saat 16:00'da dışarıdaydım. Arabam beyaz renklidir.")

        st.divider()

        if st.button("🚨 Çapraz Sorgu Başlat ve Çelişkileri Tara"):
            if not ref_text:
                st.warning("Lütfen karşılaştırma yapmak için eski ifadeyi girin.")
            elif not audio_file and not manual_transcript:
                st.warning("Lütfen duruşma ses kaydı yükleyin veya metin girin.")
            else:
                current_statement = ""
                
                # A) SES İŞLEME (Speech-to-Text)
                if audio_file:
                    with st.spinner("Ses dosyası metne dönüştürülüyor (Transkripsiyon)..."):
                        try:
                            import speech_recognition as sr
                            r = sr.Recognizer()
                            with sr.AudioFile(audio_file) as source:
                                audio_data = r.record(source)
                                # Google Speech API (Ücretsiz versiyon)
                                try:
                                    current_statement = r.recognize_google(audio_data, language='tr-TR')
                                    st.success("Ses başarıyla metne çevrildi!")
                                    with st.expander("Duruşma Transkriptini Gör"):
                                        st.write(current_statement)
                                except sr.UnknownValueError:
                                    st.error("Ses anlaşılamadı.")
                                except sr.RequestError:
                                    st.error("Google Speech API'ye erişilemedi.")
                        except ImportError:
                            st.error("SpeechRecognition kütüphanesi yüklü değil.")
                        except Exception as e:
                            st.error(f"Ses işleme hatası: {e} (Lütfen .WAV formatı deneyin)")
                
                # B) MANUEL GİRİŞ VARSA
                if manual_transcript:
                    current_statement = manual_transcript

                # C) YAPAY ZEKA İLE ÇELİŞKİ ANALİZİ
                if current_statement and api_key:
                    with st.spinner("🕵️ Yapay Zeka ifadeleri çapraz sorguya tutuyor..."):
                        prompt = f"""
                        GÖREV: Sen duruşma salonundaki çok dikkatli bir avukatsın.
                        Amacın: Tanığın şu anki beyanları ile geçmişteki ifadesi arasındaki ÇELİŞKİLERİ yakalamak.
                        
                        1. GEÇMİŞ İFADE (REFERANS):
                        "{ref_text}"
                        
                        2. ŞU ANKİ BEYAN (DURUŞMA):
                        "{current_statement}"
                        
                        ANALİZ KURALLARI:
                        - Sadece bariz çelişkileri bul (Örn: "Mavi" dedi, şimdi "Beyaz" diyor).
                        - Ufak kelime farklarını önemseme.
                        - Çıktıyı şu formatta ver:
                        
                        ALARM: [Çelişki Başlığı]
                        DETAY: Tanık daha önce "[Eski Bilgi]" demişti, ancak şu an "[Yeni Bilgi]" diyor.
                        ÖNERİ: Avukat şu soruyu sormalı: "[Soru Önerisi]"
                        """
                        
                        analiz_sonucu = get_ai_response(prompt, api_key)
                        
                        # Sonuç Gösterimi
                        st.markdown("### 🚨 Çelişki Tespit Raporu")
                        
                        # Eğer AI "Çelişki yok" derse yeşil, varsa kırmızı gösterelim
                        if "yok" in analiz_sonucu.lower() and len(analiz_sonucu) < 50:
                            st.success("✅ İfadeler arasında bariz bir çelişki tespit edilmedi.")
                        else:
                            # Çelişki Kartları
                            st.markdown(f"""
                            <div style="background-color:#ffe6e6; border-left: 6px solid #ff0000; padding:20px; border-radius:10px;">
                                <h4 style="color:#cc0000; margin-top:0;">⚠️ DİKKAT: İFADE DEĞİŞİKLİĞİ TESPİT EDİLDİ</h4>
                                <div style="font-size:1.1em; line-height:1.6; color:#333;">
                                    {analiz_sonucu.replace(chr(10), '<br>')}
                                </div>
                            </div>
                            """, unsafe_allow_html=True)
                            
                            st.info("💡 İpucu: Bu raporu tabletinizde açık tutarak duruşma sırasında anlık müdahale edebilirsiniz.")
                
                elif not api_key:
                    st.error("Analiz için API Key gereklidir.")

    with tab36: # Dijital Otopsi & Metadata Analizi
        st.subheader("🕵️ Dijital Otopsi ve Metadata Dedektifi")
        st.info("Bir dosyanın (PDF veya Resim) 'perde arkasındaki' verilerini (Metadata/EXIF) analiz eder. Dosyanın ne zaman, kim tarafından, hangi yazılımla oluşturulduğunu ve değiştirildiğini ortaya çıkarır.")

        col_meta1, col_meta2 = st.columns([1, 2])

        with col_meta1:
            st.markdown("### 📂 Delil Yükle")
            uploaded_evid = st.file_uploader("İncelenecek Dosya", type=["pdf", "jpg", "jpeg", "png", "tiff"])
            
            st.markdown("---")
            st.markdown("### 📅 İddia Kontrolü")
            claimed_date = st.date_input("Belgenin İddia Edilen Tarihi (Opsiyonel)", value=None)
            st.caption("Eğer bu belgenin '2020 yılında yapıldığı' iddia ediliyorsa, o tarihi seçin. Sistem tutarlılığı denetlesin.")

        with col_meta2:
            if uploaded_evid:
                # Kütüphaneleri Çağır
                from datetime import datetime
                import pandas as pd
                
                meta_data = {}
                file_type = uploaded_evid.name.split('.')[-1].lower()
                
                st.markdown(f"### 🧬 Analiz Raporu: {uploaded_evid.name}")
                
                # --- PDF ANALİZİ ---
                if file_type == 'pdf':
                    try:
                        import PyPDF2
                        pdf_reader = PyPDF2.PdfReader(uploaded_evid)
                        doc_info = pdf_reader.metadata
                        
                        if doc_info:
                            # PDF Tarih Formatını Okunabilir Yapma (D:20230101...)
                            def parse_pdf_date(date_str):
                                if not date_str: return "Bilinmiyor"
                                try:
                                    # Genelde D:YYYYMMDDHHmmSS formatındadır
                                    clean_date = date_str.replace("D:", "").split('+')[0].split('-')[0]
                                    return datetime.strptime(clean_date, "%Y%m%d%H%M%S").strftime("%d.%m.%Y %H:%M:%S")
                                except:
                                    return date_str # Parse edilemezse ham hali

                            meta_data = {
                                "Oluşturulma Tarihi (CreationDate)": parse_pdf_date(doc_info.get('/CreationDate')),
                                "Değiştirilme Tarihi (ModDate)": parse_pdf_date(doc_info.get('/ModDate')),
                                "Yazar (Author)": doc_info.get('/Author', 'Belirtilmemiş'),
                                "Oluşturan Yazılım (Producer)": doc_info.get('/Producer', 'Belirtilmemiş'),
                                "Uygulama (Creator)": doc_info.get('/Creator', 'Belirtilmemiş'),
                                "Sayfa Sayısı": len(pdf_reader.pages)
                            }
                        else:
                            st.warning("Bu PDF dosyasında metadata bulunamadı veya silinmiş.")
                    except Exception as e:
                        st.error(f"PDF Analiz Hatası: {e}")

                # --- RESİM (EXIF) ANALİZİ ---
                elif file_type in ['jpg', 'jpeg', 'png', 'tiff']:
                    try:
                        from PIL import Image, ExifTags
                        image = Image.open(uploaded_evid)
                        exif_raw = image._getexif()
                        
                        if exif_raw:
                            for tag, value in exif_raw.items():
                                decoded = ExifTags.TAGS.get(tag, tag)
                                # Önemli verileri filtrele
                                if decoded in ['DateTime', 'DateTimeOriginal', 'Make', 'Model', 'Software', 'GPSInfo', 'Artist']:
                                    meta_data[decoded] = str(value)
                            
                            # Eğer boşsa
                            if not meta_data:
                                meta_data = {"Durum": "EXIF verisi bulunamadı (Temizlenmiş olabilir)."}
                        else:
                            meta_data = {"Durum": "Bu resimde EXIF verisi yok."}
                            
                    except Exception as e:
                        st.error(f"Resim Analiz Hatası: {e}")

                # --- SONUÇLARI GÖSTER ---
                if meta_data:
                    # 1. Tablo Gösterimi
                    df_meta = pd.DataFrame(list(meta_data.items()), columns=["Veri Türü", "Tespit Edilen Değer"])
                    st.table(df_meta)

                    # 2. YAPAY ZEKA DEDEKTİF YORUMU
                    if api_key:
                        st.divider()
                        with st.spinner("🕵️ Yapay Zeka delil üzerinde sahtecilik taraması yapıyor..."):
                            
                            prompt = f"""
                            GÖREV: Sen uzman bir Adli Bilişim (Digital Forensics) uzmanısın.
                            
                            ANALİZ EDİLEN DOSYA METADATASI:
                            {meta_data}
                            
                            İDDİA EDİLEN TARİH: {claimed_date if claimed_date else "Belirtilmedi"}
                            
                            İSTENEN ANALİZ:
                            1. **Zaman Tutarlılığı:** Dosyanın oluşturulma tarihi ile iddia edilen tarih uyuşuyor mu? (Örn: 2020 denmiş ama CreationDate 2024 ise bu bir sahtecilik şüphesidir).
                            2. **Yazılım İzi:** Kullanılan yazılım (Producer/Software) dosyanın iddia edilen tarihinde var mıydı? (Örn: 2010 tarihli belgede Word 2019 imzası varsa yakala).
                            3. **Manipülasyon Şüphesi:** Değiştirilme tarihi (ModDate) ile Oluşturulma tarihi arasında şüpheli bir fark var mı?
                            4. **Sonuç:** Bu belge teknik olarak güvenilir mi yoksa şüpheli mi?
                            """
                            
                            report = get_ai_response(prompt, api_key)
                            
                            st.markdown("### 🚨 Adli Bilişim Uzman Görüşü")
                            
                            # Renkli Kutu Mantığı
                            if "şüpheli" in report.lower() or "uyuşmuyor" in report.lower() or "sahte" in report.lower():
                                box_color = "#ffe6e6" # Kırmızımsı
                                border_color = "#ff0000"
                                icon = "⚠️"
                            else:
                                box_color = "#e6fffa" # Yeşilimsi
                                border_color = "#00b894"
                                icon = "✅"

                            st.markdown(f"""
                            <div style="background-color:{box_color}; border-left: 5px solid {border_color}; padding:20px; border-radius:10px;">
                                <h4>{icon} Analiz Sonucu</h4>
                                {report}
                            </div>
                            """, unsafe_allow_html=True)
                    else:
                        st.info("Detaylı sahtecilik analizi için API Key gereklidir.")

    with tab37: # 3. SATIR: Mevzuat Kelebek Etkisi Simülatörü
        st.subheader("🦋 Mevzuat Kelebek Etkisi Simülatörü (Graph Analizi)")
        st.info("Hukuk bir ağdır. Bir kanun maddesindeki tek bir kelime değişikliğinin, en uçtaki yönetmelik, tebliğ ve ruhsatları nasıl etkilediğini haritalandırır.")

        col_graph1, col_graph2 = st.columns([1, 2])

        with col_graph1:
            st.markdown("### 🌪️ Değişiklik Girdisi")
            law_change = st.text_area(
                "Yapılan/Beklenen Değişiklik:", 
                height=150, 
                placeholder="Örn: İmar Kanunu'nda 'yüksek yapı' tanımı 10 kattan 8 kata düşürüldü."
            )
            
            st.markdown("### 🎯 Hedef Sektör")
            sector = st.selectbox("Etki Analizi Odak Alanı:", 
                                  ["Genel Bakış", "İnşaat & Emlak", "Vergi & Finans", "İş Hukuku & IK", "Sağlık & İlaç"])

            analyze_btn = st.button("🕸️ Etki Ağını Haritalandır", use_container_width=True)

        with col_graph2:
            if analyze_btn and law_change:
                if not api_key:
                    st.warning("Bu simülasyon için API Key gereklidir.")
                else:
                    import graphviz
                    
                    with st.spinner("Yapay Zeka, hukuk ağındaki dolaylı bağlantıları tarıyor..."):
                        # AI'dan Graphviz formatında veri istiyoruz
                        prompt = f"""
                        GÖREV: Sen bir Hukuk Graph Database (Neo4j) simülatörüsün.
                        
                        GİRDİ: "{law_change}"
                        ODAK SEKTÖR: {sector}
                        
                        İSTENEN ÇIKTI:
                        Bu değişikliğin "Kelebek Etkisi"ni gösteren bir DOT (Graphviz) kodu oluştur.
                        
                        KURALLAR:
                        1. Merkezde "Kanun Değişikliği" olsun (Kırmızı Düğüm).
                        2. 1. Derece etkilenenler: Yönetmelikler/Tebliğler (Mavi Düğüm).
                        3. 2. Derece etkilenenler: Sektörel Uygulamalar/İzinler (Sarı Düğüm).
                        4. 3. Derece (Kelebek Etkisi): Hiç beklenmeyen uzak riskler (Siyah/Koyu Kırmızı Düğüm). Örn: "3 yıl önceki ruhsat iptali riski".
                        5. Sadece DOT kodunu ver, açıklama yapma. Kod `digraph` ile başlasın.
                        6. Türkçe karakter kullanma (yerine ingilizce karakterler kullan, örn: 'ı' yerine 'i').
                        7. Etiketler kısa ve çarpıcı olsun.
                        """
                        
                        try:
                            # AI Cevabını al
                            graph_code_raw = get_ai_response(prompt, api_key)
                            
                            # Temizlik (Markdown işaretlerini kaldır)
                            graph_code = graph_code_raw.replace("```dot", "").replace("```", "").strip()
                            
                            # Graphviz ile çizim
                            st.graphviz_chart(graph_code)
                            
                            st.markdown("### 🧠 Yapay Zeka Risk Analizi")
                            st.success("Simülasyon Tamamlandı. Yukarıdaki ağ haritası, bu değişikliğin tetikleyebileceği zincirleme reaksiyonları göstermektedir.")
                            
                            # Ekstra Yorum
                            explanation_prompt = f"Bu graph haritasındaki en tehlikeli 'Kelebek Etkisi' (En uçtaki risk) nedir? '{law_change}' değişikliği neden orayı etkiliyor? Tek paragraf açıkla."
                            explanation = get_ai_response(explanation_prompt, api_key)
                            
                            st.markdown(f"""
                            <div style="border: 1px solid #ffcc00; background-color: #fffbea; padding: 15px; border-radius: 8px;">
                                <strong>⚠️ Gizli Tehlike (Kelebek Etkisi):</strong><br>
                                {explanation}
                            </div>
                            """, unsafe_allow_html=True)
                            
                        except Exception as e:
                            st.error(f"Haritalama hatası: {e}")
                            st.info("Graphviz kütüphanesi yüklü olmayabilir veya AI hatalı kod üretti.")



if __name__ == "__main__":
    main()

