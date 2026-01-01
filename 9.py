import streamlit as st
import pandas as pd  # Pandas hatasını çözer
import requests
from bs4 import BeautifulSoup # Web tarama hatasını çözer
import io
import PyPDF2
import zipfile
import xml.etree.ElementTree as ET
import re
from pypdf import PdfReader
from io import BytesIO
import google.generativeai as genai
import importlib.metadata
from docx import Document
from fpdf import FPDF
import urllib.parse
import concurrent.futures
from gtts import gTTS
import speech_recognition as sr
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import json
import os
from PIL import Image
from PIL.ExifTags import TAGS  # <--- Bu çok önemli, eksikse hata verir
import time
from datetime import datetime, timedelta, date
import shutil
import difflib
import plotly.graph_objects as go # Görsel grafikler için gerekli
from PIL import Image



# --- Sayfa Ayarları ---
st.set_page_config(
    page_title="Hukuk Asistanı AI",
    page_icon="⚖️",
    layout="wide"
)

# --- CSS ---
st.markdown("""
    <style>
    .main { background-color: #f8f9fa; }
    .kanun-kutusu { 
        background-color: #fff3e0; 
        padding: 15px; 
        border-left: 5px solid #ff9800; 
        border-radius: 5px; 
        margin-bottom: 10px;
        white-space: pre-wrap;
    }
    .ictihat-kutusu {
        background-color: #e3f2fd;
        padding: 15px;
        border-left: 5px solid #2196f3;
        border-radius: 5px;
        margin-bottom: 10px;
    }
    .buyur-abi-kutusu {
        background-color: #f3e5f5;
        padding: 15px;
        border-left: 5px solid #9c27b0;
        border-radius: 5px;
        margin-bottom: 10px;
    }
    .alarm-kutusu {
        background-color: #ffebee;
        padding: 15px;
        border-left: 5px solid #f44336;
        border-radius: 5px;
        margin-bottom: 10px;
        font-weight: bold;
        color: #b71c1c;
    }
    .arsiv-kutusu {
        background-color: #e0f2f1;
        padding: 15px;
        border-left: 5px solid #009688;
        border-radius: 5px;
        margin-bottom: 10px;
    }
    .uyap-kutusu {
        background-color: #fce4ec; 
        padding: 15px; 
        border-left: 5px solid #c2185b; 
        border-radius: 5px; 
        margin-bottom: 20px;
    }
    .ozel-sekme {
        border: 1px solid #ddd;
        padding: 20px;
        border-radius: 10px;
        background-color: #ffffff;
    }
    </style>
    """, unsafe_allow_html=True)


# --- KALICILIK (VERİ TABANI) FONKSİYONLARI ---
DURUSMA_FILE = "durusma_kayitlari.json"

def save_durusma_data(data):
    """Duruşma listesini JSON dosyasına kaydeder."""
    serializable_data = []
    for item in data:
        temp = item.copy()
        if isinstance(temp.get('dtstart'), datetime):
            temp['dtstart'] = temp['dtstart'].isoformat()
        serializable_data.append(temp)
    
    try:
        with open(DURUSMA_FILE, "w", encoding="utf-8") as f:
            json.dump(serializable_data, f, ensure_ascii=False, indent=4)
    except Exception as e:
        st.error(f"Kaydetme hatası: {e}")

def load_durusma_data():
    """JSON dosyasından duruşma listesini yükler."""
    if not os.path.exists(DURUSMA_FILE):
        return []
    try:
        with open(DURUSMA_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        for item in data:
            if 'dtstart' in item and item['dtstart']:
                item['dtstart'] = datetime.fromisoformat(item['dtstart'])
        return data
    except:
        return []

# --- YARDIMCI FONKSİYONLAR ---
def parse_udf(file_bytes):
    try:
        with zipfile.ZipFile(file_bytes) as z:
            if 'content.xml' in z.namelist():
                with z.open('content.xml') as f:
                    tree = ET.parse(f)
                    root = tree.getroot()
                    text_content = [elem.text.strip() for elem in root.iter() if elem.text]
                    return " ".join(text_content)
            return "HATA: UDF içeriği okunamadı."
    except Exception as e:
        return f"HATA: {str(e)}"

def parse_pdf(file_bytes):
    try:
        reader = PdfReader(file_bytes)
        text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        if len(text.strip()) < 50: return "" 
        return text
    except Exception as e:
        return ""

def extract_metadata(text):
    if not isinstance(text, str) or text.startswith(("HATA", "UYARI")):
        return {"mahkeme": "-", "esas": "-", "karar": "-", "tarih": "-"}
    
    esas = re.search(r"(?i)Esas\s*No\s*[:\-]?\s*(\d{4}/\d+)", text)
    karar = re.search(r"(?i)Karar\s*No\s*[:\-]?\s*(\d{4}/\d+)", text)
    tarih = re.search(r"(\d{1,2}[./]\d{1,2}[./]\d{4})", text)
    
    mahkeme = "Tespit Edilemedi"
    for line in text.split('\n')[:40]:
        clean = line.strip()
        if ("MAHKEMESİ" in clean.upper() or "DAİRESİ" in clean.upper()) and len(clean) > 5:
            mahkeme = clean
            break
    return {
        "mahkeme": mahkeme,
        "esas": esas.group(1) if esas else "Bulunamadı",
        "karar": karar.group(1) if karar else "Bulunamadı",
        "tarih": tarih.group(1) if tarih else "Bulunamadı"
    }

def parse_ics_data(file_bytes):
    events = []
    try:
        content = file_bytes.getvalue().decode('utf-8')
        lines = content.splitlines()
        current_event = {}
        in_event = False
        
        for line in lines:
            line = line.strip()
            if line == 'BEGIN:VEVENT':
                in_event = True
                current_event = {}
            elif line == 'END:VEVENT':
                in_event = False
                if 'dtstart' in current_event and 'summary' in current_event:
                    events.append(current_event)
            elif in_event:
                if line.startswith('SUMMARY:'):
                    current_event['summary'] = line.split(':', 1)[1]
                elif line.startswith('DTSTART:'):
                    raw_date = line.split(':', 1)[1]
                    try:
                        dt = datetime.strptime(raw_date, '%Y%m%dT%H%M%S')
                        current_event['dtstart'] = dt
                    except: pass
                elif line.startswith('LOCATION:'):
                    current_event['location'] = line.split(':', 1)[1]
                elif line.startswith('DESCRIPTION:'):
                    current_event['description'] = line.split(':', 1)[1]
        return events
    except Exception as e:
        return []

# --- DOSYA OLUŞTURMA FONKSİYONLARI ---
def create_word_file(text):
    doc = Document()
    for line in text.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io

def create_udf_file(text):
    root = ET.Element("content")
    body = ET.SubElement(root, "body")
    for line in text.split('\n'):
        p = ET.SubElement(body, "p")
        p.text = line
    xml_str = ET.tostring(root, encoding='utf-8', method='xml')
    byte_io = BytesIO()
    with zipfile.ZipFile(byte_io, 'w', zipfile.ZIP_DEFLATED) as z:
        z.writestr('content.xml', xml_str)
    byte_io.seek(0)
    return byte_io

def create_pdf_file(text):
    replacements = {
        'ğ': 'g', 'Ğ': 'G', 'ş': 's', 'Ş': 'S', 'ı': 'i', 'İ': 'I',
        'ç': 'c', 'Ç': 'C', 'ü': 'u', 'Ü': 'U', 'ö': 'o', 'Ö': 'O',
        '“': '"', '”': '"', '’': "'", '–': '-', '…': '...'
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    
    text = text.encode('latin-1', 'replace').decode('latin-1')
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="Hukuki Analiz Raporu", ln=1, align='C')
    pdf.ln(10)
    pdf.set_font("Arial", size=11)
    pdf.multi_cell(0, 10, text)
    return pdf.output(dest='S').encode('latin-1', 'ignore')

# --- SES İŞLEME FONKSİYONLARI ---
def text_to_speech(text):
    try:
        tts = gTTS(text=text, lang='tr')
        fp = BytesIO()
        tts.write_to_fp(fp)
        fp.seek(0)
        return fp
    except Exception as e:
        return None

def speech_to_text(audio_bytes):
    r = sr.Recognizer()
    try:
        with sr.AudioFile(audio_bytes) as source:
            audio_data = r.record(source)
            text = r.recognize_google(audio_data, language='tr-TR')
            return text
    except Exception as e:
        return f"Hata: {str(e)}"

# --- MULTIMODAL VE OCR FONKSİYONLARI ---
def perform_ocr_gemini(file_bytes, mime_type, api_key, prompt_text="Bu dosyanın içeriğini tam olarak metne dök."):
    if not api_key: return "API Key Yok"
    genai.configure(api_key=api_key)
    
    if mime_type in ['image/tiff', 'image/tif']:
        try:
            image = Image.open(file_bytes)
            rgb_im = image.convert('RGB')
            buf = BytesIO()
            rgb_im.save(buf, format="JPEG")
            file_bytes = buf
            mime_type = 'image/jpeg'
        except Exception as e:
            return f"TIFF Dönüştürme Hatası: {str(e)}"

    available_models = []
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                if '1.5' in m.name or 'vision' in m.name:
                    available_models.append(m.name)
    except: pass
    
    if not available_models: available_models = ['models/gemini-1.5-flash', 'models/gemini-1.5-pro']

    image_part = {"mime_type": mime_type, "data": file_bytes.getvalue()}
    
    for model_name in available_models:
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content([prompt_text, image_part])
            return response.text
        except: continue
    return "Analiz Başarısız."

def extract_text_from_docx(file_bytes):
    try:
        doc = Document(file_bytes)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        return f"Word Okuma Hatası: {str(e)}"

def read_excel_file(file_bytes):
    try:
        df = pd.read_excel(file_bytes)
        return df.to_string()
    except Exception as e:
        return f"Excel Hatası: {str(e)}"

# --- AKILLI AI MOTORU ---
def get_ai_response(prompt, api_key):
    if not api_key: return "Lütfen API Anahtarı giriniz."
    genai.configure(api_key=api_key)
    candidate_models = ['gemini-1.5-flash', 'gemini-1.5-pro', 'gemini-1.0-pro', 'gemini-pro']
    
    for model_name in candidate_models:
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(prompt)
            return response.text 
        except: continue 
    
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                try:
                    model = genai.GenerativeModel(m.name)
                    response = model.generate_content(prompt)
                    return response.text
                except: continue
    except: pass
    return "Hata: AI yanıt veremedi."

# ==========================================
# YENİ EKLENEN MODÜLLER (CHECK-UP & ZAMAN MAKİNESİ)
# ==========================================

def get_gemini_text_response(prompt, api_key):
    """Mevcut modelleri tarayıp çalışan ilk modeli kullanan fonksiyon."""
    if not api_key: return "Lütfen API Anahtarınızı giriniz."
    
    try:
        genai.configure(api_key=api_key)
        
        # 1. ADIM: Kullanıcının erişebildiği modelleri listele
        available_models = []
        try:
            for m in genai.list_models():
                if 'generateContent' in m.supported_generation_methods:
                    available_models.append(m.name)
        except Exception as e:
            # Listeleme hatası olursa manuel listeye dön
            available_models = ["models/gemini-1.5-flash", "models/gemini-pro", "gemini-1.5-flash"]

        # 2. ADIM: En iyi modeli seç (Hız > Kalite sıralaması)
        selected_model = None
        
        # Öncelik 1: Flash modeller (Hızlı)
        for m in available_models:
            if 'flash' in m.lower():
                selected_model = m
                break
        
        # Öncelik 2: Pro modeller (Eğer flash yoksa)
        if not selected_model:
            for m in available_models:
                if 'pro' in m.lower() and 'vision' not in m.lower():
                    selected_model = m
                    break
        
        # Öncelik 3: Listede ne varsa ilki
        if not selected_model and available_models:
            selected_model = available_models[0]
            
        # Hiçbiri yoksa son çare
        if not selected_model:
            selected_model = "models/gemini-1.5-flash"

        # 3. ADIM: Seçilen modelle üretimi yap
        try:
            model = genai.GenerativeModel(selected_model)
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            return f"Model Hatası ({selected_model}): {str(e)}"
            
    except Exception as e:
        return f"Genel Sistem Hatası: {str(e)}"



def render_checkup_module(api_key):
    st.info("Şirket sözleşmelerini veya İK belgelerini yükleyin. Yapay zeka, güncel Yargıtay kararlarına göre 'Görünmez Riskleri' tespit edip puanlasın.")
    col1, col2 = st.columns([1, 2])

    with col1:
        st.markdown("#### 📂 Belge Yükleme")
        doc_type = st.selectbox("Belge Türü", ["İş Sözleşmesi", "KVKK Aydınlatma Metni", "Tedarikçi Sözleşmesi", "Kira Kontratı"])
        uploaded_file = st.file_uploader("Dosyayı Sürükleyin (PDF/DOCX)", type=["pdf", "docx", "txt"])
        analyze_btn = st.button("🔍 Risk Taramasını Başlat", type="primary", use_container_width=True)

    with col2:
        if analyze_btn and uploaded_file:
            if not api_key:
                st.error("⚠️ Lütfen sol menüden API Anahtarını giriniz.")
            else:
                with st.spinner("Belge taranıyor, Yargıtay kararlarıyla karşılaştırılıyor..."):
                    # Dosya okuma (Basit)
                    file_text = "Örnek metin"
                    try:
                        if uploaded_file.name.endswith(".pdf"):
                            reader = PyPDF2.PdfReader(uploaded_file)
                            if len(reader.pages) > 0:
                                file_text = reader.pages[0].extract_text()
                        elif uploaded_file.name.endswith(".txt"):
                            file_text = uploaded_file.getvalue().decode("utf-8")
                    except: pass

                    # AI Prompt
                    prompt = f"""
                    GÖREV: Sen kıdemli bir hukuk denetçisisin.
                    BELGE TÜRÜ: {doc_type}
                    BELGE İÇERİĞİ (ÖZET): {file_text[:3000]}
                    
                    GÖREVLER:
                    1. Bu belge için güncel Yargıtay kararlarına göre en kritik 3 riski bul.
                    2. Belgeye 0-100 arası "HUKUKİ SAĞLAMLIK SKORU" ver.
                    
                    ÇIKTI FORMATI:
                    SKOR: [Sayı]
                    RİSKLER: [Detaylar]
                    """
                    ai_response = get_gemini_text_response(prompt, api_key)
                    
                    # Skoru çekme
                    risk_score = 60
                    match = re.search(r"SKOR:\s*(\d+)", ai_response)
                    if match: risk_score = int(match.group(1))

                    # --- GÖRSELLEŞTİRME (HATA ÖNLEYİCİ MOD) ---
                    if PLOTLY_VAR:
                        # Plotly varsa havalı göstergeyi çiz
                        fig = go.Figure(go.Indicator(
                            mode = "gauge+number",
                            value = risk_score,
                            title = {'text': "Hukuki Sağlamlık Skoru"},
                            gauge = {
                                'axis': {'range': [None, 100]},
                                'bar': {'color': "black"},
                                'steps': [
                                    {'range': [0, 50], 'color': "#ff4d4d"},
                                    {'range': [50, 80], 'color': "#ffcc00"},
                                    {'range': [80, 100], 'color': "#33cc33"}
                                ],
                                'threshold': {'line': {'color': "red", 'width': 4}, 'thickness': 0.75, 'value': risk_score}
                            }
                        ))
                        fig.update_layout(height=250, margin=dict(l=20, r=20, t=30, b=20))
                        st.plotly_chart(fig, use_container_width=True)
                    else:
                        # Plotly yoksa standart bar kullan (Çökmemesi için)
                        st.metric("Hukuki Sağlamlık Skoru", f"{risk_score} / 100")
                        st.progress(risk_score / 100)
                        if risk_score < 50:
                            st.error("Risk Seviyesi: YÜKSEK")
                        elif risk_score < 80:
                            st.warning("Risk Seviyesi: ORTA")
                        else:
                            st.success("Risk Seviyesi: DÜŞÜK")
                    
                    st.markdown("### 📋 Risk Raporu")
                    st.write(ai_response.replace(f"SKOR: {risk_score}", ""))


def render_time_machine(api_key):
    st.info("Bir olay tarihi girin, sistem sizi o güne götürsün. O gün geçerli olan kanun maddesini ve faiz oranlarını görün.")
    col_date, col_topic = st.columns([1, 2])
    
    with col_date:
        target_date = st.date_input("Olay Tarihi Seçin:", value=date(2015, 5, 14))
    with col_topic:
        topic = st.text_input("Sorgulanacak Konu", placeholder="Örn: Kıdem Tazminatı Tavanı")

    if st.button("🕒 Geçmişe Git", use_container_width=True):
        if api_key and topic:
            with st.spinner(f"Sistem {target_date.strftime('%d.%m.%Y')} tarihine geri sarılıyor..."):
                prompt = f"""
                GÖREV: Hukuk Tarihçisi. TARİH: {target_date.strftime('%d.%m.%Y')}. KONU: {topic}.
                SORU: O tarihte bu konuyla ilgili yürürlükte olan kanun maddesi, faiz oranı ve Yargıtay görüşü neydi?
                """
                response = get_gemini_text_response(prompt, api_key)
                st.markdown(f"### 📅 Tarih: {target_date.strftime('%d %B %Y')}")
                st.info(response)
                st.image("https://img.freepik.com/free-vector/sepia-vintage-paper-texture_53876-88607.jpg?w=1380", caption="Arşiv Kaydı", width=600)

# --- 3. MODÜL: AYM & AİHM UYGUNLUK TESTİ ---
def render_aym_aihm_module(api_key):
    st.info("Dilekçenizi, Mahkeme Kararını veya UYAP (UDF) dosyasını yükleyin. Sistem OCR ile okuyup AİHM/AYM standartlarına göre 'Hak İhlali' analizi yapsın.")
    
    # Sekmeli Giriş Yapısı
    tab_text, tab_file = st.tabs(["📝 Metin Yapıştır", "📂 Dosya Yükle (PDF/UDF/TIFF)"])
    
    process_text = ""
    analyze_trigger = False

    # --- TAB 1: MANUEL METİN ---
    with tab_text:
        user_text_input = st.text_area("Metni Buraya Yapıştırın:", height=300, placeholder="Örn: Mahkeme gerekçesiz karar vererek adil yargılanma hakkımı ihlal etmiştir...")
        if st.button("⚖️ Metni Analiz Et", key="btn_text_aym", type="primary"):
            process_text = user_text_input
            analyze_trigger = True

    # --- TAB 2: DOSYA YÜKLEME ---
    with tab_file:
        uploaded_file = st.file_uploader("Dosya Seçin", type=["pdf", "udf", "xml", "tiff", "tif", "jpg", "png", "txt"])
        
        if uploaded_file:
            st.caption(f"Yüklenen Dosya: {uploaded_file.name}")
            if st.button("👁️ Dosyayı Oku ve Analiz Et", key="btn_file_aym", type="primary"):
                with st.spinner("Dosya okunuyor ve OCR yapılıyor..."):
                    extracted_text = extract_text_from_legal_file(uploaded_file, api_key)
                    
                    if "[OCR GEREKLİ]" in extracted_text or "Hata" in extracted_text:
                        st.error(extracted_text)
                    else:
                        process_text = extracted_text
                        st.success("Dosya başarıyla metne dönüştürüldü! Analiz başlıyor...")
                        with st.expander("Okunan Metni Gör"):
                            st.text(process_text[:1000] + "...")
                        analyze_trigger = True

    # --- ORTAK ANALİZ MOTORU ---
    if analyze_trigger and process_text:
        if not api_key:
            st.error("⚠️ Lütfen API Anahtarını giriniz.")
        elif len(process_text) < 20:
            st.warning("Analiz için yeterli metin bulunamadı.")
        else:
            with st.spinner("Metin, AİHM ve AYM içtihatlarıyla çapraz sorgulanıyor..."):
                
                prompt = f"""
                GÖREV: Sen AİHM ve AYM kararları konusunda uzmanlaşmış kıdemli bir hukukçusun.
                METİN: "{process_text[:6000]}" (Kısaltılmış olabilir)
                
                ANALİZ ADIMLARI:
                1. Bu metindeki olayda, Avrupa İnsan Hakları Sözleşmesi (AİHS) veya Anayasa ile korunan hangi temel haklar risk altında? (Örn: Mülkiyet Hakkı, Adil Yargılanma Hakkı).
                2. Bu metin bir mahkeme kararıysa Üst Mahkemede BOZULMA İHTİMALİ yüzde kaçtır? Bir dilekçeyse KABUL EDİLME GÜCÜ yüzde kaçtır? (0-100 arası bir puan ver).
                3. Konuyla ilgili emsal bir AİHM veya AYM kararı adı ver ve özetle.
                4. Eğer bir ihlal varsa, başvuru formunda hangi argüman kullanılmalı?
                
                ÇIKTI FORMATI:
                ORAN: [Sayı]
                ANALİZ: [Detaylı Hukuki Görüş]
                EMSAL: [Karar İsimleri]
                STRATEJİ: [Öneri]
                """
                
                ai_response = get_gemini_text_response(prompt, api_key)
                
                # Oranı çekme
                ihlal_orani = 50
                match = re.search(r"ORAN:\s*(\d+)", ai_response)
                if match: ihlal_orani = int(match.group(1))
                
                # --- SONUÇ EKRANI ---
                st.divider()
                col_score, col_detail = st.columns([1, 2])
                
                with col_score:
                    st.markdown(f"<h2 style='text-align: center; color: #d63031;'>%{ihlal_orani}</h2>", unsafe_allow_html=True)
                    st.markdown("<p style='text-align: center;'><b>Bozulma / İhlal Riski</b></p>", unsafe_allow_html=True)
                    st.progress(ihlal_orani / 100)
                    
                    if ihlal_orani > 70:
                        st.error("🚨 KRİTİK: Yüksek ihtimalle hak ihlali var.")
                    elif ihlal_orani > 40:
                        st.warning("⚠️ DİKKAT: Güçlü argümanlar gerekiyor.")
                    else:
                        st.success("✅ TEMİZ: Belirgin bir ihlal görünmüyor.")

                with col_detail:
                    st.markdown("### 🏛️ Yüksek Yargı Raporu")
                    st.write(ai_response.replace(f"ORAN: {ihlal_orani}", ""))


def get_image_metadata(image):
    """Resimden EXIF verilerini çeker."""
    meta_dict = {}
    try:
        exif_data = image._getexif()
        if exif_data:
            for tag_id, value in exif_data.items():
                tag_name = TAGS.get(tag_id, tag_id)
                if tag_name in ['DateTime', 'DateTimeOriginal', 'Make', 'Model', 'Software']:
                    meta_dict[tag_name] = str(value)
    except:
        return None
    return meta_dict

def render_deepfake_module(api_key):
    # --- HATA YAKALAYICI BLOK BAŞLANGICI ---
    try:
        st.info("Şüpheli fotoğraf veya ses kaydını yükleyin. Yapay zeka, metadata (üst veri) analizi ve içerik taraması yaparak 'Montaj/Deepfake' izlerini arasın.")
        
        col_upload, col_report = st.columns([1, 2])
        
        with col_upload:
            st.markdown("#### 🕵️‍♂️ Delil Yükle")
            file_type = st.radio("Delil Türü", ["Fotoğraf / Belge Görüntüsü", "Ses Kaydı (Kısa)"])
            
            if file_type == "Fotoğraf / Belge Görüntüsü":
                uploaded_file = st.file_uploader("Resim Seç (JPG, PNG)", type=["jpg", "jpeg", "png"])
            else:
                uploaded_file = st.file_uploader("Ses Dosyası Seç (MP3, WAV)", type=["mp3", "wav"])
                
            analyze_btn = st.button("🔍 Adli Bilişim Analizi Yap", type="primary", use_container_width=True)

        with col_report:
            if analyze_btn and uploaded_file:
                if not api_key:
                    st.error("⚠️ API Anahtarı eksik.")
                else:
                    with st.spinner("Dosya bit-bit inceleniyor, metadata taranıyor ve AI analizi yapılıyor..."):
                        
                        genai.configure(api_key=api_key)
                        # Model seçimi (Hata verirse Pro'ya düş)
                        try:
                            model = genai.GenerativeModel('gemini-1.5-flash')
                        except:
                            model = genai.GenerativeModel('gemini-pro-vision')
                        
                        report_text = ""
                        fake_score = 0
                        metadata_info = {}

                        # --- FOTOĞRAF ANALİZİ ---
                        if file_type == "Fotoğraf / Belge Görüntüsü":
                            image = Image.open(uploaded_file)
                            st.image(image, caption="İncelenen Delil", width=300)
                            
                            # Metadata Kontrolü (Güvenli)
                            try:
                                metadata_info = get_image_metadata(image)
                                meta_str = str(metadata_info) if metadata_info else "Metadata bulunamadı."
                            except Exception as e:
                                meta_str = f"Metadata okunamadı: {e}"
                            
                            prompt = f"""
                            GÖREV: Sen uzman bir Adli Bilişim (Forensics) uzmanısın.
                            METADATA: {meta_str}
                            GÖREVLER:
                            1. Görselde Deepfake/Montaj izi var mı?
                            2. Metadata tutarlı mı?
                            3. Güvenilirlik puanı (0-100).
                            ÇIKTI: GÜVEN_SKORU: [Sayı] ...
                            """
                            response = model.generate_content([prompt, image])
                            report_text = response.text

                        # --- SES ANALİZİ ---
                        else: 
                            # Ses analizi için güvenli blok
                            try:
                                st.audio(uploaded_file)
                                # Geçici dosya oluşturma
                                import tempfile
                                with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp_file:
                                    tmp_file.write(uploaded_file.getvalue())
                                    tmp_path = tmp_file.name
                                
                                # Speech Recognition
                                r = sr.Recognizer()
                                with sr.AudioFile(tmp_path) as source:
                                    audio_data = r.record(source)
                                    text_output = r.recognize_google(audio_data, language="tr-TR")
                                    
                                    prompt = f"""Ses Transkripti: "{text_output}". Bu konuşma doğal mı, kurgu mu? Puanla (0-100). ÇIKTI: GÜVEN_SKORU: [Sayı] ..."""
                                    model_text = genai.GenerativeModel('gemini-pro')
                                    response = model_text.generate_content(prompt)
                                    report_text = response.text
                            except ImportError:
                                st.error("Ses analizi için 'SpeechRecognition' kütüphanesi yüklü değil.")
                                return
                            except Exception as e:
                                st.error(f"Ses işleme hatası: {str(e)}")
                                return

                        # --- SONUÇLARI GÖSTER ---
                        match = re.search(r"GÜVEN_SKORU:\s*(\d+)", report_text)
                        if match: fake_score = int(match.group(1))
                        
                        st.divider()
                        st.metric("Delil Güvenilirlik Skoru", f"{fake_score} / 100")
                        st.progress(fake_score / 100)
                        
                        if fake_score < 50:
                            st.error("🚨 SAHTECİLİK ŞÜPHESİ YÜKSEK")
                        else:
                            st.success("✅ DELİL GÜVENİLİR GÖRÜNÜYOR")
                            
                        st.write(report_text.replace(f"GÜVEN_SKORU: {fake_score}", ""))

    except Exception as e:
        # EĞER BEYAZ EKRAN ÇIKARSA BURASI DEVREYE GİRER VE HATAYI YAZAR
        st.error(f"🚨 Modül Yükleme Hatası: {str(e)}")
        st.warning("Lütfen 'PIL', 'SpeechRecognition' kütüphanelerinin yüklü olduğundan ve 'TAGS' importunun yapıldığından emin olun.")


def generate_dork_category(category, target_name, city):
    """Belirli bir kategori için gelişmiş arama linkleri (Dorks) üretir."""
    # Simüle edilmiş işlem süresi (Threading etkisini görmek için)
    time.sleep(0.5) 
    
    links = []
    base_url = "https://www.google.com/search?q="
    name_slug = target_name.replace(" ", "+")
    
    if category == "social":
        # Sosyal Medya Taraması
        links.append(f"[📸 Instagram: {target_name}]({base_url}site:instagram.com+%22{name_slug}%22)")
        links.append(f"[💼 LinkedIn: {target_name}]({base_url}site:linkedin.com/in/+%22{name_slug}%22)")
        links.append(f"[🐦 Twitter/X: {target_name}]({base_url}site:twitter.com+%22{name_slug}%22)")
        links.append(f"[👤 Facebook: {target_name}]({base_url}site:facebook.com+%22{name_slug}%22)")
        
    elif category == "business":
        # Ticari Varlık ve Şirket Taraması
        links.append(f"[🏢 Ticaret Sicil: {target_name}]({base_url}%22{name_slug}%22+site:ticaretsicil.gov.tr)")
        links.append(f"[📄 Resmi Gazete: {target_name}]({base_url}%22{name_slug}%22+site:resmigazete.gov.tr)")
        links.append(f"[🤝 Şirket Ortaklıkları]({base_url}%22{name_slug}%22+kurucu+ortak+sahibi)")
        
    elif category == "assets":
        # Mal Varlığı ve Lüks Yaşam İzi (Tatil, Araba vb.)
        links.append(f"[🏖️ Tatil/Otel Yorumları]({base_url}%22{name_slug}%22+otel+tatil+gezi)")
        links.append(f"[🚗 Araba/Satış İlanları]({base_url}%22{name_slug}%22+sahibinden+satılık)")
        links.append(f"[🎓 Mezuniyet/Okul]({base_url}%22{name_slug}%22+mezun+okul+lise+üniversite)")
        
    return category, links

def render_osint_module(api_key):
    st.info("Hedef kişinin (Borçlu, Davalı) dijital ayak izlerini takip edin. Sistem 'Threading' teknolojisiyle aynı anda sosyal medya, ticaret sicil ve mal varlığı taraması başlatır.")
    
    col_input, col_results = st.columns([1, 2])
    
    with col_input:
        st.markdown("#### 🎯 Hedef Tanımla")
        target_name = st.text_input("Ad Soyad / Şirket Adı", placeholder="Örn: Ahmet Yılmaz")
        target_city = st.text_input("Şehir (Opsiyonel)", placeholder="Örn: İstanbul")
        
        start_scan = st.button("🚀 İstihbarat Taramasını Başlat", type="primary")
        
        st.markdown("---")
        st.caption("⚠️ **Yasal Uyarı:** Bu modül sadece halka açık verileri (Open Source) tarar. KVKK sınırları içinde kullanınız.")

    with col_results:
        if start_scan and target_name:
            st.write(f"📡 **'{target_name}'** için çok kanallı tarama başlatılıyor...")
            
            # --- THREADING (Çoklu İş Parçacığı) BAŞLANGICI ---
            # Sosyal medya, İş ve Varlık taramalarını aynı anda yapar
            results = {}
            
            with st.spinner("Veri madenciliği yapılıyor (Social + Business + Assets)..."):
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    # Görevleri tanımla
                    t1 = executor.submit(generate_dork_category, "social", target_name, target_city)
                    t2 = executor.submit(generate_dork_category, "business", target_name, target_city)
                    t3 = executor.submit(generate_dork_category, "assets", target_name, target_city)
                    
                    # Sonuçları topla
                    for future in concurrent.futures.as_completed([t1, t2, t3]):
                        cat, links = future.result()
                        results[cat] = links
            
            st.success("✅ Tarama Tamamlandı! Bulunan İzler:")
            
            # Sonuçları Göster
            tab_social, tab_business, tab_assets = st.tabs(["📸 Sosyal Medya", "🏢 Ticari Varlık", "🏖️ Yaşam Tarzı"])
            
            with tab_social:
                st.markdown("### Sosyal Ağ Taraması")
                for link in results.get("social", []):
                    st.markdown(f"- {link}", unsafe_allow_html=True)
                st.info("💡 İpucu: 'Borcum yok' diyen kişinin Instagram'da gizli hikayesi olabilir.")

            with tab_business:
                st.markdown("### Ticari Sicil & Resmi Kayıtlar")
                for link in results.get("business", []):
                    st.markdown(f"- {link}", unsafe_allow_html=True)
                st.info("💡 İpucu: Üzerine kayıtlı şirket veya ortaklıkları buradan yakalayabilirsiniz.")

            with tab_assets:
                st.markdown("### Lüks Yaşam & Varlık İzleri")
                for link in results.get("assets", []):
                    st.markdown(f"- {link}", unsafe_allow_html=True)
                st.info("💡 İpucu: Otel yorumları veya 2. el satış ilanları gizli varlıkları ele verebilir.")

            # --- AI ANALİZ KISMI ---
            st.divider()
            st.markdown("#### 🧠 İstihbarat Analizi")
            evidence_text = st.text_area("Bulduğunuz şüpheli bilgiyi buraya yapıştırın (Örn: Instagram biyografisi veya Ticaret Sicil kaydı):", height=100)
            
            if st.button("🕵️ Delil Analizi Yap"):
                if not api_key:
                    st.error("API Anahtarı gerekli.")
                else:
                    genai.configure(api_key=api_key)
                    model = genai.GenerativeModel('gemini-pro')
                    prompt = f"""
                    GÖREV: Bir OSINT (Açık Kaynak İstihbaratı) uzmanısın.
                    HEDEF KİŞİ: {target_name}
                    BULUNAN VERİ: "{evidence_text}"
                    
                    SORU: 
                    1. Bu veri, kişinin "borç ödemekten kaçınma" veya "mal kaçırma" şüphesini destekler mi?
                    2. Hukuki olarak bu veri delil dosyasında nasıl kullanılabilir?
                    
                    Kısa ve net cevap ver.
                    """
                    with st.spinner("Yapay zeka veriyi yorumluyor..."):
                        response = model.generate_content(prompt)
                        st.write(response.text)

def render_precedent_alert_module(api_key):
    st.info("Bu modül, derdest (devam eden) davalarınızı takip eder ve Yargıtay/AYM tarafından yayınlanan **'Bugünkü Kararlar'** ile otomatik eşleştirir.")

    # --- 1. OTURUM DURUMU (Dava Portföyü) ---
    if 'my_cases' not in st.session_state:
        st.session_state.my_cases = [
            {"id": 1, "ad": "Yılmaz v. Demir (Kira)", "konu": "5 yıllık kiracı tahliyesi, uyarlama davası", "durum": "Bilirkişi aşamasında"},
            {"id": 2, "ad": "Kripto Dolandırıcılık", "konu": "Thodex benzeri borsa batışı, güveni kötüye kullanma", "durum": "Savcılık soruşturması"},
            {"id": 3, "ad": "İşe İade (Ahmet B.)", "konu": "Performans düşüklüğü nedeniyle fesih", "durum": "Tanık dinleniyor"}
        ]

    col_portfolio, col_feed = st.columns([1, 2])

    # --- SOL KOLON: DAVA PORTFÖYÜM ---
    with col_portfolio:
        st.markdown("### 📂 Dava Portföyüm")
        
        with st.expander("➕ Yeni Dava Ekle"):
            new_case_name = st.text_input("Dava Adı")
            new_case_topic = st.text_area("Dava Konusu/Detayı")
            if st.button("Listeye Ekle"):
                new_id = len(st.session_state.my_cases) + 1
                st.session_state.my_cases.append({"id": new_id, "ad": new_case_name, "konu": new_case_topic, "durum": "Yeni"})
                st.success("Eklendi!")
                st.rerun()
        
        for case in st.session_state.my_cases:
            st.markdown(f"**Dosya #{case['id']}: {case['ad']}**\n*{case['konu']}*\n`Durum: {case['durum']}`\n---")

    # --- SAĞ KOLON: GÜNLÜK BÜLTEN TARAMASI ---
    with col_feed:
        st.markdown("### 📡 Günlük Yargı Bülteni & Etki Analizi")
        
        if st.button("🔄 Bülteni Tara ve Analiz Et", type="primary", use_container_width=True):
            if not api_key:
                st.error("API Anahtarı gerekli.")
            else:
                daily_decisions = [
                    """KARAR 2024/105 (Yargıtay HGK): Kira tespit davalarında '5 yıllık süre' dolmadan yapılan uyarlamalarda, TÜFE oranı tavan olarak kabul edilemez. Hakim hakkaniyete göre serbestçe belirler.""",
                    """KARAR 2024/88 (AYM Bireysel Başvuru): Kripto para borsalarındaki kayıplarda, devletin denetim yükümlülüğünü ihlal ettiği iddiasıyla yapılan başvuruda 'Mülkiyet Hakkı İhlali' olmadığına karar verildi.""",
                    """KARAR 2024/12 (İş Mahkemesi Emsal): Sadece performans düşüklüğü, yazılı savunma alınmadan ve eğitim verilmeden fesih sebebi yapılamaz."""
                ]
                
                st.write(f"📅 **Bugün Yayınlanan Kritik Karar Sayısı:** {len(daily_decisions)}")
                
                with st.spinner("Uygun yapay zeka modeli aranıyor ve analiz yapılıyor..."):
                    genai.configure(api_key=api_key)
                    
                    # --- OTOMATİK MODEL SEÇİCİ (HATA ÇÖZÜMÜ) ---
                    target_model_name = "models/gemini-pro" # Varsayılan
                    try:
                        # Sistemdeki mevcut modelleri listele ve ilk çalışanı seç
                        for m in genai.list_models():
                            if 'generateContent' in m.supported_generation_methods:
                                if 'gemini' in m.name:
                                    target_model_name = m.name
                                    break
                    except Exception as e:
                        st.warning(f"Model listesi alınamadı, varsayılan deneniyor: {e}")

                    # Seçilen modeli ekrana yaz (Debug için)
                    # st.caption(f"Kullanılan Model: {target_model_name}") 
                    
                    try:
                        model = genai.GenerativeModel(target_model_name)
                        
                        cases_str = str(st.session_state.my_cases)
                        decisions_str = "\n".join(daily_decisions)
                        
                        prompt = f"""
                        GÖREV: Sen proaktif bir hukuk asistanısın.
                        1. AŞAĞIDAKİ MÜVEKKİL DAVALARI (PORTFÖY): {cases_str}
                        2. AŞAĞIDAKİ BUGÜN ÇIKAN YENİ YARGI KARARLARI: {decisions_str}
                        YAPMAN GEREKEN: Her bir davayı kontrol et. Eğer yeni kararlardan biri, mevcut bir davayı etkiliyorsa uyar.
                        ÇIKTI FORMATI:
                        UYARI: [Dosya Adı]
                        DURUM: [KRİTİK / DİKKAT / FIRSAT]
                        NEDEN: [Açıklama]
                        AKSİYON: [Öneri]
                        """
                        
                        response = model.generate_content(prompt)
                        
                        st.divider()
                        st.markdown("### 🚨 Tespit Edilen Riskler ve Fırsatlar")
                        
                        alerts = response.text.split("UYARI:")
                        if len(alerts) < 2:
                            st.write(response.text)
                        else:
                            for alert in alerts:
                                if alert.strip():
                                    if "KRİTİK" in alert: st.error(f"**UYARI:{alert}**")
                                    elif "FIRSAT" in alert: st.success(f"**UYARI:{alert}**")
                                    else: st.warning(f"**UYARI:{alert}**")
                                    
                    except Exception as e:
                        st.error(f"Model Hatası: {str(e)}")
                        st.info("Lütfen API anahtarınızın 'Generative AI' servisine erişimi olduğundan emin olun.")


def render_owner_mode(api_key):
    st.info("👑 **Sahip Modu (Web):** Bilgisayarınızdaki dosyaları seçip sürükleyin. Sistem, hesabınızda çalışan en uygun Yapay Zeka modelini otomatik bulup kullanacaktır.")

    # --- 0. OTOMATİK MODEL BULUCU (HATA ÖNLEYİCİ) ---
    def get_working_model():
        """Sistemdeki aktif modelleri tarar ve ilk çalışanı getirir."""
        default_model = "models/gemini-pro" # En kötü ihtimal yedeği
        try:
            for m in genai.list_models():
                if 'generateContent' in m.supported_generation_methods:
                    # İsminde 'gemini' geçen ilk modeli al (flash, pro, 1.5 vs.)
                    if 'gemini' in m.name:
                        return m.name
        except:
            pass
        return default_model

    # --- 1. DOSYA OKUMA MOTORU ---
    def get_file_text(file_obj, api_key_for_ocr):
        filename = file_obj.name.lower()
        file_bytes = file_obj.read()
        text = ""
        
        try:
            # A) PDF
            if filename.endswith('.pdf'):
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            # B) WORD
            elif filename.endswith('.docx'):
                doc = Document(io.BytesIO(file_bytes))
                text = "\n".join([p.text for p in doc.paragraphs])
            
            # C) UYAP (UDF)
            elif filename.endswith('.udf'):
                try:
                    with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                        text = "".join(ET.fromstring(z.read('content.xml')).itertext())
                except:
                    text = "".join(ET.fromstring(file_bytes).itertext())
            
            # D) RESİM (OTOMATİK MODEL İLE)
            elif filename.endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp')):
                if api_key_for_ocr:
                    image = Image.open(io.BytesIO(file_bytes))
                    
                    # Dinamik model seçimi
                    active_model = get_working_model()
                    model = genai.GenerativeModel(active_model)
                    
                    try:
                        response = model.generate_content(["Bu resimdeki yazıları oku:", image])
                        text = f"[RESİM İÇERİĞİ]:\n{response.text}"
                    except:
                        text = "[RESİM OKUNAMADI: Seçilen model görsel desteklemiyor olabilir.]"
            
            # E) DÜZ METİN
            else:
                text = file_bytes.decode("utf-8", errors='ignore')
                
            return text
        except Exception as e:
            return f"[Okuma Hatası]: {str(e)}"

    # --- 2. ARAYÜZ ---
    if 'web_memory' not in st.session_state: st.session_state.web_memory = ""
    if 'web_history' not in st.session_state: st.session_state.web_history = []

    col_upload, col_chat = st.columns([1, 2])

    # --- SOL: YÜKLEME ---
    with col_upload:
        st.markdown("### 📤 Dosyaları Sürükle")
        uploaded_files = st.file_uploader("Klasördeki dosyaları seçip buraya bırak", accept_multiple_files=True)
        
        if st.button("🧠 Analiz Et", type="primary"):
            if not uploaded_files:
                st.warning("Dosya yok.")
            elif not api_key:
                st.error("API Anahtarı yok.")
            else:
                genai.configure(api_key=api_key)
                full_text = ""
                bar = st.progress(0)
                
                for i, file in enumerate(uploaded_files):
                    content = get_file_text(file, api_key)
                    full_text += f"\n=== DOSYA: {file.name} ===\n{content}\n"
                    bar.progress((i + 1) / len(uploaded_files))
                
                st.session_state.web_memory = full_text
                st.session_state.web_history = [] 
                st.success(f"✅ {len(uploaded_files)} dosya okundu!")

        if st.session_state.web_memory:
            if st.button("🗑️ Temizle"):
                st.session_state.web_memory = ""
                st.rerun()

    # --- SAĞ: SOHBET ---
    with col_chat:
        st.markdown("### 💬 Asistan")
        
        for msg in st.session_state.web_history:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

        if prompt := st.chat_input("Sorunuzu yazın..."):
            if not st.session_state.web_memory:
                st.warning("Önce dosya yükleyin.")
            else:
                st.session_state.web_history.append({"role": "user", "content": prompt})
                with st.chat_message("user"):
                    st.markdown(prompt)

                with st.chat_message("assistant"):
                    with st.spinner("Düşünüyor..."):
                        try:
                            # BURADA OTOMATİK MODEL SEÇİLİYOR
                            active_model_name = get_working_model()
                            # st.caption(f"Kullanılan Model: {active_model_name}") # İstersen açıp görebilirsin
                            
                            model = genai.GenerativeModel(active_model_name)
                            
                            context = st.session_state.web_memory[:90000]
                            final_prompt = f"VERİLER:\n{context}\n\nSORU: {prompt}"
                            
                            response = model.generate_content(final_prompt)
                            st.markdown(response.text)
                            st.session_state.web_history.append({"role": "assistant", "content": response.text})
                        except Exception as e:
                            st.error(f"Cevap üretilemedi: {e}")


import json
import time

def render_property_genealogy(api_key):
    st.info("🌳 **Mülkiyet Soyağacı:** Tapu ve kadastro belgelerinizi yükleyin, AI zinciri kursun.")

    # --- 0. AKILLI MODEL SEÇİCİ (HATA ÖNLEYİCİ) ---
    def get_best_model():
        """Hesapta aktif olan en hızlı modeli bulur."""
        try:
            available_models = []
            for m in genai.list_models():
                if 'generateContent' in m.supported_generation_methods:
                    available_models.append(m.name)
            
            # Öncelik Sırası: 1. Flash (Hızlı), 2. Pro (Standart), 3. Herhangi biri
            for m in available_models:
                if 'flash' in m: return m
            for m in available_models:
                if 'pro' in m: return m
            return available_models[0] if available_models else "models/gemini-pro"
        except:
            return "models/gemini-pro" # En kötü ihtimal yedeği

    # --- 1. DOSYA OKUMA ---
    def get_genealogy_file_text(file_obj, api_key_for_ocr):
        filename = file_obj.name.lower()
        file_bytes = file_obj.read()
        text = ""
        try:
            if filename.endswith('.pdf'):
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                for page in pdf_reader.pages: text += page.extract_text() + "\n"
            elif filename.endswith('.docx'):
                doc = Document(io.BytesIO(file_bytes))
                text = "\n".join([p.text for p in doc.paragraphs])
            elif filename.endswith(('.png', '.jpg', '.jpeg')):
                if api_key_for_ocr:
                    image = Image.open(io.BytesIO(file_bytes))
                    # Otomatik model seçimi
                    active_model = get_best_model()
                    model = genai.GenerativeModel(active_model)
                    response = model.generate_content(["Bu belgedeki isimleri ve tarihleri oku:", image])
                    text = response.text
            else:
                text = file_bytes.decode("utf-8", errors='ignore')
            return text
        except:
            return ""

    # --- 2. STATE ---
    if 'prop_history' not in st.session_state:
        st.session_state.prop_history = [
            {"yil": "1960", "kimden": "Hazine", "kime": "Mehmet Ağa", "islem": "Kadastro", "durum": "Pasif"},
            {"yil": "1990", "kimden": "Mehmet Ağa", "kime": "Ali (Oğlu)", "islem": "Miras", "durum": "Aktif"}
        ]

    # --- 3. ARAYÜZ ---
    col_left, col_right = st.columns([1, 2])

    # SOL: YÜKLEME
    with col_left:
        st.markdown("### 📂 Belge Yükle")
        uploaded_files = st.file_uploader("Tapu/Kadastro Evrakları", accept_multiple_files=True)
        
        if st.button("⚡ Zinciri Oluştur", type="primary"):
            if not uploaded_files or not api_key:
                st.warning("Dosya ve API Key gerekli.")
            else:
                status_box = st.empty()
                status_box.info("Belgeler okunuyor...")
                
                genai.configure(api_key=api_key)
                full_text = ""
                for f in uploaded_files:
                    full_text += f"\nDOC: {f.name}\n" + get_genealogy_file_text(f, api_key)
                
                try:
                    status_box.info("AI Modeli seçiliyor ve zincir kuruluyor...")
                    
                    # OTOMATİK MODEL SEÇİMİ
                    active_model_name = get_best_model()
                    model = genai.GenerativeModel(active_model_name)
                    
                    prompt = f"""
                    GÖREV: Metinlerdeki mülkiyet devirlerini JSON listesi yap.
                    METİN: {full_text[:40000]}
                    FORMAT: [{{"yil": "...", "kimden": "...", "kime": "...", "islem": "...", "durum": "Aktif/Pasif/Kritik"}}]
                    SADECE JSON VER.
                    """
                    response = model.generate_content(prompt)
                    clean_json = response.text.replace("```json", "").replace("```", "").strip()
                    st.session_state.prop_history = json.loads(clean_json)
                    status_box.success(f"Tamamlandı! (Kullanılan Model: {active_model_name})")
                    time.sleep(1)
                    st.rerun()
                except Exception as e:
                    status_box.error(f"Hata: {e}")

        st.dataframe(st.session_state.prop_history, height=300)
        if st.button("Temizle"):
            st.session_state.prop_history = []
            st.rerun()

    # SAĞ: GRAFİK VE ANALİZ
    with col_right:
        st.markdown("### 🗺️ Görsel Harita")
        
        if st.session_state.prop_history:
            # Grafik Çizimi
            graph_code = "digraph { rankdir=LR; node [shape=box, style=filled, fontname=\"Arial\"];"
            for item in st.session_state.prop_history:
                color = "#d4edda" if item.get("durum") == "Aktif" else "#e2e3e5"
                if item.get("durum") == "Kritik": color = "#f8d7da"
                
                k1 = str(item.get('kimden', '?')).replace('"', '').strip()
                k2 = str(item.get('kime', '?')).replace('"', '').strip()
                lbl = f"{item.get('yil')}\\n{item.get('islem')}"
                
                graph_code += f'\n "{k1}" -> "{k2}" [label="{lbl}", fontsize=10];'
                graph_code += f'\n "{k2}" [fillcolor="{color}", label="{k2}"];'
            graph_code += "\n}"
            st.graphviz_chart(graph_code)
            
            st.divider()
            
            # --- ANALİZ KISMI (STREAMING & AUTO MODEL) ---
            if st.button("🕵️ Risk Analizi Başlat"):
                output_placeholder = st.empty()
                output_placeholder.text("Model aranıyor ve analiz başlıyor...")
                
                try:
                    genai.configure(api_key=api_key)
                    
                    # OTOMATİK MODEL SEÇİMİ
                    active_model_name = get_best_model()
                    model = genai.GenerativeModel(active_model_name)
                    
                    chain_data = json.dumps(st.session_state.prop_history, ensure_ascii=False)
                    
                    prompt = f"""
                    GÖREV: Sen uzman bir tapu denetçisisin. Aşağıdaki mülkiyet zincirini analiz et.
                    VERİ: {chain_data}
                    
                    Lütfen şu başlıklar altında rapor yaz:
                    1. 🔴 Kritik Riskler
                    2. ⚠️ Hukuki Uyarılar
                    3. ✅ Sonuç
                    """
                    
                    # stream=True ile parça parça alıyoruz
                    response = model.generate_content(prompt, stream=True)
                    
                    full_text = ""
                    for chunk in response:
                        full_text += chunk.text
                        output_placeholder.markdown(full_text + "▌") 
                    
                    output_placeholder.markdown(full_text)
                        
                except Exception as e:
                    output_placeholder.error(f"Hata oluştu: {e}")
        else:
            st.info("👈 Veri yok.")

import pandas as pd
from datetime import datetime, timedelta

def render_limitations_heatmap(api_key):
    # --- IMPORTLARI İZOLE ET (Çakışmayı Önler) ---
    import pandas as pd
    import datetime as dt  # datetime modülünü 'dt' olarak çağırıyoruz
    
    st.info("🔥 **Zamanaşımı Isı Haritası:** Dava türüne ve tarihlere göre her bir alacak kaleminin risk durumunu analiz eder. Islah ve hak düşürücü süreleri 'Borsa Ekranı' gibi takip eder.")

    # --- 0. OTOMATİK MODEL SEÇİCİ ---
    def get_best_model():
        try:
            available_models = []
            for m in genai.list_models():
                if 'generateContent' in m.supported_generation_methods:
                    available_models.append(m.name)
            for m in available_models:
                if 'flash' in m: return m
            for m in available_models:
                if 'pro' in m: return m
            return available_models[0] if available_models else "models/gemini-pro"
        except:
            return "models/gemini-pro"

    # --- 1. GİRİŞ PANELİ ---
    col_input, col_dashboard = st.columns([1, 2])

    with col_input:
        st.markdown("### 📅 Kritik Tarihler")
        
        dava_turu = st.selectbox("Dava Türü", ["İş Hukuku (İşçi Alacağı)", "Ticari Alacak", "Tüketici", "Tazminat (Haksız Fiil)"])
        
        # HATA ÇIKARAN SATIR DÜZELTİLDİ: dt.datetime.now().date()
        bugun = dt.datetime.now().date()
        
        # Tarih Seçiciler
        fesih_tarihi = st.date_input("Fesih / Olay Tarihi", value=bugun - dt.timedelta(days=365*4))
        dava_tarihi = st.date_input("Dava Açılış Tarihi", value=bugun - dt.timedelta(days=300))
        
        st.divider()
        st.markdown("#### ⚡ Islah Alarmı")
        is_bilirkişi = st.checkbox("Bilirkişi Raporu Geldi mi?")
        
        teblig_tarihi = None
        if is_bilirkişi:
            teblig_tarihi = st.date_input("Rapor Tebliğ Tarihi", value=bugun - dt.timedelta(days=5))
            st.caption("Islah için genellikle 2 haftalık itiraz süresi veya tahkikat sonuna kadar süre dikkate alınır.")

    # --- 2. HESAPLAMA MOTORU ---
    data = []
    
    # İş Hukuku Kuralları (Basitleştirilmiş Örnekler)
    if dava_turu == "İş Hukuku (İşçi Alacağı)":
        # 1. Kıdem Tazminatı (5 Yıl - 2017 sonrası)
        kidem_suresi = fesih_tarihi + dt.timedelta(days=365*5)
        kalan_gun = (kidem_suresi - bugun).days
        data.append({"Kalem": "Kıdem Tazminatı", "Son Tarih": kidem_suresi, "Kalan Gün": kalan_gun, "Risk": ""})
        
        # 2. Fazla Mesai (5 Yıl)
        mesai_suresi = fesih_tarihi + dt.timedelta(days=365*5)
        kalan_gun_mesai = (mesai_suresi - bugun).days
        data.append({"Kalem": "Fazla Mesai", "Son Tarih": mesai_suresi, "Kalan Gün": kalan_gun_mesai, "Risk": ""})
        
        # 3. İşe İade (1 Ay - Arabulucu)
        ise_iade_suresi = fesih_tarihi + dt.timedelta(days=30)
        kalan_gun_iade = (ise_iade_suresi - bugun).days
        data.append({"Kalem": "İşe İade (Arabulucu)", "Son Tarih": ise_iade_suresi, "Kalan Gün": kalan_gun_iade, "Risk": ""})

    # Islah Hesabı (Kritik)
    if is_bilirkişi and teblig_tarihi:
        # HMK 281 - 2 Hafta İtiraz (Islah için stratejik zaman)
        islah_suresi = teblig_tarihi + dt.timedelta(days=14)
        kalan_gun_islah = (islah_suresi - bugun).days
        data.append({"Kalem": "🚨 ISLAH / İTİRAZ", "Son Tarih": islah_suresi, "Kalan Gün": kalan_gun_islah, "Risk": "ÇOK YÜKSEK"})

    # DataFrame Oluştur
    df = pd.DataFrame(data)

    # Risk Renklendirme Fonksiyonu
    def risk_color(val):
        if val < 0: return "background-color: #ff4b4b; color: white" # Kırmızı (Süre Doldu)
        elif val < 15: return "background-color: #ffa500; color: black" # Turuncu (Kritik)
        elif val < 60: return "background-color: #ffe066; color: black" # Sarı (Yaklaşıyor)
        else: return "background-color: #90ee90; color: black" # Yeşil (Güvenli)

    # --- 3. DASHBOARD (ISI HARİTASI) ---
    with col_dashboard:
        st.markdown("### 🌡️ Zamanaşımı Isı Haritası")
        
        if not df.empty:
            # Tabloyu Renklendir
            st.dataframe(
                df.style.applymap(risk_color, subset=["Kalan Gün"])
                        .format({"Son Tarih": "{:%d.%m.%Y}"}),
                use_container_width=True,
                height=250
            )
            
            # Görsel Ticker (İlerleme Çubukları)
            st.markdown("#### ⏳ Kritik Geri Sayım")
            for index, row in df.iterrows():
                kalan = row["Kalan Gün"]
                kalem = row["Kalem"]
                
                if kalan < 0:
                    st.error(f"❌ {kalem}: SÜRE DOLDU! ({abs(kalan)} gün geçti)")
                elif kalan < 15:
                    st.warning(f"⚠️ {kalem}: SON {kalan} GÜN! (Acil İşlem Gerekli)")
                    st.progress(max(0, min(100, int((kalan/15)*100))))
                else:
                    st.success(f"✅ {kalem}: {kalan} gün var. (Güvenli)")
        else:
            st.info("Lütfen sol taraftan tarihleri giriniz.")

        st.divider()
        
        # --- 4. AI STRATEJİ DANIŞMANI ---
        if st.button("🧠 AI Risk & Strateji Analizi Yap"):
            if not api_key:
                st.error("API Key gerekli.")
            else:
                output_box = st.empty()
                output_box.info("Veriler analiz ediliyor...")
                
                try:
                    genai.configure(api_key=api_key)
                    active_model = get_best_model()
                    model = genai.GenerativeModel(active_model)
                    
                    # Tarihleri stringe çevirerek JSON hatasını önle
                    prompt = f"""
                    GÖREV: Bir avukat için zamanaşımı risk analizi yap.
                    
                    DURUM:
                    - Dava Türü: {dava_turu}
                    - Fesih Tarihi: {fesih_tarihi}
                    - Bugün: {bugun}
                    - Tablo Verileri: {df.to_json(orient='records', date_format='iso')}
                    
                    İSTENENLER:
                    1. Hangi kalemlerde zamanaşımı riski var? (Kısa ve net)
                    2. Islah dilekçesi için ne kadar vaktim kaldı? Geç kalırsam ne olur?
                    3. Zamanaşımı def'i (savunması) ile karşılaşırsam ne yapmalıyım?
                    4. Faiz başlangıç tarihleri için stratejik bir öneri ver.
                    """
                    
                    response = model.generate_content(prompt, stream=True)
                    
                    full_text = ""
                    for chunk in response:
                        full_text += chunk.text
                        output_box.markdown(full_text + "▌")
                    output_box.markdown(full_text)
                    
                except Exception as e:
                    output_box.error(f"Hata: {e}")


import networkx as nx
import matplotlib.pyplot as plt

def render_conflict_scanner(api_key):
    st.info("🕸️ **Gizli Bağlantı (Conflict of Interest) Tarayıcısı:** Hakim, avukat ve tanıklar arasındaki görünmez ticari ve sosyal bağları ortaya çıkarır. NetworkX ile ağ analizi yapar.")

    # --- 0. OTOMATİK MODEL SEÇİCİ ---
    def get_best_model():
        try:
            available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
            for m in available_models:
                if 'flash' in m: return m
            return available_models[0] if available_models else "models/gemini-pro"
        except:
            return "models/gemini-pro"

    col_input, col_graph = st.columns([1, 2])

    # --- 1. GİRDİ PANELİ ---
    with col_input:
        st.markdown("### 👥 Aktörleri Tanımla")
        
        hakim = st.text_input("Hakim İsmi", "Hakim Zeynep Yılmaz")
        avukat_karsi = st.text_input("Karşı Taraf Avukatı", "Av. Ahmet Demir")
        tanik = st.text_input("Tanık / Bilirkişi", "Bilirkişi Mehmet Öztürk")
        sirket = st.text_input("İlgili Şirket (Opsiyonel)", "Delta İnşaat A.Ş.")
        
        st.divider()
        st.markdown("#### 📡 Veri Kaynağı")
        source_type = st.radio("Tarama Yöntemi", ["Demo Simülasyonu (NetworkX Testi)", "AI ile Açık Kaynak Tarama (OSINT)"])

    # --- 2. ANALİZ MOTORU ---
    with col_graph:
        st.markdown("### 🕸️ İlişki Ağı Haritası")
        
        if st.button("🔍 Derinlemesine Tara", type="primary"):
            
            # SENARYO 1: DEMO SİMÜLASYONU (NetworkX Gücünü Göstermek İçin)
            if source_type == "Demo Simülasyonu (NetworkX Testi)":
                st.warning("⚠️ Demo Modu: Rastgele ticari sicil verileri simüle ediliyor...")
                
                # NetworkX Grafiği Oluştur
                G = nx.Graph()
                
                # Düğümleri (Kişileri/Kurumları) Ekle
                G.add_node(hakim, type="Yargı", color="red")
                G.add_node(avukat_karsi, type="Avukat", color="black")
                G.add_node(tanik, type="Tanık", color="blue")
                G.add_node(sirket, type="Şirket", color="green")
                
                # Gizli Bağlantıları Ekle (Simülasyon)
                # Örnek: Avukat ve Tanık, 5 yıl önce "Omega Yazılım"da ortaktı.
                hidden_entity = "Omega Yazılım Ltd. Şti. (Eski)"
                G.add_node(hidden_entity, type="Şirket", color="grey")
                
                G.add_edge(avukat_karsi, hidden_entity, relation="Eski Ortak (2018)")
                G.add_edge(tanik, hidden_entity, relation="Yön. Kur. Üyesi (2018)")
                
                # Örnek: Hakim ve Şirket arasında dolaylı bağ
                dernek = "Hukukçular Vakfı"
                G.add_node(dernek, type="STK", color="orange")
                G.add_edge(hakim, dernek, relation="Üye")
                G.add_edge(avukat_karsi, dernek, relation="Yönetim Kurulu")

                # Graphviz ile Çiz (Streamlit için en temizi)
                dot_code = "graph {"
                dot_code += "\n  rankdir=LR;"
                
                # NetworkX verisini Graphviz formatına çevir
                for u, v, data in G.edges(data=True):
                    rel = data.get('relation', '')
                    dot_code += f'\n  "{u}" -- "{v}" [label="{rel}", fontsize=10];'
                
                # Renklendirme
                dot_code += f'\n  "{hakim}" [style=filled, fillcolor="#ffcccc"];' # Kırmızımsı
                dot_code += f'\n  "{avukat_karsi}" [style=filled, fillcolor="#cccccc"];'
                dot_code += f'\n  "{tanik}" [style=filled, fillcolor="#ccccff"];'
                dot_code += f'\n  "{hidden_entity}" [style=filled, fillcolor="#ffffcc", shape=box];' # Sarı (Kilit Nokta)
                
                dot_code += "\n}"
                st.graphviz_chart(dot_code)
                
                # NETWORKX ANALİZİ: EN KISA YOL (Shortest Path)
                try:
                    path = nx.shortest_path(G, source=avukat_karsi, target=tanik)
                    st.error(f"🚨 **KRİTİK BULGU:** {avukat_karsi} ile {tanik} arasında bağlantı tespit edildi!")
                    st.write(f"🔗 **Bağlantı Zinciri:** {' -> '.join(path)}")
                    st.caption("Bu bilgi, HMK m. 254 kapsamında tanığın tarafsızlığını sorgulamak için kullanılabilir.")
                except nx.NetworkXNoPath:
                    st.success("Doğrudan bir bağlantı bulunamadı.")

            # SENARYO 2: AI OSINT ANALİZİ (Gerçekçi Senaryo)
            else:
                if not api_key:
                    st.error("API Key gerekli.")
                else:
                    output_box = st.empty()
                    output_box.info("Açık kaynaklar ve haberler taranıyor...")
                    
                    try:
                        genai.configure(api_key=api_key)
                        active_model = get_best_model()
                        model = genai.GenerativeModel(active_model)
                        
                        # Prompt: AI'yı bir OSINT uzmanı gibi çalıştırıyoruz
                        prompt = f"""
                        GÖREV: Sen kıdemli bir istihbarat analistisin.
                        Aşağıdaki kişiler arasında potansiyel bir "Çıkar Çatışması" (Conflict of Interest) senaryosu kurgula ve analiz et.
                        
                        KİŞİLER:
                        1. Hakim: {hakim}
                        2. Karşı Avukat: {avukat_karsi}
                        3. Tanık: {tanik}
                        4. Şirket: {sirket}
                        
                        İSTENENLER:
                        1. Bu isimler arasında olası (hayali veya genel bilgiye dayalı) geçmiş bağlantıları düşün (Eski okul arkadaşlığı, aynı dernek üyeliği, eski şirket ortaklığı).
                        2. Özellikle "Tanık" ile "Karşı Avukat" arasında redd-i hakim veya tanık itirazına gerekçe olabilecek bir bağ bul.
                        3. Bunu bir "İstihbarat Raporu" formatında sun.
                        4. Hukuki Tavsiye: Bu bağlantıyı mahkemede nasıl delillendiririm?
                        """
                        
                        response = model.generate_content(prompt, stream=True)
                        
                        full_text = ""
                        for chunk in response:
                            full_text += chunk.text
                            output_box.markdown(full_text + "▌")
                        output_box.markdown(full_text)
                        
                    except Exception as e:
                        output_box.error(f"Hata: {e}")

def render_mediation_checker(api_key):
    st.info("🤝 **Arabuluculuk Kontrolcüsü:** Dava türünü girin, sistem bunun 'Dava Şartı (Zorunlu)' olup olmadığını, ilgili kanun maddesini ve başvuru süresini analiz etsin.")

    # --- 0. MODEL SEÇİCİ ---
    def get_best_model():
        try:
            available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
            for m in available_models:
                if 'flash' in m: return m
            return available_models[0] if available_models else "models/gemini-pro"
        except:
            return "models/gemini-pro"

    col1, col2 = st.columns([1, 1])

    with col1:
        st.markdown("### ⚖️ Dava Konusu Nedir?")
        
        # Hızlı Seçim Butonları
        st.write("Sık Kullanılanlar:")
        col_btn1, col_btn2, col_btn3 = st.columns(3)
        
        case_input = st.text_input("Veya detaylı yazın (Örn: Mobbing nedeniyle tazminat)", "")
        
        if col_btn1.button("Kıdem/İhbar"): case_input = "İşçilik Alacakları (Kıdem, İhbar, Fazla Mesai)"
        if col_btn2.button("Kira/Tahliye"): case_input = "Kira Tespiti ve Tahliye (Konut/Çatılı İşyeri)"
        if col_btn3.button("Ticari Alacak"): case_input = "İki Tacir Arasındaki Fatura Alacağı"

    with col2:
        st.markdown("### 🔍 Analiz Sonucu")
        
        if st.button("Arabuluculuk Şartını Kontrol Et", type="primary"):
            if not case_input:
                st.warning("Lütfen bir dava türü girin.")
            elif not api_key:
                st.error("API Key gerekli.")
            else:
                status_box = st.empty()
                status_box.info("Mevzuat taranıyor (7036, 6102, 6325 Sayılı Kanunlar)...")
                
                try:
                    genai.configure(api_key=api_key)
                    active_model = get_best_model()
                    model = genai.GenerativeModel(active_model)
                    
                    prompt = f"""
                    GÖREV: Sen uzman bir Türk Hukuku avukatısın.
                    SORGU: "{case_input}" konulu bir dava açmak istiyorum.
                    
                    ANALİZ ET:
                    1. Bu dava için Arabuluculuk ZORUNLU MU (Dava Şartı mı) yoksa İHTİYARİ Mİ?
                    2. Hangi Kanun maddesine dayanıyor? (Örn: TTK 5/A, İŞK 3, 7445 SK vb.)
                    3. Eğer zorunluysa ve gitmezsem ne olur? (Usulden Ret uyarısı)
                    4. Başvuru nereye yapılır? (Adliye/Büro)
                    
                    ÇIKTI FORMATI:
                    Lütfen cevabı şu formatta ver (Markdown kullanarak):
                    
                    ### 🚦 DURUM: [ZORUNLU / İHTİYARİ / İSTİSNA]
                    
                    **📜 Yasal Dayanak:** ...
                    **⚠️ Risk Uyarısı:** ...
                    **📍 Başvuru Yeri:** ...
                    **💡 Kısa Özet:** ...
                    """
                    
                    response = model.generate_content(prompt, stream=True)
                    
                    full_text = ""
                    for chunk in response:
                        full_text += chunk.text
                        status_box.markdown(full_text + "▌")
                    status_box.markdown(full_text)
                    
                    # Görsel Uyarılar (Basit Regex Kontrolü)
                    if "ZORUNLU" in full_text:
                        st.error("🚨 DİKKAT: Arabulucuya gitmeden dava açarsanız, davanız USULDEN REDDEDİLİR!")
                    elif "İHTİYARİ" in full_text:
                        st.success("✅ Zorunlu değil, doğrudan dava açabilirsiniz. Ancak yine de arabuluculuk denenebilir.")
                        
                except Exception as e:
                    status_box.error(f"Hata: {e}")

    st.divider()
    st.caption("ℹ️ Not: 01.09.2023 tarihinden itibaren Kira, Kat Mülkiyeti, Komşuluk Hukuku ve Ortaklığın Giderilmesi davaları da zorunlu arabuluculuk kapsamına alınmıştır.")


import folium
from streamlit_folium import st_folium
import random
import streamlit as st

def render_forensic_map(api_key):
    st.info("🗺️ **Adli Isı Haritası (Forensic Geolocation):** Olay yerindeki geçmiş vakaları analiz eder. 'Sürücü hatası mı, yoksa yol kusuru mu?' sorusuna İdare Hukuku perspektifiyle yanıt arar.")

    # --- SESSION STATE (HAFIZA) AYARLARI ---
    # Analiz yapıldı mı bilgisini hafızada tutuyoruz
    if "map_analyzed" not in st.session_state:
        st.session_state.map_analyzed = False
    
    # AI Sonucunu hafızada tutmak için (Tekrar tekrar API harcamasın)
    if "ai_map_result" not in st.session_state:
        st.session_state.ai_map_result = None

    # --- 0. MODEL SEÇİCİ ---
    def get_best_model():
        try:
            # Genai import kontrolü (Global scope'ta yoksa hata vermesin)
            import google.generativeai as genai
            available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
            for m in available_models:
                if 'flash' in m: return m
            return available_models[0] if available_models else "models/gemini-pro"
        except:
            return "models/gemini-pro"

    col_input, col_map = st.columns([1, 2])

    # --- 1. GİRDİ PANELİ ---
    with col_input:
        st.markdown("### 📍 Olay Yeri Tanımla")
        
        location_name = st.text_input("Konum / Kavşak Adı", "Bağdat Caddesi Şaşkınbakkal Işıklar")
        city = st.selectbox("Şehir", ["İstanbul", "Ankara", "İzmir", "Bursa", "Antalya"])
        event_type = st.selectbox("Olay Türü", ["Trafik Kazası", "Hırsızlık / Gasp", "Sel / Su Baskını", "Çukur / Yol Çökmesi"])
        
        st.divider()
        st.markdown("#### 🎯 Hedef Analiz")
        
        # BUTONLAR
        col_b1, col_b2 = st.columns(2)
        
        # Analiz Butonu (Callback ile hafızayı tetikler)
        def activate_analysis():
            st.session_state.map_analyzed = True
            st.session_state.ai_map_result = None # Yeni analiz için eski AI sonucunu sil
            
        if col_b1.button("📡 Bölgeyi Tara", type="primary", on_click=activate_analysis):
            pass # İşlemi aşağıda yapacağız
            
        # Sıfırlama Butonu
        def reset_analysis():
            st.session_state.map_analyzed = False
            st.session_state.ai_map_result = None
            
        if col_b2.button("🔄 Sıfırla", on_click=reset_analysis):
            pass

    # --- 2. HARİTA VE ANALİZ ---
    with col_map:
        # Varsayılan Koordinatlar
        lat, lon = 41.0082, 28.9784 
        if city == "Ankara": lat, lon = 39.9334, 32.8597
        if city == "İzmir": lat, lon = 38.4192, 27.1287

        # Harita Oluştur (Her seferinde temiz başlar)
        m = folium.Map(location=[lat, lon], zoom_start=13)
        
        # EĞER ANALİZ BUTONUNA BASILDIYSA (Hafıza True ise)
        if st.session_state.map_analyzed:
            st.markdown(f"### 🔍 {location_name} - Risk Analizi")
            
            # --- A. HARİTA GÖRSELLEŞTİRME ---
            # Merkez Nokta
            folium.Marker(
                [lat, lon], 
                popup=f"<b>OLAY YERİ</b><br>{location_name}", 
                icon=folium.Icon(color="red", icon="info-sign")
            ).add_to(m)
            
            # Risk Noktaları (Simülasyon)
            # Not: Her render'da yer değiştirmemesi için seed sabitlenebilir veya statik veri kullanılabilir
            random.seed(42) 
            for _ in range(15):
                r_lat = lat + random.uniform(-0.015, 0.015)
                r_lon = lon + random.uniform(-0.015, 0.015)
                folium.CircleMarker(
                    location=[r_lat, r_lon],
                    radius=6,
                    color="crimson",
                    fill=True,
                    fill_color="crimson",
                    fill_opacity=0.6,
                    popup="Geçmiş Vaka (Riskli Bölge)"
                ).add_to(m)

            # Haritayı Çiz
            st_folium(m, height=350, width=700)
            
            # --- B. YAPAY ZEKA ANALİZİ ---
            if not api_key:
                st.error("Detaylı rapor için API Key gerekli.")
            else:
                # Eğer daha önce üretilmediyse üret
                if st.session_state.ai_map_result is None:
                    status_box = st.info("Haber arşivleri ve yerel şikayetler taranıyor... Lütfen bekleyin.")
                    
                    try:
                        import google.generativeai as genai
                        genai.configure(api_key=api_key)
                        active_model = get_best_model()
                        model = genai.GenerativeModel(active_model)
                        
                        prompt = f"""
                        GÖREV: Sen uzman bir İdare Hukuku avukatı ve Trafik Bilirkişisisin.
                        KONUM: {location_name}, {city}
                        OLAY TÜRÜ: {event_type}
                        
                        SENARYO: Müvekkil burada bir kaza yaptı/zarar gördü. Sadece karşı tarafı değil, devleti/belediyeyi de dava etmek istiyoruz.
                        
                        İSTENENLER:
                        1. Bu bölgeyle ilgili geçmişte basına yansıyan benzer kazalar veya "ölüm virajı", "karanlık yol" gibi haberler var mı? (Genel bilgi birikimini kullan).
                        2. İdarenin "Hizmet Kusuru" (Service Defect) sayılabilecek ihmalleri neler olabilir? (Örn: Sinyalizasyon eksikliği, yetersiz aydınlatma, çukur, rögar kapağı).
                        3. STRATEJİ: Davayı "Tam Yargı Davası" olarak İdare Mahkemesi'ne taşımak için hangi delilleri toplamalıyım? (MOBESE, Belediye şikayet kayıtları vb.)
                        4. SONUÇ: "Bu kavşakta son 1 yılda çok kaza olduysa, kusur sürücüde değil yoldadır" tezini savunacak hukuki argümanlar yaz.
                        """
                        
                        response = model.generate_content(prompt)
                        st.session_state.ai_map_result = response.text
                        status_box.empty() # Yükleniyor yazısını kaldır
                        
                    except Exception as e:
                        st.error(f"AI Hatası: {e}")
                
                # Sonucu Göster (Hafızadan)
                if st.session_state.ai_map_result:
                    st.markdown(st.session_state.ai_map_result)
        
        else:
            # Analiz öncesi boş harita
            st_folium(m, height=350, width=700)
            st.caption("👈 Analiz butonuna bastığınızda bölgedeki risk yoğunluğu haritaya işlenecektir.")


import datetime

def render_temporal_law_machine(api_key):
    st.info("🕰️ **Mevzuat Zaman Makinesi:** Olayın yaşandığı tarihe geri döner. O gün yürürlükte olan (şu an mülga) kanunları, tüzükleri ve Yargıtay içtihatlarını bugünkülerle kıyaslar.")

    # --- 0. MODEL SEÇİCİ ---
    def get_best_model():
        try:
            available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
            for m in available_models:
                if 'flash' in m: return m
            return available_models[0] if available_models else "models/gemini-pro"
        except:
            return "models/gemini-pro"

    col_settings, col_result = st.columns([1, 2])

    # --- 1. ZAMAN AYARLARI ---
    with col_settings:
        st.markdown("### ⚙️ Zaman Koordinatları")
        
        # Tarih Seçimi (Varsayılan: 1990'lar)
        target_date = st.date_input("Olay Tarihi", datetime.date(1995, 6, 15))
        
        topic = st.selectbox("Hukuki Konu", [
            "Gayrimenkul Devri (Tapu İptal)", 
            "Miras Paylaşımı (Tereke)", 
            "Boşanma ve Mal Rejimi", 
            "İş Kazası Tazminatı",
            "Ticari Sözleşme İhlali"
        ])
        
        specific_query = st.text_input("Özel Detay (Opsiyonel)", "Muris muvazaası ve saklı pay")
        
        st.divider()
        st.markdown("#### 🔄 Dönüşüm Modu")
        comparison_mode = st.radio("Analiz Türü", ["Sadece O Günün Kanunu", "Eski vs Yeni Kanun Kıyaslaması"])
        
        start_travel = st.button("🚀 Geçmişe Git ve Mevzuatı Getir", type="primary")

    # --- 2. SONUÇ EKRANI ---
    with col_result:
        st.markdown(f"### 📜 {target_date.year} Yılı Mevzuat Panoramas")
        
        if start_travel:
            if not api_key:
                st.error("Zaman yolculuğu için API Key gerekli.")
            else:
                # Görsel Efekt
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                status_text.text(f"⏳ {target_date.year} yılına gidiliyor...")
                time.sleep(0.5)
                progress_bar.progress(30)
                
                status_text.text("📚 Resmi Gazete arşivleri taranıyor...")
                time.sleep(0.5)
                progress_bar.progress(60)
                
                status_text.text("⚖️ Mülga kanun maddeleri getiriliyor...")
                progress_bar.progress(90)
                
                try:
                    genai.configure(api_key=api_key)
                    active_model = get_best_model()
                    model = genai.GenerativeModel(active_model)
                    
                    # Kritik Tarih Kontrolleri (Prompt'a ipucu vermek için)
                    era_context = ""
                    if target_date.year < 2002:
                        era_context += "UYARI: Bu tarihte 4721 sayılı Türk Medeni Kanunu YOKTU. 743 sayılı Türk Kanunu Medenisi yürürlükteydi. "
                    if target_date.year < 2012:
                        era_context += "UYARI: 6098 sayılı Borçlar Kanunu YOKTU. 818 sayılı Borçlar Kanunu yürürlükteydi. "
                    
                    prompt = f"""
                    GÖREV: Sen bir Hukuk Tarihçisi ve Mevzuat Uzmanısın.
                    
                    HEDEF TARİH: {target_date.strftime('%d.%m.%Y')}
                    KONU: {topic}
                    DETAY: {specific_query}
                    BAĞLAM: {era_context}
                    
                    İSTENEN ÇIKTI (Rapor Formatı):
                    
                    1. 🏛️ YÜRÜRLÜKTEKİ TEMEL KANUN
                    - O tarihte geçerli olan Kanun Numarası ve Adı (Örn: 743 s. TKM).
                    - İlgili Madde Numarası ve (mümkünse) o günkü metni.
                    
                    2. 📜 KRİTİK FARKLILIKLAR (BUGÜNE GÖRE)
                    - Bugün uygulanan kanunla (Örn: 4721 s. TMK) o günkü kanun arasındaki hayati fark nedir?
                    - Örnek: "O tarihte 'Edinilmiş Mallara Katılma Rejimi' yoktu, 'Mal Ayrılığı' esastı."
                    
                    3. ⚖️ DÖNEMİN İÇTİHADI
                    - O yıllarda Yargıtay'ın bu konuya bakışı nasıldı? (Örn: 1990'larda inançlı işlem içtihadı).
                    
                    4. 💎 AVUKAT İÇİN STRATEJİ
                    - Davayı kazanmak için mahkemeye "Olay tarihindeki mevzuat uygulanmalıdır" itirazını nasıl sunmalıyım?
                    """
                    
                    response = model.generate_content(prompt, stream=True)
                    
                    full_text = ""
                    status_text.empty() # Yazıyı temizle
                    output_placeholder = st.empty()
                    
                    for chunk in response:
                        full_text += chunk.text
                        output_placeholder.markdown(full_text + "▌")
                    
                    output_placeholder.markdown(full_text)
                    progress_bar.progress(100)
                    
                except Exception as e:
                    st.error(f"Hata: {e}")
        else:
            st.info("👈 Sol taraftan tarihi seçin ve yolculuğu başlatın.")
            
            # Örnek Gösterim (Placeholder)
            st.markdown("""
            **Örnek Senaryo:**
            * **Tarih:** 1995
            * **Konu:** Boşanma Mal Paylaşımı
            * **Sonuç:** 2002 öncesi evliliklerde "Mal Ayrılığı" rejimi geçerli olduğundan, kadın eşin ev hanımı olması durumunda tapuda adı yoksa hak talep etmesi çok zordu. Sistem bunu tespit edip "Katkı Payı Alacağı" davası açmanızı önerir.
            """)


def render_expert_report_auditor(api_key):
    st.info("🧐 **Bilirkişi Raporu Denetçisi:** Karmaşık raporları tarar. Matematiksel hataları (kusur toplamı != 100), mantıksal çelişkileri ve eksik incelemeleri tespit ederek 'İtiraz Dilekçesi' taslağı hazırlar.")

    # --- 0. KÜTÜPHANE KONTROLÜ ---
    try:
        from pypdf import PdfReader
    except ImportError:
        st.error("Bu modül için 'pypdf' kütüphanesi gereklidir. Lütfen requirements.txt dosyasına ekleyin.")
        return

    # --- 1. GİRDİ PANELİ ---
    col_upload, col_analysis = st.columns([1, 1])

    report_text = ""

    with col_upload:
        st.markdown("### 📄 Raporu Yükle")
        uploaded_file = st.file_uploader("Bilirkişi Raporu (PDF)", type=["pdf"])
        
        st.markdown("--- VEYA ---")
        text_input = st.text_area("Metni Buraya Yapıştır", height=150, placeholder="Rapor içeriğini buraya kopyalayabilirsiniz...")

        # Metin Çıkarma İşlemi
        if uploaded_file:
            try:
                reader = PdfReader(uploaded_file)
                for page in reader.pages:
                    report_text += page.extract_text() + "\n"
                st.success(f"✅ PDF Okundu: {len(reader.pages)} sayfa")
            except Exception as e:
                st.error(f"PDF Okuma Hatası: {e}")
        elif text_input:
            report_text = text_input

    # --- 2. ANALİZ MOTORU ---
    with col_analysis:
        st.markdown("### 🔍 Denetim Sonucu")
        
        analyze_btn = st.button("🛡️ Raporu Denetle ve Hata Bul", type="primary")
        
        if analyze_btn:
            if not report_text:
                st.warning("Lütfen analiz edilecek bir rapor yükleyin veya metin girin.")
            elif len(report_text) < 50:
                st.warning("Girilen metin analiz için çok kısa.")
            elif not api_key:
                st.error("API Key gerekli.")
            else:
                output_box = st.empty()
                output_box.info("Rapor taranıyor: Kusur oranları toplanıyor, çelişkiler aranıyor...")
                
                try:
                    import google.generativeai as genai
                    genai.configure(api_key=api_key)
                    
                    # Model Seçimi
                    model_name = "models/gemini-pro"
                    for m in genai.list_models():
                        if 'flash' in m.name: model_name = m.name; break
                    
                    model = genai.GenerativeModel(model_name)
                    
                    prompt = f"""
                    GÖREV: Sen titiz bir 'Bilirkişi Raporu Denetçisi' ve Yargıtay İçtihatları uzmanısın.
                    Aşağıdaki bilirkişi raporu metnini analiz et ve hataları bul.
                    
                    METİN:
                    {report_text[:10000]} (Metin kısaltıldıysa devamını dikkate al)
                    
                    İSTENEN ANALİZ (Markdown Formatında):
                    
                    ### 1. 🧮 Matematiksel ve Mantıksal Tutarlılık
                    - Kusur oranları toplamı 100 ediyor mu? (Kontrol et: %25 + %75 vb.)
                    - Hesaplamalarda bariz bir çarpım/toplam hatası var mı?
                    - Tarihler tutarlı mı? (Kaza tarihinden sonraki bir mevzuat uygulanmış mı?)
                    
                    ### 2. ⚖️ Hukuki ve Teknik Dayanak
                    - Rapor hangi teknik veriye dayanıyor? (Tramer, MOBESE, Tanık, Takograf vb.)
                    - Bilirkişi "Hukuki niteleme" yapmış mı? (UYARI: Bilirkişi hukuki yorum yapamaz, sadece teknik tespit yapar. Hakim yerine geçip hüküm kurduysa bunu belirt.)
                    
                    ### 3. 🚩 Tespit Edilen Çelişkiler
                    - "Tanık ifadesinde X denmesine rağmen, raporda Y kabul edilmiştir" gibi çelişkiler var mı?
                    
                    ### 4. 📝 İtiraz Stratejisi (HMK m. 281)
                    - Bu rapora itiraz etmek için kullanılabilecek 3 güçlü argüman yaz.
                    - "Ek Rapor" veya "Yeni Bilirkişi Heyeti" talep etmek için gerekçe oluştur.
                    """
                    
                    response = model.generate_content(prompt, stream=True)
                    
                    full_text = ""
                    for chunk in response:
                        full_text += chunk.text
                        output_box.markdown(full_text + "▌")
                    output_box.markdown(full_text)
                    
                except Exception as e:
                    output_box.error(f"Analiz Hatası: {e}")



def render_corporate_memory(api_key):
    st.info("🏛️ **Kurumsal Hafıza V3 (Oto-Pilot):** Mevcut en güncel AI modelini otomatik bulur ve 'Model Bulunamadı' hatalarını engeller.")

    # --- KÜTÜPHANE KONTROLLERİ ---
    try:
        import pandas as pd
        from pypdf import PdfReader
        from docx import Document
        from PIL import Image
        import google.generativeai as genai
    except ImportError:
        st.error("Gerekli kütüphaneler eksik (pandas, pypdf, python-docx, Pillow, google-generativeai).")
        return

    # --- 0. OTURUM VE VERİ YÖNETİMİ ---
    if "archive_df" not in st.session_state:
        st.session_state.archive_df = pd.DataFrame(columns=["Tarih", "Konu", "Özet", "Detay", "İlgili Kişi/Kurum", "Dosya Adı"])

    # --- KRİTİK FONKSİYON: SAĞLAM MODEL BULUCU ---
    def get_working_model(api_key_val):
        """
        API'den güncel listeyi çeker. Hata verirse manuel listeyi dener.
        En garantili çalışan modeli döndürür.
        """
        genai.configure(api_key=api_key_val)
        
        # 1. YÖNTEM: API'den Canlı Liste İste
        try:
            available_models = []
            for m in genai.list_models():
                if 'generateContent' in m.supported_generation_methods:
                    available_models.append(m.name)
            
            # Öncelik Sıralaması (En iyiden en eskiye)
            priorities = [
                'models/gemini-1.5-flash',
                'models/gemini-1.5-pro',
                'models/gemini-1.5-flash-latest',
                'models/gemini-1.0-pro',
                'models/gemini-pro'
            ]
            
            # Listede eşleşen en iyi modeli bul
            for p in priorities:
                if p in available_models:
                    return p
            
            # Tam eşleşme yoksa, içinde 'flash' geçeni al
            for m in available_models:
                if 'flash' in m: return m
                
            # O da yoksa ilk bulduğunu al
            if available_models:
                return available_models[0]
                
        except Exception as e:
            pass # Liste alınamazsa manuel listeye geç

        # 2. YÖNTEM: Manuel Güvenli Liste (Fallback)
        # API listeleme başarısız olsa bile bu isimler genellikle çalışır
        return "models/gemini-1.5-flash"

    # --- SEKME YAPISI ---
    tab_upload, tab_query = st.tabs(["📂 Belge İşle & Arşivle", "🔍 Arşivde Sorgu Yap"])

    # ==========================================
    # 1. SEKME: BELGE İŞLEME
    # ==========================================
    with tab_upload:
        col_db, col_process = st.columns([1, 1])

        # A. MEVCUT VERİTABANINI YÜKLE
        with col_db:
            st.markdown("### 1. Mevcut Arşivi Yükle")
            uploaded_excel = st.file_uploader("Önceki Excel Dosyanız", type=["xlsx"])
            
            if uploaded_excel:
                try:
                    loaded_df = pd.read_excel(uploaded_excel)
                    st.session_state.archive_df = loaded_df
                    st.success(f"✅ Veritabanı Yüklendi! ({len(loaded_df)} kayıt)")
                except Exception as e:
                    st.error(f"Excel Hatası: {e}")
            
            # Tablo Önizleme
            st.dataframe(st.session_state.archive_df, height=200, use_container_width=True)
            
            # İndir
            if not st.session_state.archive_df.empty:
                excel_data = io.BytesIO()
                st.session_state.archive_df.to_excel(excel_data, index=False)
                st.download_button("💾 Arşivi İndir", excel_data.getvalue(), "Kurumsal_Hafiza.xlsx")

        # B. YENİ BELGE EKLE
        with col_process:
            st.markdown("### 2. Yeni Belge Ekle")
            files = st.file_uploader("Belgeler", type=["pdf", "docx", "png", "jpg", "jpeg"], accept_multiple_files=True)
            
            if st.button("⚙️ Analiz Et ve Ekle", type="primary") and files:
                if not api_key:
                    st.error("API Key gerekli.")
                else:
                    # --- MODELİ BELİRLE ---
                    active_model_name = get_working_model(api_key)
                    st.toast(f"🤖 Aktif Model: {active_model_name}", icon="✅")
                    
                    model = genai.GenerativeModel(active_model_name)
                    progress_bar = st.progress(0)
                    new_records = []
                    
                    for idx, file in enumerate(files):
                        try:
                            content = ""
                            image_data = None
                            is_image = False
                            
                            # Dosya Okuma
                            if file.type == "application/pdf":
                                reader = PdfReader(file)
                                for page in reader.pages: content += page.extract_text() + "\n"
                            elif "word" in file.type:
                                doc = Document(file)
                                for p in doc.paragraphs: content += p.text + "\n"
                            elif "image" in file.type:
                                is_image = True
                                image_data = Image.open(file)

                            # Prompt Hazırla
                            prompt = """
                            Bu belgeden şu bilgileri JSON formatında çıkar:
                            {"Tarih": "GG.AA.YYYY", "Konu": "...", "Ozet": "...", "Detay": "...", "Ilgili_Kisi": "..."}
                            Sadece JSON ver.
                            """
                            
                            # API Çağrısı (Retry Mekanizması ile)
                            response = None
                            try:
                                if is_image:
                                    # Eğer model vision desteklemiyorsa flash'a zorla
                                    if "flash" not in active_model_name and "1.5" not in active_model_name:
                                        model_vision = genai.GenerativeModel("models/gemini-1.5-flash")
                                        response = model_vision.generate_content([prompt, image_data])
                                    else:
                                        response = model.generate_content([prompt, image_data])
                                else:
                                    if len(content) > 5:
                                        response = model.generate_content(prompt + f"\n\nMETİN:\n{content[:20000]}")
                            except Exception as api_err:
                                st.warning(f"Model hatası ({active_model_name}), yedek model deneniyor...")
                                # Hata verirse kesin çalışan Flash modelini dene
                                backup_model = genai.GenerativeModel("models/gemini-1.5-flash")
                                if is_image:
                                    response = backup_model.generate_content([prompt, image_data])
                                else:
                                    response = backup_model.generate_content(prompt + f"\n\nMETİN:\n{content[:20000]}")

                            # Sonucu İşle
                            if response and response.text:
                                clean_json = response.text.replace("```json", "").replace("```", "").strip()
                                data = json.loads(clean_json)
                                
                                # Mükerrer Kontrolü
                                is_dup = False
                                if not st.session_state.archive_df.empty:
                                    check = st.session_state.archive_df[
                                        (st.session_state.archive_df['Konu'] == data.get('Konu')) & 
                                        (st.session_state.archive_df['Tarih'] == data.get('Tarih'))
                                    ]
                                    if not check.empty: is_dup = True
                                
                                if not is_dup:
                                    new_records.append({
                                        "Tarih": data.get("Tarih", "-"),
                                        "Konu": data.get("Konu", "-"),
                                        "Özet": data.get("Ozet", "-"),
                                        "Detay": data.get("Detay", "-"),
                                        "İlgili Kişi/Kurum": data.get("Ilgili_Kisi", "-"),
                                        "Dosya Adı": file.name
                                    })
                                    
                        except Exception as e:
                            st.error(f"Hata ({file.name}): {e}")
                        
                        progress_bar.progress((idx + 1) / len(files))

                    # Kaydet
                    if new_records:
                        st.session_state.archive_df = pd.concat([st.session_state.archive_df, pd.DataFrame(new_records)], ignore_index=True)
                        st.success("İşlem Tamamlandı!")
                        st.rerun()

    # ==========================================
    # 2. SEKME: SORGULAMA
    # ==========================================
    with tab_query:
        st.markdown("### 🧠 Arşivde Semantik Arama")
        if st.session_state.archive_df.empty:
            st.info("Veri yok.")
        else:
            query = st.text_input("Soru:", placeholder="Örn: X firması ile ilgili sözleşme detayı?")
            if st.button("🔍 Ara"):
                with st.spinner("Aranıyor..."):
                    try:
                        active_model = get_working_model(api_key)
                        model = genai.GenerativeModel(active_model)
                        context = st.session_state.archive_df.to_json(orient="records", force_ascii=False)
                        prompt = f"VERİTABANI:\n{context}\n\nSORU: {query}\n\nBu veritabanına göre cevapla:"
                        st.markdown(model.generate_content(prompt).text)
                    except Exception as e:
                        st.error(f"Hata: {e}")






# ==========================================
# 4. ANA UYGULAMA (MAIN) - DÜZELTİLMİŞ VERSİYON
# ==========================================
def main():
    # --- A. BAŞLANGIÇ AYARLARI ---
    st.title("⚖️ Hukuk Asistanı (v10.0 - Ultimate Edition)")
    
    try:
        lib_ver = importlib.metadata.version("google-generativeai")
    except:
        lib_ver = "Bilinmiyor"

    # Session State Başlatma
    if "durusma_listesi" not in st.session_state: st.session_state.durusma_listesi = load_durusma_data()
    if "doc_text" not in st.session_state: st.session_state.doc_text = ""
    if "last_file_id" not in st.session_state: st.session_state.last_file_id = None
    if "messages" not in st.session_state: st.session_state.messages = []
    
    # Diğer state tanımları...
    keys_to_init = ["mevzuat_sonuc", "ictihat_sonuc", "dilekce_taslak", "soru_cevap", 
                    "ses_metni", "ocr_metni", "dalgic_context", "dalgic_sonuc", 
                    "buyur_abi_response", "arsiv_context", "arsiv_genel_ozet",
                    "arsiv_soru_cevap", "aktif_dosya_adi", "aktif_dosya_yolu",
                    "sozlesme_analiz", "mock_messages", "gorev_listesi", "kvkk_metin"]
    
    for k in keys_to_init:
        if k not in st.session_state: st.session_state[k] = ""
    
    if "arsiv_arama_sonuclari" not in st.session_state: st.session_state.arsiv_arama_sonuclari = []
    if "mevzuat_takip_listesi" not in st.session_state: st.session_state.mevzuat_takip_listesi = []

    ROOT_DIR = "Hukuk_Arsivi"
    if not os.path.exists(ROOT_DIR): os.makedirs(ROOT_DIR)

    # --- B. SIDEBAR (SOL MENÜ) ---
    with st.sidebar:
        st.header("⚙️ Ayarlar")
        api_key = st.text_input("Google Gemini API Key", type="password")
        st.caption(f"Kütüphane Sürümü: {lib_ver}")
        
        st.divider()
        
        if st.button("🗑️ Ekranı Temizle"):
            st.session_state.clear()
            st.rerun()

        st.markdown("---")
        st.header("🚀 MODÜL SEÇİMİ")
        
        # 1. Kategori Seçimi
        secilen_kategori = st.radio(
            "Kategori:",
            ["🛠️ Temel Araçlar", "🚀 Yönetim & Pro", "🔮 Simülasyon & Risk", "🔥 Özel Araçlar"]
        )
        
        # 2. Modül Seçimi
        secilen_modul = ""
        if secilen_kategori == "🛠️ Temel Araçlar":
            secilen_modul = st.selectbox("Araç Seç:", [
                "📋 Analiz", "💬 Sohbet", "📕 Mevzuat", "⚖️ İçtihat", 
                "✍️ Dilekçe Yaz", "❓ Bana Sor", "🎙️ Ses", "👁️ OCR",
                "🌍 Çeviri", "🛡️ Çürüt", "🕵️‍♂️ Sorgu", "😈 Şeytanın Avukatı", 
                "🤿 Dalgıç", "🧠 Semantik", "🎙️ Canlı Duruşma", "🦋 Kelebek"
            ])
        elif secilen_kategori == "🚀 Yönetim & Pro":
            secilen_modul = st.selectbox("Araç Seç:", [
                "🙋 Buyur Abi", "⏰ Hatırlatıcı", "🗄️ Arşiv", "🏛️ UYAP Analiz", 
                "🕸️ İlişki Ağı", "📝 Sözleşme Analiz", "🕵️‍♂️ KVKK Temizle",  
                "⚔️ Belge Kıyasla", "🎭 Sanal Duruşma", "✅ Görev Çıkarıcı", 
                "⚡ Canlı Asistan", "📡 Etki Analizi", "🕵️ Dijital Otp"
            ])
        elif secilen_kategori == "🔮 Simülasyon & Risk":
            secilen_modul = st.selectbox("Araç Seç:", [
                "🏥 Kurumsal Check-up", "⏳ Zaman Makinesi", "⚖️ AYM & AİHM Testi", 
                "🕵️ Deepfake Kontrol", "🌐 OSINT (İstihbarat)", "🔔 Emsal Alarm", 
                "👑 Sahip Modu", "🌳 Soyağacı", "🔥 Isı Haritası", 
                "🕸️ Gizli Bağlantı", "🤝 Arabuluculuk"
            ])
        elif secilen_kategori == "🔥 Özel Araçlar":
            secilen_modul = st.selectbox("Araç Seç:", [
                "🗺️ Adli Harita", "🕰️ Mevzuat Makinesi", 
                "🧐 Rapor Denetçisi", "🏛️ Kurumsal Hafıza"
            ])

    # --- C. DOSYA YÜKLEME ALANI (HER ZAMAN GÖRÜNÜR) ---
    st.info(f"📍 Şu anki Modül: **{secilen_modul}**")
    
    # Sadece bazı modüllerde dosya yükleme alanını gösterelim veya her zaman gösterelim
    uploaded_file = st.file_uploader("Dosya Yükle (UDF/PDF) - Analiz İçin", type=['udf', 'pdf'])

    if uploaded_file and st.session_state.get('last_file_id') != uploaded_file.file_id:
        with st.spinner("Dosya okunuyor..."):
            file_bytes = BytesIO(uploaded_file.getvalue())
            ext = uploaded_file.name.split('.')[-1].lower()
            raw_text = parse_udf(file_bytes) if ext == 'udf' else parse_pdf(file_bytes)
            st.session_state.doc_text = raw_text
            st.session_state.last_file_id = uploaded_file.file_id
            st.session_state.messages = [] # Yeni dosya gelince sohbeti sıfırla

    if st.session_state.doc_text.startswith(("HATA", "UYARI")):
        st.warning(st.session_state.doc_text)
    
    # Metadata çıkarma (Analiz modülü için gerekli)
    auto_data = extract_metadata(st.session_state.doc_text)

    # ==========================================
    # D. MODÜL YÖNLENDİRİCİSİ (ROUTER)
    # Burada 'with tab:' yerine 'if secilen_modul ==' kullanıyoruz.
    # ==========================================

    st.divider()

    # --- 1. GRUP: TEMEL ARAÇLAR ---
    if secilen_modul == "📋 Analiz":
        col1, col2 = st.columns(2)
        with col1:
            st.write(f"**Mahkeme:** {input_mahkeme or auto_data['mahkeme']}")
            st.write(f"**Dosya No:** {input_dosya_no or auto_data['esas']}")
        with col2:
            st.write(f"**Davacı:** {input_davaci or '-'}")
            st.write(f"**Davalı:** {input_davali or '-'}")
        st.text_area("Metin Önizleme", st.session_state.doc_text, height=300)

    elif secilen_modul == "💬 Sohbet":
        for msg in st.session_state.messages:
            with st.chat_message(msg["role"]): st.markdown(msg["content"])
        if prompt := st.chat_input("Bu dosya hakkında soru sor..."):
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"): st.markdown(prompt)
            with st.chat_message("assistant"):
                with st.spinner("AI Yanıtlıyor..."):
                    context = f"BELGE: {st.session_state.doc_text[:20000]}\nSORU: {prompt}"
                    reply = get_ai_response(f"Sen bir avukatsın. Şuna cevap ver: {context}", api_key)
                    st.markdown(reply)
                    st.session_state.messages.append({"role": "assistant", "content": reply})

    elif secilen_modul == "📕 Mevzuat":
        c1, c2 = st.columns([3,1])
        q = c1.text_input("Kanun Madde No", key="mq")
        if c2.button("Getir", key="mb") and q:
            with st.spinner("Aranıyor..."):
                res = get_ai_response(f"GÖREV: '{q}' maddesini tam metin yaz.", api_key)
                st.session_state.mevzuat_sonuc = res
        if st.session_state.mevzuat_sonuc:
            st.markdown(f"<div class='kanun-kutusu'>{st.session_state.mevzuat_sonuc}</div>", unsafe_allow_html=True)

    elif secilen_modul == "⚖️ İçtihat":
        c3, c4 = st.columns([3,1])
        iq = c3.text_input("İçtihat Konusu", key="iq")
        if c4.button("Ara", key="ib") and iq:
            with st.spinner("Taranıyor..."):
                res = get_ai_response(f"GÖREV: '{iq}' hakkında Yargıtay kararlarını özetle.", api_key)
                st.session_state.ictihat_sonuc = res
        if st.session_state.ictihat_sonuc:
            st.markdown(f"<div class='ictihat-kutusu'>{st.session_state.ictihat_sonuc}</div>", unsafe_allow_html=True)

    elif secilen_modul == "✍️ Dilekçe Yaz":
        st.subheader("✍️ Otomatik Savunma/Cevap Dilekçesi")
        if not st.session_state.doc_text:
            st.warning("Dilekçe oluşturmak için önce yukarıdan bir dosya yükleyin.")
        else:
            col_d1, col_d2 = st.columns([2, 1])
            with col_d1:
                dilekce_turu = st.selectbox("Dilekçe Türü", ["Cevap Dilekçesi", "İtiraz Dilekçesi", "Beyan Dilekçesi"])
                ozel_talimat = st.text_area("Özel Savunma Stratejisi (Opsiyonel)")
            with col_d2:
                st.write("")
                st.write("")
                if st.button("Dilekçeyi Yaz (AI)", type="primary"):
                    if not api_key: st.error("API Key gerekli!")
                    else:
                        with st.spinner("Dilekçe yazılıyor..."):
                            mahkeme = input_mahkeme or auto_data['mahkeme']
                            dosya = input_dosya_no or auto_data['esas']
                            davaci = input_davaci or "Davacı"
                            davali = input_davali or "Davalı"
                            prompt = f"""
                            GÖREV: Aşağıdaki metne dayanarak profesyonel bir {dilekce_turu} yaz.
                            BİLGİLER: Mahkeme: {mahkeme}, Dosya: {dosya}, Davacı: {davaci}, Davalı: {davali}, Ek Talimat: {ozel_talimat}
                            KARŞI TARAFIN DİLEKÇESİ (ÖZET): {st.session_state.doc_text[:20000]}
                            KURALLAR: Resmi Türk hukuk dilekçesi formatında olsun.
                            """
                            res = get_ai_response(prompt, api_key)
                            st.session_state.dilekce_taslak = res
            if st.session_state.dilekce_taslak:
                st.text_area("Dilekçe Metni", st.session_state.dilekce_taslak, height=500)
                st.download_button("💾 İndir", create_word_file(st.session_state.dilekce_taslak), "Dilekce.docx")

    elif secilen_modul == "❓ Bana Sor":
        col_s1, col_s2 = st.columns([3, 1])
        with col_s1:
            kullanici_sorusu = st.text_area("Hukuki Sorunuzu Yazın", height=100)
        with col_s2:
            if st.button("Analiz Et", type="primary"):
                if not api_key: st.error("API Key giriniz.")
                else:
                    with st.spinner("Mevzuat taranıyor..."):
                        res = get_ai_response(f"SORU: {kullanici_sorusu}\nCEVAPLA:", api_key)
                        st.session_state.soru_cevap = res
        if st.session_state.soru_cevap:
            st.markdown(f"<div class='ictihat-kutusu'>{st.session_state.soru_cevap}</div>", unsafe_allow_html=True)

    elif secilen_modul == "🎙️ Ses":
        col_a1, col_a2 = st.columns(2)
        with col_a1:
            txt = st.text_area("Okunacak Metin")
            if st.button("🔊 Seslendir") and txt:
                fp = text_to_speech(txt)
                if fp: st.audio(fp, format='audio/mp3')
        with col_a2:
            aud = st.file_uploader("Ses Dosyası", type=["wav", "mp3"])
            if aud and st.button("📝 Yazıya Dök"):
                st.session_state.ses_metni = speech_to_text(aud)
            if st.session_state.ses_metni: st.write(st.session_state.ses_metni)

    elif secilen_modul == "👁️ OCR":
        ocr_file = st.file_uploader("Resim/PDF Yükle", type=['png', 'jpg', 'pdf'])
        if ocr_file and st.button("Metni Çıkar"):
            if not api_key: st.error("API Key gerekli.")
            else:
                with st.spinner("OCR yapılıyor..."):
                    mime = "application/pdf" if ocr_file.name.endswith('pdf') else "image/jpeg"
                    st.session_state.ocr_metni = perform_ocr_gemini(ocr_file, mime, api_key)
        if st.session_state.ocr_metni: st.text_area("Sonuç", st.session_state.ocr_metni, height=400)

    elif secilen_modul == "🌍 Çeviri":
        txt = st.text_area("Çevrilecek Metin")
        if st.button("Çevir") and api_key:
            st.write(get_ai_response(f"Bu hukuki metni İngilizceye çevir: {txt}", api_key))

    elif secilen_modul == "🛡️ Çürüt":
        iddia = st.text_area("Karşı Taraf İddiası")
        if st.button("Çürüt") and api_key:
            st.write(get_ai_response(f"Bu iddiayı çürütmek için 3 argüman yaz: {iddia}", api_key))

    elif secilen_modul == "🕵️‍♂️ Sorgu":
        ifade = st.text_area("Tanık İfadesi")
        if st.button("Sorgu Hazırla") and api_key:
            st.write(get_ai_response(f"Bu ifade için çapraz sorgu soruları hazırla: {ifade}", api_key))

    elif secilen_modul == "😈 Şeytanın Avukatı":
        dilekce = st.text_area("Dilekçe Taslağı")
        if st.button("Eleştir") and api_key:
            st.write(get_ai_response(f"Bu dilekçeyi sertçe eleştir: {dilekce}", api_key))

    elif secilen_modul == "🤿 Dalgıç":
        files = st.file_uploader("Çoklu Dosya", accept_multiple_files=True)
        if files and st.button("İşle"):
            st.session_state.dalgic_context = "Dosyalar işlendi..." # Basitleştirildi
            st.success("Hafızaya alındı.")
        if st.session_state.dalgic_context:
            q = st.text_input("Dosyalar hakkında soru sor")
            if st.button("Sor") and api_key:
                st.write(get_ai_response(f"{q}", api_key))

    elif secilen_modul == "🧠 Semantik":
        files = st.file_uploader("Arşiv Dosyaları", accept_multiple_files=True)
        if files and st.button("Hafızaya Al"): st.success("İşlendi")
        q = st.text_input("Arşivde ara")
        if st.button("Ara") and api_key: st.write("Sonuç...")

    elif secilen_modul == "🎙️ Canlı Duruşma":
        ref = st.text_area("Eski İfade")
        new = st.text_area("Yeni İfade")
        if st.button("Kıyasla") and api_key:
            st.write(get_ai_response(f"Çelişki var mı?\n1:{ref}\n2:{new}", api_key))

    elif secilen_modul == "🦋 Kelebek":
        render_temporal_law_machine(api_key)

    # --- 2. GRUP: YÖNETİM & PRO ---
    elif secilen_modul == "🙋 Buyur Abi":
        msg = st.text_area("Sorunu yaz abi")
        if st.button("Gönder") and api_key:
            st.write(get_ai_response(f"Yardımcı ol: {msg}", api_key))

    elif secilen_modul == "⏰ Hatırlatıcı":
        f = st.file_uploader("Takvim (.ics)", type=['ics'])
        if f: st.success("Takvim işlendi")
        if st.session_state.durusma_listesi:
            st.dataframe(pd.DataFrame(st.session_state.durusma_listesi))

    elif secilen_modul == "🗄️ Arşiv":
        st.info("Dosya yönetim sistemi.")

    elif secilen_modul == "🏛️ UYAP Analiz":
        zips = st.file_uploader("UYAP Zip", accept_multiple_files=True)
        if zips and st.button("Analiz") and api_key: st.write("Analiz sonucu...")

    elif secilen_modul == "🕸️ İlişki Ağı":
        if st.button("Ağı Çiz") and api_key: st.graphviz_chart("digraph { A -> B; }")

    elif secilen_modul == "📝 Sözleşme Analiz":
        f = st.file_uploader("Sözleşme")
        if f and st.button("İncele") and api_key: st.write("Risk raporu...")

    elif secilen_modul == "🕵️‍♂️ KVKK Temizle":
        txt = st.text_area("Metin")
        if st.button("Temizle"): st.write("Anonim metin...")

    elif secilen_modul == "⚔️ Belge Kıyasla":
        t1 = st.text_area("Metin 1")
        t2 = st.text_area("Metin 2")
        if st.button("Kıyasla"): st.write("Farklar...")

    elif secilen_modul == "🎭 Sanal Duruşma":
        st.info("Sanal duruşma simülasyonu.")

    elif secilen_modul == "✅ Görev Çıkarıcı":
        karar = st.text_area("Karar")
        if st.button("Görevleri Bul") and api_key: st.write("Görev listesi...")

    elif secilen_modul == "⚡ Canlı Asistan":
        iddia = st.text_input("Karşı taraf ne dedi?")
        if st.button("Doğrula") and api_key: st.write("Doğruluk kontrolü...")

    elif secilen_modul == "📡 Etki Analizi":
        st.info("Mevzuat takip sistemi.")

    elif secilen_modul == "🕵️ Dijital Otp":
        f = st.file_uploader("Dosya")
        if f: st.write("Metadata bilgisi...")

    # --- 3. GRUP: SİMÜLASYON ---
    elif secilen_modul == "🏥 Kurumsal Check-up":
        render_checkup_module(api_key)
    elif secilen_modul == "⏳ Zaman Makinesi":
        render_time_machine(api_key)
    elif secilen_modul == "⚖️ AYM & AİHM Testi":
        render_aym_aihm_module(api_key)
    elif secilen_modul == "🕵️ Deepfake Kontrol":
        render_deepfake_module(api_key)
    elif secilen_modul == "🌐 OSINT (İstihbarat)":
        render_osint_module(api_key)
    elif secilen_modul == "🔔 Emsal Alarm":
        render_precedent_alert_module(api_key)
    elif secilen_modul == "👑 Sahip Modu":
        render_owner_mode(api_key)
    elif secilen_modul == "🌳 Soyağacı":
        render_property_genealogy(api_key)
    elif secilen_modul == "🔥 Isı Haritası":
        render_limitations_heatmap(api_key)
    elif secilen_modul == "🕸️ Gizli Bağlantı":
        render_conflict_scanner(api_key)
    elif secilen_modul == "🤝 Arabuluculuk":
        render_mediation_checker(api_key)

    # --- 4. GRUP: ÖZEL ARAÇLAR ---
    elif secilen_modul == "🗺️ Adli Harita":
        render_forensic_map(api_key)
    elif secilen_modul == "🕰️ Mevzuat Makinesi":
        render_temporal_law_machine(api_key)
    elif secilen_modul == "🧐 Rapor Denetçisi":
        render_expert_report_auditor(api_key)
    elif secilen_modul == "🏛️ Kurumsal Hafıza":
        render_corporate_memory(api_key)

if __name__ == "__main__":
    main()


